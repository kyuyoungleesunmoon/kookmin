from docx import Document
from docx.shared import Pt
import re

input_path = r"C:\국민대프로젝트\논문작성\IEEE_DPF_Paper_Final_v5.docx"
output_path = r"C:\국민대프로젝트\논문작성\IEEE_DPF_Paper_Clean_v6.docx"

doc = Document(input_path)

# Patterns to identify code blocks
code_patterns = [
    r'^\s*(def |class |import |from |if __name__|#!/)',
    r'^\s*#\s*\d*\.\s*',  # Comment with numbering
    r'^\s*\w+\s*=\s*\{$',  # Dict start
    r'^\s*\w+\s*=\s*\[$',  # List start
    r"^\s*'[\w_]+'\s*:\s*",  # Dict key-value
    r'^\s*plt\.',  # matplotlib
    r'^\s*np\.',  # numpy
    r'^\s*torch\.',  # pytorch
    r'^\s*model\s*=',  # model assignment
    r'^\s*return\s+',  # return statements
    r'^\s*print\s*\(',  # print calls
    r'^\s*for\s+\w+\s+in\s+',  # for loops
    r'^\s*while\s+',  # while loops
    r'^\s*try\s*:',  # try blocks
    r'^\s*except\s*',  # except blocks
    r'^\s*with\s+open',  # file operations
    r'^\s*df\s*=',  # dataframe
    r'^\s*\w+_config\s*=\s*\{',  # config dicts
    r'^\s*augmentation\s*=\s*\{',  # augmentation config
    r'^\s*logs\s*=\s*\{',  # logs dict
]

# Tree structure patterns (├─ └─ │)
tree_patterns = [
    r'^\s*├─',
    r'^\s*└─',
    r'^\s*│\s+├─',
    r'^\s*│\s+└─',
]

# Essential equations to KEEP (by content keywords)
essential_equation_keywords = [
    'mAP50',  # Main performance metric
    'd_{\\mathcal{A}}',  # Domain distance (Proxy A-distance)
    '\\mathcal{L}_{YOLO}',  # Main loss function
    'IoU',  # Intersection over Union
    'Precision',  # Core metrics
    'Recall',
    'F1',
    'AP_c',  # Average Precision
]

# Equation block patterns
equation_start_patterns = [
    r'\$\$\s*$',  # $$ on its own line
    r'\$\$\s*\\',  # $$ followed by latex
]

equation_content_patterns = [
    r'\\begin\{',
    r'\\end\{',
    r'\\tag\{',
    r'\\frac\{',
    r'\\sum',
    r'\\mathcal\{',
    r'\\left',
    r'\\right',
    r'\\eta',
    r'\\lambda',
    r'\\epsilon',
]

def is_code_block(text):
    """Check if text looks like source code"""
    for pattern in code_patterns:
        if re.search(pattern, text):
            return True
    return False

def is_tree_structure(text):
    """Check if text is ASCII tree structure"""
    for pattern in tree_patterns:
        if re.search(pattern, text):
            return True
    return False

def is_equation_block(text):
    """Check if text is part of an equation block"""
    if text.strip() == '$$':
        return True
    for pattern in equation_content_patterns:
        if re.search(pattern, text):
            return True
    return False

def is_essential_equation(text):
    """Check if equation contains essential keywords"""
    for keyword in essential_equation_keywords:
        if keyword in text:
            return True
    return False

# Process document
paragraphs_to_remove = []
in_equation_block = False
equation_buffer = []
current_equation_essential = False

print("Processing paragraphs...")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    if not text:
        continue
    
    # Handle equation blocks
    if text == '$$' or text.startswith('$$'):
        if not in_equation_block:
            # Starting equation block
            in_equation_block = True
            equation_buffer = [i]
            current_equation_essential = is_essential_equation(text)
        else:
            # Ending equation block
            equation_buffer.append(i)
            if not current_equation_essential:
                paragraphs_to_remove.extend(equation_buffer)
            in_equation_block = False
            equation_buffer = []
            current_equation_essential = False
        continue
    
    if in_equation_block:
        equation_buffer.append(i)
        if is_essential_equation(text):
            current_equation_essential = True
        continue
    
    # Check for standalone equation content (inline $$...$$)
    if is_equation_block(text) and not is_essential_equation(text):
        paragraphs_to_remove.append(i)
        continue
    
    # Check for code blocks
    if is_code_block(text):
        paragraphs_to_remove.append(i)
        continue
    
    # Check for tree structures
    if is_tree_structure(text):
        paragraphs_to_remove.append(i)
        continue

print(f"Paragraphs to remove: {len(paragraphs_to_remove)}")

# Remove paragraphs in reverse order to maintain indices
for idx in sorted(paragraphs_to_remove, reverse=True):
    try:
        para = doc.paragraphs[idx]
        p = para._element
        p.getparent().remove(p)
    except Exception as e:
        print(f"Error removing paragraph {idx}: {e}")

print(f"Saving cleaned document to: {output_path}")
doc.save(output_path)

# Verify
doc2 = Document(output_path)
print(f"Original paragraphs: {len(Document(input_path).paragraphs)}")
print(f"Cleaned paragraphs: {len(doc2.paragraphs)}")
print("Done!")
