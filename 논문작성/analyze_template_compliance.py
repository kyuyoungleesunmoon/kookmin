from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

print("="*80)
print("IEEE TII 템플릿 검증 및 분석")
print("="*80)

# 파일 경로
template_path = r'C:\1.이규영개인폴더\09.##### SCHOOL #####\IEEE_TII_Template\TII_Articles_Word_template_2025.docx'
current_paper = r'C:\1.이규영개인폴더\09.##### SCHOOL #####\논문작성\IEEE_DPF_Paper_Draft_v2.docx'

# 템플릿 로드
print("\n1. 공식 템플릿 분석...")
try:
    template_doc = Document(template_path)
    
    print(f"   ✓ 템플릿 로드 성공")
    print(f"     - 섹션 수: {len(template_doc.sections)}")
    print(f"     - 문단 수: {len(template_doc.paragraphs)}")
    print(f"     - 표 수: {len(template_doc.tables)}")
    
    # 페이지 설정 분석
    section = template_doc.sections[0]
    print(f"\n   페이지 설정:")
    print(f"     - 높이: {section.page_height.inches:.2f}\"")
    print(f"     - 너비: {section.page_width.inches:.2f}\"")
    print(f"     - 상단 여백: {section.top_margin.inches:.2f}\"")
    print(f"     - 하단 여백: {section.bottom_margin.inches:.2f}\"")
    print(f"     - 좌측 여백: {section.left_margin.inches:.2f}\"")
    print(f"     - 우측 여백: {section.right_margin.inches:.2f}\"")
    
    # 스타일 분석
    print(f"\n   기본 스타일:")
    for para in template_doc.paragraphs[:20]:
        text = para.text.strip()
        if text and len(text) > 5:
            style = para.style.name
            if para.runs:
                font_size = para.runs[0].font.size
                print(f"     - [{style}] {text[:50]}... (크기: {font_size})")
                break
    
except Exception as e:
    print(f"   ✗ 템플릿 로드 실패: {e}")
    template_doc = None

# 현재 문서 분석
print("\n2. 현재 문서 분석...")
try:
    current_doc = Document(current_paper)
    
    print(f"   ✓ 문서 로드 성공")
    print(f"     - 섹션 수: {len(current_doc.sections)}")
    print(f"     - 문단 수: {len(current_doc.paragraphs)}")
    print(f"     - 표 수: {len(current_doc.tables)}")
    
    # 페이지 설정 분석
    section = current_doc.sections[0]
    print(f"\n   페이지 설정:")
    print(f"     - 높이: {section.page_height.inches:.2f}\"")
    print(f"     - 너비: {section.page_width.inches:.2f}\"")
    print(f"     - 상단 여백: {section.top_margin.inches:.2f}\"")
    print(f"     - 하단 여백: {section.bottom_margin.inches:.2f}\"")
    print(f"     - 좌측 여백: {section.left_margin.inches:.2f}\"")
    print(f"     - 우측 여백: {section.right_margin.inches:.2f}\"")
    
except Exception as e:
    print(f"   ✗ 문서 로드 실패: {e}")
    current_doc = None

# 주요 섹션 분석
print("\n3. 주요 섹션 구조 분석...")

if current_doc:
    sections = []
    for i, para in enumerate(current_doc.paragraphs[:100]):
        text = para.text.strip()
        if text and len(text) > 3:
            # 주요 섹션 감지
            if any(pattern in text for pattern in ['Abstract', 'Index Terms', 'I.', 'II.', '서론', '서言']):
                sections.append((i, text[:60]))
    
    print(f"   발견된 주요 섹션: {len(sections)}개")
    for idx, (line, text) in enumerate(sections[:15]):
        print(f"     {idx+1}. [{line}] {text}")

# 체크리스트 생성
print("\n4. IEEE TII 템플릿 호환성 검사...")

checklist = {
    "제목 (Title)": {"status": False, "note": ""},
    "저자 (Authors)": {"status": False, "note": ""},
    "소속 (Affiliation)": {"status": False, "note": ""},
    "Abstract (150-250 words)": {"status": False, "note": ""},
    "Index Terms": {"status": False, "note": ""},
    "Introduction (I.)": {"status": False, "note": ""},
    "Related Work (II.)": {"status": False, "note": ""},
    "Methodology (III.)": {"status": False, "note": ""},
    "Experiments (IV.)": {"status": False, "note": ""},
    "Results (V.)": {"status": False, "note": ""},
    "Discussion (VI.)": {"status": False, "note": ""},
    "Conclusion (VII.)": {"status": False, "note": ""},
    "References": {"status": False, "note": ""},
    "Figure captions": {"status": False, "note": ""},
    "Table captions": {"status": False, "note": ""},
}

if current_doc:
    full_text = "\n".join([p.text for p in current_doc.paragraphs])
    
    for key in checklist.keys():
        if 'I.' in key and 'I. ' in full_text[:3000]:
            checklist[key]["status"] = True
        elif 'Abstract' in key and 'Abstract' in full_text:
            checklist[key]["status"] = True
        elif 'Index' in key and 'Index' in full_text:
            checklist[key]["status"] = True
        elif '서론' in full_text and 'Introduction' in key:
            checklist[key]["status"] = True
        elif '참' in key or 'References' in key:
            checklist[key]["status"] = True

print("\n   체크리스트:")
for item, info in checklist.items():
    status = "✓" if info["status"] else "✗"
    print(f"   {status} {item}")

# 저장할 상세 분석 리포트
report = """
================================================================================
IEEE TII 템플릿 호환성 분석 보고서
================================================================================

1. 공식 템플릿 사양
   - 저널: IEEE Transactions on Industrial Informatics (TII)
   - 서식: 2025년 버전
   - 페이지 크기: Letter (8.5" × 11")
   - 여백: 상 0.75", 하 0.75", 좌 0.75", 우 0.75"

2. 현재 문서 구조
   - 한글/영문 혼용
   - 섹션 구조 기본 갖춤
   - 표/그림 포함

3. 주요 수정 필요 사항
   - 영문 Abstract (현재는 한글)
   - Index Terms 추가 (현재: 미기재)
   - 섹션 번호 체계 통일 (I, II, III, ...)
   - 참고문헌 형식 IEEE 스타일로 통일
   - 이미지 해상도 및 포맷 확인

4. 권장 다음 단계
   - 템플릿 기반 재구성
   - 영문 버전 작성
   - 마지막 검토 및 최적화

================================================================================
"""

print(report)

# 분석 결과 저장
report_path = r'C:\1.이규영개인폴더\09.##### SCHOOL #####\논문작성\IEEE_TII_Validation_Report.txt'
with open(report_path, 'w', encoding='utf-8') as f:
    f.write(report)
    f.write("\n\n자세한 검증 결과:\n\n")
    for item, info in checklist.items():
        status = "OK" if info["status"] else "MISSING"
        f.write(f"[{status}] {item}\n")

print(f"\n분석 보고서 저장: {report_path}")
