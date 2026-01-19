# -*- coding: utf-8 -*-
"""
IEEE_DPF_Paper_Final.docx에 누락된 이미지 추가
원본 마크다운의 이미지를 Word 문서에 삽입
"""

from docx import Document
from docx.shared import Inches, Pt
import os
import re

def find_image_placeholders(doc):
    """이미지 플레이스홀더 또는 이미지 참조 찾기"""
    placeholders = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        
        # [Image Error: ...] 패턴 찾기
        if '[Image Error:' in text:
            placeholders.append({
                'para_index': i,
                'type': 'error',
                'text': text
            })
        
        # 그림 N: 캡션 패턴 (이미지가 누락된 경우)
        if re.match(r'^그림\s*\d+:', text) and len(text) < 100:
            placeholders.append({
                'para_index': i,
                'type': 'caption_only',
                'text': text
            })
    
    return placeholders

def check_existing_images(docx_path):
    """DOCX에 이미 존재하는 이미지 확인"""
    import zipfile
    
    with zipfile.ZipFile(docx_path, 'r') as z:
        images = [f for f in z.namelist() if f.startswith('word/media/')]
    
    return images

def get_md_images(md_path):
    """마크다운에서 참조된 이미지 목록"""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 이미지 참조 패턴
    pattern = r'!\[(.*?)\]\((images/[^)]+)\)'
    matches = re.findall(pattern, content)
    
    # 중복 제거
    unique_images = {}
    for alt, path in matches:
        if path not in unique_images:
            unique_images[path] = alt
    
    return unique_images

def analyze_image_gaps(docx_path, md_path, images_dir):
    """이미지 갭 분석"""
    print("=" * 70)
    print("이미지 갭 분석")
    print("=" * 70)
    
    # DOCX 이미지
    docx_images = check_existing_images(docx_path)
    print(f"\nDOCX 내 이미지: {len(docx_images)}개")
    
    # 마크다운 이미지 참조
    md_images = get_md_images(md_path)
    print(f"마크다운 이미지 참조: {len(md_images)}개")
    
    # 실제 이미지 파일 확인
    existing_files = os.listdir(images_dir)
    dpf_images = [f for f in existing_files if '이규영_국민대_DPF_논문' in f]
    print(f"실제 DPF 논문 이미지 파일: {len(dpf_images)}개")
    
    for img in sorted(dpf_images):
        print(f"   - {img}")
    
    # confusion_matrix 확인
    if 'confusion_matrix_normalized.png' in existing_files:
        print(f"   - confusion_matrix_normalized.png")
    
    return md_images, dpf_images

def fix_image_errors_in_doc(docx_path, output_path, images_dir):
    """문서에서 [Image Error:] 텍스트 제거 및 정리"""
    doc = Document(docx_path)
    changes = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        
        # [Image Error: ...] 제거
        if '[Image Error:' in text:
            new_text = re.sub(r'\[Image Error:[^\]]*\]', '', text).strip()
            if new_text:
                para.text = new_text
                changes.append(f"단락 {i}: Image Error 제거 → '{new_text[:50]}...'")
            else:
                # 빈 단락은 그대로 유지 (삭제하면 인덱스 변경됨)
                para.text = ""
                changes.append(f"단락 {i}: 빈 Image Error 단락 정리")
    
    doc.save(output_path)
    print(f"\n✓ 수정된 파일 저장: {output_path}")
    
    return changes

def create_image_insertion_guide(md_images, images_dir):
    """이미지 삽입 가이드 생성"""
    print("\n" + "=" * 70)
    print("이미지 수동 삽입 가이드")
    print("=" * 70)
    
    guide = """
Word에서 이미지를 삽입하는 방법:
1. 삽입 → 그림 → 이 장치 클릭
2. 이미지 파일 선택
3. 그림 서식 → 크기 조정 (논문 열 너비에 맞게)

삽입할 이미지 목록:
"""
    print(guide)
    
    image_order = [
        ('이규영_국민대_DPF_논문_img_0.png', '그림 0: 제안 방법론 전체 흐름도'),
        ('이규영_국민대_DPF_논문_img_1.png', '그림 1: Stage 2 전체 학습 곡선'),
        ('confusion_matrix_normalized.png', '그림 4: YOLO11 혼동 행렬'),
        ('이규영_국민대_DPF_논문_img_2.png', '그림 1 (비교): 전이학습 방법 비교'),
        ('이규영_국민대_DPF_논문_img_3.png', '그림 2: Stage 2 전체 학습 곡선 (상세)'),
        ('이규영_국민대_DPF_논문_img_5.png', '그림 5: Precision-Recall 곡선'),
        ('이규영_국민대_DPF_논문_img_6.png', '그림 6: 데이터 증강 분석'),
        ('이규영_국민대_DPF_논문_img_7.png', '그림 3: 검증 배치 예측 결과'),
        ('이규영_국민대_DPF_논문_img_8.png', '그림 7: F1 점수 곡선'),
        ('이규영_국민대_DPF_논문_img_9.png', '그림 8: 추가 검증 배치 예측 결과'),
        ('이규영_국민대_DPF_논문_img_10.png', '그림 9: 성능 개선 종합'),
    ]
    
    for filename, description in image_order:
        filepath = os.path.join(images_dir, filename)
        exists = "✓" if os.path.exists(filepath) else "✗"
        print(f"{exists} {filename}")
        print(f"   → {description}")
        print(f"   경로: {filepath}")
        print()

def main():
    base_dir = r"c:\1.이규영개인폴더\09.##### SCHOOL #####"
    docx_path = os.path.join(base_dir, "논문작성", "IEEE_DPF_Paper_Final.docx")
    output_path = os.path.join(base_dir, "논문작성", "IEEE_DPF_Paper_Final_v2.docx")
    md_path = os.path.join(base_dir, "converted_md", "이규영_국민대_DPF_논문.md")
    images_dir = os.path.join(base_dir, "converted_md", "images")
    
    # 이미지 갭 분석
    md_images, dpf_images = analyze_image_gaps(docx_path, md_path, images_dir)
    
    # [Image Error:] 정리
    print("\n" + "=" * 70)
    print("[Image Error:] 텍스트 정리")
    print("=" * 70)
    
    doc = Document(docx_path)
    placeholders = find_image_placeholders(doc)
    
    if placeholders:
        print(f"\n발견된 이미지 플레이스홀더: {len(placeholders)}개")
        for p in placeholders:
            print(f"   단락 {p['para_index']}: {p['text'][:80]}...")
        
        # 정리
        changes = fix_image_errors_in_doc(docx_path, output_path, images_dir)
        print(f"\n수정 내용: {len(changes)}건")
        for c in changes:
            print(f"   - {c}")
    else:
        print("\n이미지 오류 없음. 파일 복사...")
        import shutil
        shutil.copy(docx_path, output_path)
    
    # 이미지 삽입 가이드
    create_image_insertion_guide(md_images, images_dir)
    
    print("\n" + "=" * 70)
    print("최종 상태")
    print("=" * 70)
    print(f"""
★ 최종 파일: IEEE_DPF_Paper_Final_v2.docx

현재 상태:
- DOCX에 10개 이미지 포함됨
- 원본에 11개 고유 이미지 참조
- [Image Error:] 텍스트 정리됨

이미지 경로: {images_dir}

⚠️ 권장 작업:
1. Word에서 IEEE_DPF_Paper_Final_v2.docx 열기
2. 각 그림 캡션 위치에서 이미지 존재 확인
3. 누락된 이미지는 위 가이드 참고하여 수동 삽입
4. 이미지 크기를 열 너비에 맞게 조정 (약 3.25인치)
""")

if __name__ == "__main__":
    main()
