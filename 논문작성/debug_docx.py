from docx import Document

template_path = r"C:\국민대프로젝트\IEEE_TII_Template\TII_Articles_Word_template_2025.docx"
output_path = r"C:\국민대프로젝트\논문작성\debug_output.docx"

print("Opening template...")
doc = Document(template_path)
print(f"Initial paragraphs: {len(doc.paragraphs)}")

print("Clearing content...")
doc._body.clear_content()
print(f"Post-clear paragraphs: {len(doc.paragraphs)}")

print("Adding one line...")
doc.add_paragraph("Hello World Test String 12345")

print("Saving...")
doc.save(output_path)

print("Reading back...")
doc2 = Document(output_path)
paras = [p.text for p in doc2.paragraphs]
print(f"Paragraphs found: {len(paras)}")
for p in paras:
    print(f"Content: '{p}'")
