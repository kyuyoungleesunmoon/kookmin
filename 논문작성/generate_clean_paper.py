from docx import Document
from docx.shared import Pt
import re
import os

# Paths
md_path = r"C:\국민대프로젝트\converted_md\이규영_국민대_DPF_논문.md"
template_path = r"C:\국민대프로젝트\IEEE_TII_Template\TII_Articles_Word_template_2025.docx"
output_path = r"C:\국민대프로젝트\논문작성\IEEE_DPF_Paper_Submission_Ready.docx"

print(f"Reading Markdown: {md_path}")
with open(md_path, 'r', encoding='utf-8') as f:
    lines = f.readlines()

print(f"Opening Template: {template_path}")
try:
    doc = Document(template_path)
except Exception as e:
    print(f"Template load failed: {e}. creating new.")
    doc = Document()

# Clear existing content from template (optional, but safer to append if it's a blank template)
# But TII template usually has dummy text. We should probably keep the header/footer but clear body.
# For now, let's just append to the end, or clear paragraphs if there are many.
if len(doc.paragraphs) > 10:
    print("Clearing template dummy content...")
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

# Helper to add styled paragraph
def add_para(text, style=None):
    try:
        if style:
            p = doc.add_paragraph(text, style=style)
        else:
            p = doc.add_paragraph(text)
    except:
        p = doc.add_paragraph(text) # Fallback if style missing
    return p

current_section_level = 0

print("Processing content...")
for line in lines:
    line = line.strip()
    if not line:
        continue
        
    # Headers
    if line.startswith('# '):
        # Title
        text = line.replace('# ', '').replace('*', '')
        add_para(text, style='Title')
    elif line.startswith('## '):
        # Abstract or Main Title (The md has ## for Title sometimes?)
        text = line.replace('## ', '').replace('*', '')
        # Check if it looks like a section
        if "초록" in text or "Abstract" in text:
            add_para("Abstract", style='Heading 1')
            add_para(text, style='Body Text')
        else:
            add_para(text, style='Heading 1')
            
    elif line.startswith('### '):
        # Section Level 1 (I. 서론)
        text = line.replace('### ', '').replace('*', '')
        add_para(text, style='Heading 1')
        
    elif line.startswith('#### '):
        # Section Level 2 (A. 연구 배경)
        text = line.replace('#### ', '').replace('*', '')
        add_para(text, style='Heading 2')
        
    elif line.startswith('##### '):
        text = line.replace('##### ', '').replace('*', '')
        add_para(text, style='Heading 3')
        
    # Images
    elif line.startswith('!['):
        # ![Caption](path)
        match = re.search(r'!\[(.*?)\]\((.*?)\)', line)
        if match:
            caption = match.group(1)
            img_path = match.group(2)
            # Fix relative paths
            if not os.path.isabs(img_path):
                 # Assuming relative to md file location
                 base_dir = os.path.dirname(md_path)
                 img_full_path = os.path.join(base_dir, img_path)
            else:
                img_full_path = img_path
                
            if os.path.exists(img_full_path):
                try:
                    doc.add_picture(img_full_path, width=Pt(400))
                    add_para(f"Fig. {caption}", style='Caption')
                except Exception as e:
                    add_para(f"[Image: {caption} - {img_path}]")
            else:
                 add_para(f"[Image Missing: {caption} - {img_path}]")
                 
    # Tables (Simulated)
    elif line.startswith('|'):
        # Just add as text for now, converting tables is hard
        add_para(line, style='Normal')
        
    # Normal Text
    else:
        # Check if Metadata
        if line.startswith('저자:') or line.startswith('소속:'):
            add_para(line, style='Author')
        else:
            add_para(line, style='Body Text')

print(f"Saving to: {output_path}")
doc.save(output_path)
print("Done.")
