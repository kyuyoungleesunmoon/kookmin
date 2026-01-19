from docx import Document
import re

input_path = r"C:\국민대프로젝트\논문작성\IEEE_DPF_Paper_Clean_v7.docx"
output_path = r"C:\국민대프로젝트\논문작성\IEEE_DPF_Paper_Final_Clean.docx"

doc = Document(input_path)

# Final cleanup patterns - very specific
final_code_patterns = [
    r'^import\s+',
    r'^from\s+\w+\s+import',
    r"^'\w+'\s*:\s*\[\]",  # Empty list dict entries
    r"^'\w+'\s*:\s*$",  # Dict keys only
    r'^SEED\s*=',
    r'^random\.seed',
    r'^np\.random\.seed',
    r'^torch\.manual_seed',
    r'^torch\.backends',
    r'^os\.environ',
    r'^dataloader_config\s*=',
    r'^logs\s*=\s*\{',
]

to_remove = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue
    
    for pat in final_code_patterns:
        if re.search(pat, text):
            to_remove.append(i)
            break

print(f"Final removal: {len(to_remove)} paragraphs")

for idx in sorted(set(to_remove), reverse=True):
    try:
        p = doc.paragraphs[idx]._element
        p.getparent().remove(p)
    except:
        pass

doc.save(output_path)

# Final verification
doc2 = Document(output_path)
code_count = 0
eq_count = 0

for para in doc2.paragraphs:
    text = para.text.strip()
    if re.search(r'^(import|from\s+\w+\s+import|def\s+|class\s+)', text):
        code_count += 1
    if '$$' in text or '\\frac{' in text or '\\sum' in text:
        eq_count += 1

print(f"Final paragraph count: {len(doc2.paragraphs)}")
print(f"Remaining code blocks: {code_count}")
print(f"Remaining equation blocks: {eq_count}")
print(f"Saved to: {output_path}")
