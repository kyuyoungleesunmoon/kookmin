# -*- coding: utf-8 -*-
"""
소스코드를 서술형으로 변환하여 문서 재생성
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

def convert_code_to_prose():
    """소스코드 블록을 서술형 문장으로 변환"""
    
    # Stage 1 설정 서술
    stage1_prose = """Stage 1 학습에서는 일반 결함 데이터에서 기본적인 특징 표현을 학습하기 위해 다음과 같은 하이퍼파라미터를 설정하였다. 학습 에포크는 50회로 설정하여 일반 특징을 충분히 학습하도록 하였으며, 배치 크기는 CPU 메모리를 고려하여 8로 설정하였다. 입력 이미지 크기는 640×640 픽셀로 통일하였다. 학습률은 초기 0.01에서 cosine decay 스케줄을 적용하여 최종 0.01까지 감소하도록 설정하였다. SGD 모멘텀은 0.937, L2 정규화를 위한 weight decay는 0.0005로 설정하였다. 학습 초기 3 에포크 동안 학습률 워밍업을 적용하였으며, 옵티마이저는 AdamW를 사용하였다. 조기 종료 patience는 10 에포크로 설정하였고, 실험 재현성을 위해 무작위 시드를 42로 고정하였다."""

    # Stage 1 augmentation 서술
    stage1_aug_prose = """Stage 1의 데이터 증강 전략은 일반 결함 데이터의 특성을 고려하여 보수적으로 설정하였다. 색조 변화(hsv_h)는 0.015로 최소화하였으며, 채도 변화(hsv_s)는 0.7, 명도 변화(hsv_v)는 0.4로 설정하였다. 기하학적 변환으로는 ±10도 회전(degrees), 10% 평행 이동(translate), 50-150% 크기 조절(scale), ±5도 전단 변환(shear)을 적용하였다. 수평 및 수직 반전은 각각 50% 확률로 적용하였다. Mosaic 증강은 100% 적용하였으나, Mixup과 Copy-Paste는 사용하지 않았다."""

    # Stage 2 설정 서술
    stage2_prose = """Stage 2 학습에서는 DPF 도메인 적응을 위해 학습 설정을 조정하였다. 학습 에포크는 100회로 확장하여 충분한 수렴 시간을 확보하였다. 배치 크기와 이미지 크기는 Stage 1과 동일하게 유지하였다. 학습률, 모멘텀, weight decay 등의 기본 하이퍼파라미터는 Stage 1과 동일하게 유지하여 학습 안정성을 확보하였다. 다만 조기 종료 patience는 15 에포크로 증가시켜 후반부 학습(늦은 개화 현상)의 기회를 보장하였다. 또한 마지막 10 에포크에서는 Mosaic 증강을 해제(close_mosaic=10)하여 실제 이미지 분포에 대한 미세 조정이 이루어지도록 하였다."""

    # Stage 2 augmentation 서술
    stage2_aug_prose = """Stage 2의 데이터 증강은 DPF 이미지의 특성을 반영하여 설정하였다. 기하학적 변환은 Stage 1과 동일하게 적용하였으나, X-ray 이미지의 그레이스케일 특성을 고려하여 색조 변화는 최소(0.015)로 유지하였다. 채도 변화(0.7)와 명도 변화(0.4)는 X-ray 이미지의 대비 변화를 시뮬레이션하기 위해 적용하였다. 가우시안 블러(0.01)와 가우시안 노이즈(0.02)는 최소한으로 적용하여 과도한 이미지 품질 저하를 방지하였다."""

    # 검증 설정 서술
    validation_prose = """검증 단계에서는 재현성과 정확한 성능 측정을 위해 데이터 로더를 구성하였다. 워커 수(num_workers)는 0으로 설정하여 단일 워커로 순서를 보장하였고, 검증 시에는 데이터 셔플을 비활성화하였다. 모든 검증 샘플을 사용하기 위해 drop_last는 False로 설정하였다."""

    # 학습 모니터링 서술
    monitoring_prose = """학습 과정 모니터링을 위해 각 에포크별로 다음 지표들을 기록하였다: 학습 손실(train_loss), 검증 손실(val_loss), 정밀도(precision), 재현율(recall), mAP50, mAP50-95, 학습률(lr), 그리고 소요 시간(time). 이러한 상세 로깅은 학습 동역학 분석 및 늦은 개화 현상 발견에 핵심적인 역할을 하였다."""

    return {
        'stage1_config': stage1_prose,
        'stage1_aug': stage1_aug_prose,
        'stage2_config': stage2_prose,
        'stage2_aug': stage2_aug_prose,
        'validation': validation_prose,
        'monitoring': monitoring_prose
    }

def recreate_document_with_prose(input_path, output_path):
    """소스코드를 서술형으로 변환하여 문서 재생성"""
    doc = Document(input_path)
    prose = convert_code_to_prose()
    
    changes = 0
    skip_until = -1
    paragraphs_to_modify = []
    
    # 코드 블록 위치 및 변환 내용 매핑
    code_sections = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Stage 1 config 시작
        if 'stage1_config' in text and '=' in text:
            code_sections.append({
                'start': i,
                'type': 'stage1_config',
                'prose': prose['stage1_config']
            })
        
        # Stage 1 augmentation
        elif i < 400 and 'augmentation' in text and '=' in text and '{' in text:
            code_sections.append({
                'start': i,
                'type': 'stage1_aug',
                'prose': prose['stage1_aug']
            })
        
        # Stage 2 config
        elif 'stage2_config' in text and '=' in text:
            code_sections.append({
                'start': i,
                'type': 'stage2_config',
                'prose': prose['stage2_config']
            })
        
        # Stage 2 augmentation (위치 기반)
        elif i > 600 and i < 700 and 'augmentation' in text and '=' in text:
            code_sections.append({
                'start': i,
                'type': 'stage2_aug', 
                'prose': prose['stage2_aug']
            })
        
        # dataloader_config
        elif 'dataloader_config' in text and '=' in text:
            code_sections.append({
                'start': i,
                'type': 'validation',
                'prose': prose['validation']
            })
        
        # training_history
        elif "'epoch'" in text and "'train_loss'" in text.replace(' ', ''):
            code_sections.append({
                'start': i,
                'type': 'monitoring',
                'prose': prose['monitoring']
            })
    
    print(f"발견된 코드 섹션: {len(code_sections)}")
    for sec in code_sections:
        print(f"  - {sec['type']} at paragraph {sec['start']}")
    
    # 코드 블록 범위 계산
    for sec in code_sections:
        start_idx = sec['start']
        end_idx = start_idx
        
        # 닫는 괄호까지 찾기
        for j in range(start_idx, min(start_idx + 30, len(doc.paragraphs))):
            text = doc.paragraphs[j].text.strip()
            if text == '}' or text == '},':
                end_idx = j
                break
            elif j > start_idx and re.match(r'^[a-zA-Z]', text) and '=' not in text and ':' not in text:
                end_idx = j - 1
                break
        
        sec['end'] = end_idx
        print(f"  {sec['type']}: paragraphs {start_idx}-{end_idx}")
    
    # 역순으로 처리 (인덱스 변화 방지)
    code_sections.sort(key=lambda x: x['start'], reverse=True)
    
    for sec in code_sections:
        start_idx = sec['start']
        end_idx = sec['end']
        
        # 첫 번째 단락에 서술형 텍스트 삽입
        doc.paragraphs[start_idx].text = sec['prose']
        
        # 나머지 코드 단락 삭제 (내용만 비우기)
        for j in range(start_idx + 1, end_idx + 1):
            if j < len(doc.paragraphs):
                doc.paragraphs[j].text = ""
        
        changes += 1
    
    # 추가: 남은 코드 패턴 정리
    additional_cleanup = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # 단독 코드 라인 패턴
        if re.match(r"^'[a-z_]+'\s*:\s*[\d\.\[\]]+,?$", text):
            para.text = ""
            additional_cleanup += 1
        elif re.match(r"^[\[\{}\]],?$", text):
            para.text = ""
            additional_cleanup += 1
        elif text in ['training_history = {', '}']:
            para.text = ""
            additional_cleanup += 1
    
    print(f"\n서술형 변환: {changes}건")
    print(f"추가 정리: {additional_cleanup}건")
    
    # 늦은 개화 용어 복원
    late_bloom = 0
    for para in doc.paragraphs:
        text = para.text
        if 'Late Blooming' in text and '늦은 개화' not in text:
            para.text = text.replace('Late Blooming', '늦은 개화(Late Blooming)')
            late_bloom += 1
    
    print(f"늦은 개화 복원: {late_bloom}건")
    
    doc.save(output_path)
    return changes

def verify_document(docx_path):
    """문서 검증"""
    doc = Document(docx_path)
    full_text = ' '.join([p.text for p in doc.paragraphs])
    
    print("\n" + "=" * 60)
    print("검증 결과")
    print("=" * 60)
    
    # 서술형 변환 확인
    prose_checks = [
        ('학습 에포크는 50회', 'Stage 1 설정 서술'),
        ('학습 에포크는 100회', 'Stage 2 설정 서술'),
        ('Mosaic 증강은 100%', 'Augmentation 서술'),
        ('cosine decay', '학습률 스케줄 서술'),
        ('워커 수', '검증 설정 서술'),
    ]
    
    print("\n서술형 변환 확인:")
    for term, desc in prose_checks:
        found = term in full_text
        status = "✓" if found else "✗"
        print(f"  {status} {desc}: '{term}'")
    
    # 코드 패턴 잔존 확인
    code_count = 0
    for para in doc.paragraphs:
        if re.search(r"'[a-z_]+'\s*:\s*\d", para.text):
            code_count += 1
    
    # 핵심 내용 확인
    checks = [
        ('91.7%', 'mAP50'),
        ('늦은 개화', 'Late Blooming'),
        ('도메인 브리지', '핵심 방법론'),
        ('34.8%p', '총 성능 향상'),
    ]
    
    print("\n핵심 내용 확인:")
    for term, desc in checks:
        found = term in full_text
        status = "✓" if found else "✗"
        print(f"  {status} {term} ({desc})")
    
    print(f"\n잔존 코드 패턴: {code_count}개")
    
    return code_count == 0

def main():
    base_dir = r"c:\1.이규영개인폴더\09.##### SCHOOL #####"
    input_path = os.path.join(base_dir, "논문작성", "IEEE_DPF_Paper_Final_v2.docx")
    output_path = os.path.join(base_dir, "논문작성", "IEEE_DPF_Paper_Final_v4.docx")
    
    print("=" * 60)
    print("소스코드 → 서술형 변환")
    print("=" * 60)
    
    recreate_document_with_prose(input_path, output_path)
    
    print(f"\n✓ 저장 완료: {output_path}")
    
    verify_document(output_path)

if __name__ == "__main__":
    main()
