from docx import Document

template_path = r"C:\국민대프로젝트\IEEE_TII_Template\TII_Articles_Word_template_2025.docx"
output_path = r"C:\국민대프로젝트\논문작성\debug_style_dupe.docx"

print("Opening template...")
doc = Document(template_path)
doc._body.clear_content()

def add_para(text, style=None):
    print(f"Adding: {text} with style {style}")
    try:
        p = doc.add_paragraph(text, style=style)
        print("Success")
    except Exception as e:
        print(f"Error: {e}")
        p = doc.add_paragraph(text)
        print("Fallback Success")
    return p

# Test with a style that MIGHT fail or warn
# 'Title' usually exists.
# Let's try 'NonExistentStyle' which definitely fails.
add_para("Should be once (fallback)", style="NonExistentStyle")

# Let's try 'Title' (if it works, it appears once)
add_para("Should be once (Title)", style="Title")

# Let's try 'Heading 1'
add_para("Should be once (Heading 1)", style="Heading 1")

doc.save(output_path)

print("\nReading back...")
doc2 = Document(output_path)
paras = [p.text for p in doc2.paragraphs]
for p in paras:
    print(f"Content: '{p}'")
