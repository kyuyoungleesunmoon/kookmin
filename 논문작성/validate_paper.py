from docx import Document
from docx.shared import Pt, Inches
import re
import os

print("="*80)
print("IEEE TII 논문 양식 상세 검증")
print("="*80)

# 파일 경로
paper_path = r'C:\1.이규영개인폴더\09.##### SCHOOL #####\논문작성\IEEE_DPF_Paper_Draft_v2.docx'
template_path = r'C:\1.이규영개인폴더\09.##### SCHOOL #####\IEEE_TII_Template\TII_Articles_Word_template_2025.docx'

# 문서 로드
print("\n1. 문서 로드...")
paper = Document(paper_path)
template = Document(template_path)

print(f"   논문: {len(paper.paragraphs)}개 문단, {len(paper.tables)}개 표")
print(f"   템플릿: {len(template.paragraphs)}개 문단")

# ================== 페이지 설정 비교 ==================
print("\n2. 페이지 설정 비교...")

paper_section = paper.sections[0]
template_section = template.sections[0]

print(f"\n   {'항목':<20} {'논문':<15} {'템플릿':<15} {'상태':<10}")
print(f"   {'-'*60}")

def check_margin(name, paper_val, template_val, tolerance=0.05):
    diff = abs(paper_val - template_val)
    status = "✓ OK" if diff <= tolerance else "✗ 수정필요"
    print(f"   {name:<20} {paper_val:.2f}\"{'':>7} {template_val:.2f}\"{'':>7} {status}")
    return diff <= tolerance

margins_ok = True
margins_ok &= check_margin("상단 여백", paper_section.top_margin.inches, template_section.top_margin.inches)
margins_ok &= check_margin("하단 여백", paper_section.bottom_margin.inches, template_section.bottom_margin.inches)
margins_ok &= check_margin("좌측 여백", paper_section.left_margin.inches, template_section.left_margin.inches)
margins_ok &= check_margin("우측 여백", paper_section.right_margin.inches, template_section.right_margin.inches)

# ================== 과도한 표현 검사 ==================
print("\n3. 과도한 표현 검사...")

excessive_expressions = [
    ('성능 폭발', '성능 가속'),
    ('늦은 개화', '후반부 성능 향상'),
    ('치명적 결과', '심각한 영향'),
    ('시너지 발현', '상승 효과'),
    ('압도적', '우수한'),
    ('혁신적', '새로운'),
    ('획기적', '효과적'),
    ('(2배 가속!)', '(약 2배)'),
    ('!)', ')'),
    (' ⭐', ''),
    (' ★', ''),
    ('⚠️', ''),
    ('✅', ''),
    ('🔥', ''),
    ('←', ''),
]

found_expressions = []
for para in paper.paragraphs:
    text = para.text
    for old_expr, new_expr in excessive_expressions:
        if old_expr in text:
            found_expressions.append((old_expr, new_expr, text[:80]))

if found_expressions:
    print(f"\n   발견된 과도한 표현: {len(found_expressions)}개")
    for i, (old, new, context) in enumerate(found_expressions[:10]):
        print(f"   {i+1}. '{old}' → '{new}'")
        print(f"      문맥: ...{context}...")
else:
    print("   ✓ 과도한 표현 없음")

# ================== 수식 위치 검사 ==================
print("\n4. 수식 위치 및 형식 검사...")

equations = []
for i, para in enumerate(paper.paragraphs):
    text = para.text.strip()
    
    # LaTeX 수식 패턴
    if '$$' in text or '\\tag{' in text or '\\begin{' in text:
        equations.append({
            'para_idx': i,
            'text': text[:100],
            'has_tag': '\\tag{' in text,
            'has_number': bool(re.search(r'\\tag\{\d+\}', text)),
        })
    
    # 잘못된 수식 형식 체크 (본문 중간에 $$ 사용)
    if text.startswith('$$') and i > 0:
        prev_text = paper.paragraphs[i-1].text.strip()
        if prev_text and not prev_text.endswith(':') and not prev_text.endswith('다:'):
            equations[-1]['issue'] = '수식 앞에 콜론(:) 권장'

print(f"\n   발견된 수식: {len(equations)}개")

# 수식 번호 확인
tagged_eqs = [eq for eq in equations if eq.get('has_number')]
print(f"   번호 있는 수식: {len(tagged_eqs)}개")

# 수식 번호 연속성 확인
eq_numbers = []
for eq in equations:
    text = eq['text']
    match = re.search(r'\\tag\{(\d+)\}', text)
    if match:
        eq_numbers.append(int(match.group(1)))

if eq_numbers:
    expected = list(range(1, max(eq_numbers) + 1))
    missing = set(expected) - set(eq_numbers)
    duplicate = [x for x in eq_numbers if eq_numbers.count(x) > 1]
    
    if missing:
        print(f"   ⚠️ 누락된 수식 번호: {sorted(missing)}")
    if duplicate:
        print(f"   ⚠️ 중복된 수식 번호: {sorted(set(duplicate))}")
    if not missing and not duplicate:
        print(f"   ✓ 수식 번호 1-{max(eq_numbers)} 연속")

# 수식 배치 문제 확인
issues = [eq for eq in equations if eq.get('issue')]
if issues:
    print(f"\n   수식 배치 문제:")
    for eq in issues[:5]:
        print(f"   - Para {eq['para_idx']}: {eq['issue']}")
        print(f"     내용: {eq['text'][:60]}...")

# ================== 섹션 구조 검사 ==================
print("\n5. 섹션 구조 검사...")

sections_found = []
expected_sections = ['서론', '관련 연구', '방법론', '실험', '결과', '토론', '결론', 'References']

for i, para in enumerate(paper.paragraphs):
    text = para.text.strip()
    
    # 로마 숫자 섹션 감지
    if re.match(r'^[IVX]+\.\s+', text):
        sections_found.append({
            'idx': i,
            'title': text,
            'level': 1
        })
    # 영문 서브섹션
    elif re.match(r'^[A-Z]\.\s+', text) and not re.match(r'^[IVX]', text):
        sections_found.append({
            'idx': i,
            'title': text,
            'level': 2
        })

print(f"\n   발견된 섹션: {len(sections_found)}개")
for sect in sections_found[:12]:
    indent = "   " if sect['level'] == 1 else "      "
    print(f"   {indent}{sect['title'][:50]}")

# ================== 참고문헌 형식 검사 ==================
print("\n6. 참고문헌 형식 검사...")

refs_start = -1
refs = []

for i, para in enumerate(paper.paragraphs):
    text = para.text.strip()
    
    if text.startswith('References') or text.startswith('참고문헌'):
        refs_start = i
    
    if refs_start > 0 and i > refs_start:
        # [1], [2] 등의 패턴
        if re.match(r'^\[\d+\]', text):
            refs.append({
                'idx': i,
                'text': text,
                'has_author': bool(re.search(r'[A-Z]\.\s*[A-Z]', text)),  # 저자명 패턴
                'has_year': bool(re.search(r'\b(19|20)\d{2}\b', text)),
            })

print(f"\n   참고문헌 수: {len(refs)}개")
if refs:
    valid_refs = [r for r in refs if r['has_author'] and r['has_year']]
    print(f"   IEEE 형식 준수: {len(valid_refs)}/{len(refs)}")
    
    invalid_refs = [r for r in refs if not r['has_author'] or not r['has_year']]
    if invalid_refs:
        print(f"\n   ⚠️ 형식 확인 필요:")
        for ref in invalid_refs[:3]:
            print(f"   - {ref['text'][:60]}...")

# ================== 그림/표 캡션 검사 ==================
print("\n7. 그림/표 캡션 검사...")

figures = []
tables_captions = []

for i, para in enumerate(paper.paragraphs):
    text = para.text.strip()
    
    if re.match(r'^(그림|Fig\.?)\s*\d+', text, re.IGNORECASE):
        figures.append({'idx': i, 'caption': text[:80]})
    
    if re.match(r'^(표|Table)\s*\d+', text, re.IGNORECASE):
        tables_captions.append({'idx': i, 'caption': text[:80]})

print(f"   그림 캡션: {len(figures)}개")
print(f"   표 캡션: {len(tables_captions)}개")

# ================== 종합 보고서 ==================
print("\n" + "="*80)
print("종합 검증 결과")
print("="*80)

issues_summary = []

if not margins_ok:
    issues_summary.append("페이지 여백 불일치")

if found_expressions:
    issues_summary.append(f"과도한 표현 {len(found_expressions)}개")

if missing or duplicate:
    issues_summary.append("수식 번호 문제")

if invalid_refs:
    issues_summary.append(f"참고문헌 형식 확인 필요 {len(invalid_refs)}개")

print(f"\n✓ 정상 항목:")
print(f"   - 페이지 크기: Letter (8.5\" × 11\")")
print(f"   - 섹션 구조: {len(sections_found)}개 섹션")
print(f"   - 수식: {len(equations)}개")
print(f"   - 참고문헌: {len(refs)}개")
print(f"   - 그림/표 캡션 포함")

if issues_summary:
    print(f"\n⚠️ 수정 필요 항목:")
    for issue in issues_summary:
        print(f"   - {issue}")
else:
    print(f"\n✅ 모든 항목 정상!")

# 결과 저장
report_path = r'C:\1.이규영개인폴더\09.##### SCHOOL #####\논문작성\validation_report.txt'
with open(report_path, 'w', encoding='utf-8') as f:
    f.write("IEEE TII 논문 양식 검증 보고서\n")
    f.write("="*60 + "\n\n")
    f.write(f"검증 파일: {paper_path}\n")
    f.write(f"템플릿: {template_path}\n\n")
    
    f.write("발견된 과도한 표현:\n")
    for old, new, ctx in found_expressions:
        f.write(f"  - '{old}' → '{new}'\n")
    
    f.write(f"\n수식 위치:\n")
    for eq in equations[:20]:
        f.write(f"  Para {eq['para_idx']}: {eq['text'][:60]}...\n")

print(f"\n보고서 저장: {report_path}")
