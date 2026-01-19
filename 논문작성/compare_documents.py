# -*- coding: utf-8 -*-
"""
IEEE_DPF_Paper_Modern.docx와 원본 마크다운 파일 비교 분석
"""

import zipfile
import xml.etree.ElementTree as ET
from docx import Document
import os
import re

def analyze_docx(docx_path):
    """DOCX 문서의 상세 분석"""
    doc = Document(docx_path)
    
    analysis = {
        'paragraphs': [],
        'tables': [],
        'equations': [],
        'images': 0,
        'headings': [],
        'sections': []
    }
    
    # 단락 분석
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            analysis['paragraphs'].append({
                'index': i,
                'text': text[:200] + '...' if len(text) > 200 else text,
                'style': para.style.name if para.style else 'Unknown'
            })
            
            # 수식 포함 여부 확인 (괄호 안 수식 패턴)
            eq_pattern = r'\([\d]+\)'
            if re.search(eq_pattern, text):
                analysis['equations'].append({
                    'para_index': i,
                    'text': text
                })
            
            # 섹션 헤딩 확인
            if para.style and 'Heading' in para.style.name:
                analysis['headings'].append({
                    'index': i,
                    'text': text,
                    'style': para.style.name
                })
    
    # 테이블 분석
    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows[:3]:  # 처음 3행만
            cells = [cell.text.strip()[:50] for cell in row.cells]
            rows.append(cells)
        analysis['tables'].append({
            'index': i,
            'rows': len(table.rows),
            'cols': len(table.columns),
            'preview': rows
        })
    
    # 이미지 수 계산 (ZIP 방식)
    with zipfile.ZipFile(docx_path, 'r') as z:
        media_files = [f for f in z.namelist() if f.startswith('word/media/')]
        analysis['images'] = len(media_files)
    
    return analysis

def analyze_markdown(md_path):
    """마크다운 문서의 상세 분석"""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    analysis = {
        'total_lines': len(content.split('\n')),
        'headings': re.findall(r'^#{1,6}\s+(.+)$', content, re.MULTILINE),
        'equations': [],
        'tables': [],
        'images': re.findall(r'!\[.*?\]\(.*?\)', content),
        'key_sections': []
    }
    
    # 수식 찾기 (번호가 있는 수식)
    eq_patterns = [
        r'\$\$[\s\S]*?\$\$',  # 블록 수식
        r'\$[^$\n]+\$',  # 인라인 수식
        r'\\begin\{equation\}[\s\S]*?\\end\{equation\}',
        r'\([0-9]{1,2}\)',  # 수식 번호
    ]
    
    for pattern in eq_patterns[:1]:
        found = re.findall(pattern, content)
        analysis['equations'].extend(found[:20])  # 처음 20개만
    
    # 표 찾기 (마크다운 테이블)
    table_pattern = r'\|[^\n]+\|[\n\r]+\|[-:\s|]+\|'
    analysis['tables'] = len(re.findall(table_pattern, content))
    
    return analysis

def extract_key_content(md_path):
    """마크다운에서 주요 내용 추출"""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    key_items = {
        'title': '',
        'abstract': '',
        'sections': [],
        'equations_with_numbers': [],
        'key_findings': [],
        'figures_mentioned': [],
        'tables_mentioned': []
    }
    
    # 제목 추출
    title_match = re.search(r'^##\s+(.+)$', content, re.MULTILINE)
    if title_match:
        key_items['title'] = title_match.group(1)
    
    # 초록 추출
    abstract_match = re.search(r'### 초록.*?\n\n([\s\S]*?)(?=\n###|\n####)', content)
    if abstract_match:
        key_items['abstract'] = abstract_match.group(1)[:500]
    
    # 번호가 붙은 수식 찾기
    numbered_eq = re.findall(r'\((\d+)\)\s*\n', content)
    key_items['equations_with_numbers'] = list(set(numbered_eq))
    
    # Figure/Table 언급 찾기
    figures = re.findall(r'Fig\.\s*(\d+)|Figure\s*(\d+)|그림\s*(\d+)', content)
    tables = re.findall(r'Table\s*(\d+)|표\s*(\d+)', content)
    
    key_items['figures_mentioned'] = [f for t in figures for f in t if f]
    key_items['tables_mentioned'] = [t for tup in tables for t in tup if t]
    
    return key_items

def compare_documents(docx_path, md_path):
    """두 문서 비교"""
    print("=" * 80)
    print("문서 비교 분석")
    print("=" * 80)
    
    # DOCX 분석
    print("\n[1] DOCX 분석 중...")
    docx_analysis = analyze_docx(docx_path)
    
    print(f"   - 단락 수: {len(docx_analysis['paragraphs'])}")
    print(f"   - 테이블 수: {len(docx_analysis['tables'])}")
    print(f"   - 이미지 수: {docx_analysis['images']}")
    print(f"   - 수식 포함 단락: {len(docx_analysis['equations'])}")
    
    # 마크다운 분석
    print("\n[2] 마크다운 분석 중...")
    md_analysis = analyze_markdown(md_path)
    md_content = extract_key_content(md_path)
    
    print(f"   - 총 라인 수: {md_analysis['total_lines']}")
    print(f"   - 헤딩 수: {len(md_analysis['headings'])}")
    print(f"   - 번호 수식: {md_content['equations_with_numbers']}")
    
    # 비교 결과
    print("\n" + "=" * 80)
    print("비교 결과")
    print("=" * 80)
    
    # 섹션 헤딩 비교
    print("\n[섹션 헤딩 비교]")
    docx_headings = [h['text'] for h in docx_analysis['headings']]
    print(f"DOCX 헤딩: {len(docx_headings)}")
    for h in docx_headings[:10]:
        print(f"   - {h}")
    
    print(f"\n마크다운 헤딩: {len(md_analysis['headings'])}")
    for h in md_analysis['headings'][:10]:
        print(f"   - {h}")
    
    # 테이블 비교
    print("\n[테이블 비교]")
    print(f"DOCX 테이블 수: {len(docx_analysis['tables'])}")
    for i, t in enumerate(docx_analysis['tables'][:5]):
        print(f"   Table {i+1}: {t['rows']} rows x {t['cols']} cols")
        if t['preview']:
            print(f"      첫 행: {t['preview'][0][:3]}")
    
    # 수식 비교
    print("\n[수식 비교]")
    docx_eq_count = len(docx_analysis['equations'])
    md_eq_numbers = sorted([int(n) for n in md_content['equations_with_numbers'] if n.isdigit()])
    
    print(f"DOCX 수식 포함 단락: {docx_eq_count}")
    print(f"마크다운 번호 수식: {md_eq_numbers}")
    
    # 이미지 비교
    print("\n[이미지 비교]")
    print(f"DOCX 이미지 수: {docx_analysis['images']}")
    print(f"마크다운 이미지 참조: {len(md_analysis['images'])}")
    
    return {
        'docx': docx_analysis,
        'markdown': md_analysis,
        'md_content': md_content
    }

if __name__ == "__main__":
    base_dir = r"c:\1.이규영개인폴더\09.##### SCHOOL #####"
    docx_path = os.path.join(base_dir, "논문작성", "IEEE_DPF_Paper_Modern.docx")
    md_path = os.path.join(base_dir, "converted_md", "이규영_국민대_DPF_논문.md")
    
    if os.path.exists(docx_path) and os.path.exists(md_path):
        results = compare_documents(docx_path, md_path)
    else:
        print(f"파일을 찾을 수 없습니다:")
        print(f"  DOCX: {docx_path} - 존재: {os.path.exists(docx_path)}")
        print(f"  MD: {md_path} - 존재: {os.path.exists(md_path)}")
