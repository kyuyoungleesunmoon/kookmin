# -*- coding: utf-8 -*-
"""
IEEE_DPF_Paper_Final_v2.docx 수정
1. 과도한 수식 제거 (핵심 수식만 유지)
2. 소스코드를 서술형으로 변환
"""

from docx import Document
from docx.shared import Pt
import os
import re

def identify_essential_equations():
    """꼭 필요한 핵심 수식 번호 (유지할 것)"""
    # 핵심적으로 필요한 수식만 선별
    essential = {
        # 도메인 거리 측정 - 핵심 방법론
        1: "Proxy A-distance (도메인 거리)",
        
        # 손실 함수 - 모델 학습 핵심
        3: "YOLO 통합 손실 함수",
        
        # 평가 지표 - 결과 해석 필수
        7: "AP 계산식",
        8: "mAP 계산식", 
        9: "Precision",
        10: "Recall",
        11: "F1 Score",
        
        # 학습률 스케줄링 - 실험 설정
        12: "Cosine Annealing 학습률",
    }
    return essential

def convert_code_to_prose(text):
    """소스코드를 서술형으로 변환"""
    
    # 하이퍼파라미터 딕셔너리 변환
    if re.match(r"^\s*'[a-z_]+'\s*:", text):
        # 'epochs': 50, → epochs는 50으로 설정
        match = re.match(r"^\s*'([a-z_]+)'\s*:\s*([^,]+),?\s*(#.*)?", text)
        if match:
            param, value, comment = match.groups()
            comment = comment.strip('# ') if comment else ''
            
            param_names = {
                'epochs': '에포크 수',
                'batch': '배치 크기',
                'imgsz': '입력 이미지 크기',
                'lr0': '초기 학습률',
                'lrf': '최종 학습률 비율',
                'momentum': '모멘텀',
                'weight_decay': '가중치 감쇠',
                'warmup_epochs': '워밍업 에포크',
                'optimizer': '옵티마이저',
                'patience': '조기 종료 인내값',
                'device': '학습 장치',
                'seed': '무작위 시드',
                'close_mosaic': 'Mosaic 해제 에포크',
                'hsv_h': '색조 변화',
                'hsv_s': '채도 변화',
                'hsv_v': '명도 변화',
                'degrees': '회전 각도',
                'translate': '평행 이동',
                'scale': '크기 변화',
                'flipud': '상하 반전 확률',
                'fliplr': '좌우 반전 확률',
                'mosaic': 'Mosaic 증강',
                'mixup': 'Mixup 증강',
            }
            
            korean_name = param_names.get(param, param)
            return f"{korean_name}: {value}"
    
    return text

def should_remove_paragraph(text):
    """제거해야 할 단락 판단"""
    
    # 빈 텍스트
    if not text.strip():
        return False
    
    # 딕셔너리/리스트 시작/끝
    if text.strip() in ['{', '}', '[', ']', '};', '},']:
        return True
    
    # 코드 주석만 있는 경우
    if re.match(r'^\s*#\s*[=\-]+\s*$', text):
        return True
    
    # 변수 할당 시작 (빈 딕셔너리)
    if re.match(r'^\s*\w+\s*=\s*\{$', text):
        return True
    
    return False

def process_document(docx_path, output_path):
    """문서 처리"""
    doc = Document(docx_path)
    
    changes = {
        'equations_removed': 0,
        'code_converted': 0,
        'paragraphs_removed': 0,
    }
    
    essential_eqs = identify_essential_equations()
    
    # 제거할 단락 인덱스 수집
    paragraphs_to_clear = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        if not text:
            continue
        
        # 1. 불필요한 수식 제거 (번호 있는 수식 중 비핵심)
        tag_match = re.search(r'\\tag\{(\d+)\}', text)
        if tag_match:
            eq_num = int(tag_match.group(1))
            if eq_num not in essential_eqs:
                # 비핵심 수식은 제거하되, 설명만 남기기
                if eq_num in [2, 4, 5, 6, 13, 14, 15, 16, 17, 18, 19, 20]:
                    paragraphs_to_clear.append(i)
                    changes['equations_removed'] += 1
        
        # 2. 소스코드 패턴 처리
        code_patterns = [
            (r"^\s*model\s*=\s*YOLO\(['\"]([^'\"]+)['\"]\)", 
             lambda m: f"모델은 {m.group(1)} 가중치로 초기화한다."),
            
            (r"^\s*stage\d*_config\s*=\s*\{",
             lambda m: "학습 설정은 다음과 같다:"),
            
            (r"^\s*augmentation\s*=\s*\{",
             lambda m: "데이터 증강 설정:"),
            
            (r"^\s*#\s*(.+)$",
             lambda m: f""),  # 코드 주석 제거
            
            (r"^\s*'([a-z_]+)'\s*:\s*([^,]+),?\s*(#.*)?",
             lambda m: ""),  # 개별 파라미터는 통합 테이블로
        ]
        
        for pattern, replacement in code_patterns:
            match = re.match(pattern, text)
            if match:
                if callable(replacement):
                    new_text = replacement(match)
                else:
                    new_text = replacement
                
                if new_text:
                    para.text = new_text
                    changes['code_converted'] += 1
                else:
                    paragraphs_to_clear.append(i)
                break
    
    # 연속된 코드 블록을 서술형 단락으로 통합
    # (실제 구현은 복잡하므로 핵심 변환만 수행)
    
    # 단락 내용 지우기 (삭제 대신)
    for idx in paragraphs_to_clear:
        if idx < len(doc.paragraphs):
            doc.paragraphs[idx].text = ""
            changes['paragraphs_removed'] += 1
    
    doc.save(output_path)
    return changes

def create_clean_version(docx_path, output_path):
    """깔끔한 버전 생성 - 소스코드를 서술형으로 대폭 변환"""
    doc = Document(docx_path)
    
    # 변환 규칙
    code_block_replacements = {
        # Stage 1 설정
        "stage1_config = {": """
Stage 1 학습 설정은 다음과 같이 구성하였다. 에포크 수는 50으로 설정하여 일반적인 결함 특징을 충분히 학습하도록 하였다. 배치 크기는 CPU 메모리를 고려하여 8로 설정하였으며, 입력 이미지 크기는 640×640 픽셀을 사용하였다. 옵티마이저는 AdamW를 사용하였고, 초기 학습률 0.01에서 Cosine Annealing 스케줄링을 적용하였다. 조기 종료 인내값은 10 에포크로 설정하였으며, 재현성을 위해 무작위 시드를 42로 고정하였다.""",

        # Stage 2 설정
        "stage2_config = {": """
Stage 2 학습 설정에서는 충분한 수렴을 위해 에포크 수를 100으로 증가시켰다. 나머지 설정은 Stage 1과 동일하게 유지하되, 조기 종료 인내값을 15로 높여 후반부 개선을 허용하였다. 또한 마지막 10 에포크에서는 Mosaic 증강을 해제하여 정밀한 학습을 수행하였다.""",

        # 증강 설정
        "augmentation = {": """
데이터 증강은 기하학적 변환과 픽셀 변환을 조합하여 적용하였다. 기하학적 변환으로는 ±10도 회전, 10% 평행 이동, 50-150% 크기 변화, 상하/좌우 각 50% 확률의 반전을 적용하였다. 픽셀 변환으로는 명도 40% 변화를 주로 사용하였으며, 그레이스케일 특성상 색조 변화는 최소화하였다. Mosaic 증강을 100% 확률로 적용하여 다양한 객체 조합을 학습하도록 하였다.""",
    }
    
    # 제거할 코드 패턴들
    remove_patterns = [
        r"^\s*'[a-z_]+'\s*:\s*[^,]+,?\s*(#.*)?$",  # 딕셔너리 항목
        r"^\s*\}$",  # 닫는 괄호
        r"^\s*#\s*(===|---|Stage|사용).*$",  # 코드 주석
        r"^\s*(model|results|dataloader_config|logs|checkpoints)\s*=",  # 변수 할당
        r"^\s*(torch|random|np|os)\.",  # 라이브러리 호출
        r"^\s*def\s+\w+\(",  # 함수 정의
        r"^\s*for\s+\w+\s+in\s+",  # for 루프
        r"^\s*if\s+.*:$",  # if 문
        r"^\s*plt\.",  # matplotlib
        r"^\s*writer\.",  # tensorboard
        r"^\s*print\(",  # print 문
    ]
    
    changes = {'removed': 0, 'converted': 0}
    i = 0
    
    while i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        text = para.text.strip()
        
        if not text:
            i += 1
            continue
        
        # 코드 블록 시작 감지 및 대체
        replaced = False
        for trigger, replacement in code_block_replacements.items():
            if trigger in text:
                para.text = replacement.strip()
                changes['converted'] += 1
                replaced = True
                break
        
        if replaced:
            i += 1
            continue
        
        # 제거할 패턴 확인
        should_remove = False
        for pattern in remove_patterns:
            if re.match(pattern, text):
                should_remove = True
                break
        
        if should_remove:
            para.text = ""
            changes['removed'] += 1
        
        i += 1
    
    doc.save(output_path)
    return changes

def main():
    base_dir = r"c:\1.이규영개인폴더\09.##### SCHOOL #####"
    input_path = os.path.join(base_dir, "논문작성", "IEEE_DPF_Paper_Final_v2.docx")
    output_path = os.path.join(base_dir, "논문작성", "IEEE_DPF_Paper_Clean.docx")
    
    print("=" * 70)
    print("문서 정리: 과도한 수식/소스코드 제거")
    print("=" * 70)
    
    # 처리
    changes = create_clean_version(input_path, output_path)
    
    print(f"\n변환 결과:")
    print(f"  - 소스코드 → 서술형 변환: {changes['converted']}건")
    print(f"  - 코드 단락 제거: {changes['removed']}건")
    print(f"\n✓ 저장 완료: {output_path}")
    
    # 검증
    doc = Document(output_path)
    code_count = 0
    for para in doc.paragraphs:
        text = para.text
        if re.search(r"'[a-z_]+'\s*:", text) or re.match(r'\s*\w+\s*=\s*\{', text):
            code_count += 1
    
    print(f"\n[검증] 잔존 코드 패턴: {code_count}개")

if __name__ == "__main__":
    main()
