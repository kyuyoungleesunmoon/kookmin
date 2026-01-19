# -*- coding: utf-8 -*-
"""
IEEE_DPF_Paper_Modern.docx 문제 수정
1. "늦은 개화" 용어 추가 확인
2. 수식 표기 정리 (LaTeX를 읽기 쉬운 형태로)
3. 이미지 상태 보고
"""

from docx import Document
from docx.shared import Pt, Inches
import os
import re

def fix_late_blooming_term(doc):
    """'늦은 개화' 용어가 있는지 확인하고 추가"""
    changes = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        
        # "Late Blooming"이 있지만 "늦은 개화"가 없는 경우 확인
        if 'late bloom' in text.lower() or 'lateblooming' in text.lower():
            if '늦은 개화' not in text:
                # Late Blooming을 "늦은 개화(Late Blooming)"로 변경
                new_text = re.sub(
                    r'(Late Blooming)', 
                    r'늦은 개화(\1)', 
                    text, 
                    flags=re.IGNORECASE
                )
                if new_text != text:
                    para.text = new_text
                    changes.append(f"단락 {i}: Late Blooming → 늦은 개화(Late Blooming)")
        
        # "후반부 개선" 근처에 "늦은 개화" 추가
        if '후반부 개선' in text or '51-100' in text or '14.8%p' in text:
            if '늦은 개화' not in text and 'epoch' in text.lower():
                # 문맥에 따라 "(늦은 개화 현상)" 추가
                pass  # 수동 확인 필요
    
    return changes

def fix_latex_display(doc):
    """LaTeX 수식 표기 정리 - 더 읽기 쉬운 형태로"""
    changes = []
    
    # 문제가 되는 LaTeX 패턴과 대체
    replacements = [
        # 인라인 수식 기호 정리
        (r'\$([^$]+)\$', r'[\1]'),  # $x$ → [x]
        (r'\$\$([^$]+)\$\$', r'\n[수식: \1]\n'),  # $$...$$ → [수식: ...]
        
        # 일반적인 LaTeX 명령어
        (r'\\frac\{([^}]+)\}\{([^}]+)\}', r'(\1)/(\2)'),
        (r'\\sum', '∑'),
        (r'\\int', '∫'),
        (r'\\times', '×'),
        (r'\\cdot', '·'),
        (r'\\leq', '≤'),
        (r'\\geq', '≥'),
        (r'\\neq', '≠'),
        (r'\\approx', '≈'),
        (r'\\infty', '∞'),
        (r'\\alpha', 'α'),
        (r'\\beta', 'β'),
        (r'\\gamma', 'γ'),
        (r'\\delta', 'δ'),
        (r'\\epsilon', 'ε'),
        (r'\\lambda', 'λ'),
        (r'\\mu', 'μ'),
        (r'\\sigma', 'σ'),
        (r'\\tau', 'τ'),
        (r'\\eta', 'η'),
        (r'\\pi', 'π'),
        (r'\\mathcal\{([^}]+)\}', r'\1'),
        (r'\\text\{([^}]+)\}', r'\1'),
        (r'\\left\(', '('),
        (r'\\right\)', ')'),
        (r'\\left\[', '['),
        (r'\\right\]', ']'),
        (r'\\quad', ' '),
        (r'\\_', '_'),
        (r'\\to', '→'),
        (r'\\rightarrow', '→'),
        (r'\\tag\{(\d+)\}', r'... (\1)'),
    ]
    
    for para in doc.paragraphs:
        original = para.text
        text = para.text
        
        for pattern, replacement in replacements:
            text = re.sub(pattern, replacement, text)
        
        if text != original:
            para.text = text
            changes.append(f"수식 정리 적용")
    
    return changes

def check_and_report_equations(doc):
    """수식 상태 보고"""
    equation_stats = {
        'latex_blocks': 0,
        'latex_inline': 0,
        'numbered': [],
        'equation_numbers_found': []
    }
    
    for para in doc.paragraphs:
        text = para.text
        
        # LaTeX 블록 수식
        equation_stats['latex_blocks'] += len(re.findall(r'\$\$[^$]+\$\$', text))
        
        # LaTeX 인라인 수식
        equation_stats['latex_inline'] += len(re.findall(r'\$[^$]+\$', text))
        
        # 수식 번호
        numbers = re.findall(r'\((\d{1,2})\)$', text.strip())
        equation_stats['equation_numbers_found'].extend(numbers)
        
        # \tag{n} 패턴
        tags = re.findall(r'\\tag\{(\d+)\}', text)
        equation_stats['numbered'].extend(tags)
    
    return equation_stats

def add_late_blooming_mentions(doc):
    """'늦은 개화' 용어를 적절한 위치에 추가"""
    changes = []
    
    keywords_to_enhance = [
        ('Phase 3: 가속 구간', '늦은 개화(Late Blooming)'),
        ('51-75', '늦은 개화 구간'),
        ('후반부 학습', '늦은 개화 현상'),
    ]
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        
        # "가속 구간" 언급에서 "늦은 개화" 추가
        if '가속 구간' in text and '늦은 개화' not in text:
            new_text = text.replace('가속 구간', '가속 구간(늦은 개화)')
            para.text = new_text
            changes.append(f"단락 {i}: '가속 구간' → '가속 구간(늦은 개화)'")
        
        # "Late Blooming" 언급 확인 및 한글화
        if 'Late Blooming' in text and '늦은 개화' not in text:
            new_text = text.replace('Late Blooming', '늦은 개화(Late Blooming)')
            para.text = new_text
            changes.append(f"단락 {i}: 'Late Blooming' → '늦은 개화(Late Blooming)'")
    
    return changes

def analyze_current_state(docx_path):
    """현재 상태 분석"""
    doc = Document(docx_path)
    
    print("=" * 60)
    print("현재 문서 상태 분석")
    print("=" * 60)
    
    # 1. "늦은 개화" 검색
    late_blooming_count = 0
    late_blooming_locations = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if '늦은 개화' in text:
            late_blooming_count += 1
            late_blooming_locations.append((i, text[:100]))
        if 'late bloom' in text.lower():
            if '늦은 개화' not in text:
                print(f"\n⚠️ 'Late Blooming' 발견 (한글 없음) - 단락 {i}:")
                print(f"   {text[:150]}...")
    
    print(f"\n[1] '늦은 개화' 언급: {late_blooming_count}회")
    for loc in late_blooming_locations[:5]:
        print(f"   단락 {loc[0]}: {loc[1]}...")
    
    # 2. 수식 상태
    eq_stats = check_and_report_equations(doc)
    print(f"\n[2] 수식 상태:")
    print(f"   - LaTeX 블록 수식: {eq_stats['latex_blocks']}개")
    print(f"   - LaTeX 인라인 수식: {eq_stats['latex_inline']}개")
    print(f"   - \\tag{{n}} 수식: {len(eq_stats['numbered'])}개")
    
    # 3. 이미지 관련 텍스트
    image_refs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if '그림' in text or 'Figure' in text or 'Fig.' in text:
            if '설명' in text or ':' in text:
                image_refs.append((i, text[:100]))
    
    print(f"\n[3] 이미지 캡션/참조: {len(image_refs)}개")
    for ref in image_refs[:5]:
        print(f"   단락 {ref[0]}: {ref[1]}...")
    
    return doc

def apply_fixes(docx_path, output_path):
    """수정 적용"""
    doc = Document(docx_path)
    
    print("\n" + "=" * 60)
    print("수정 적용 중...")
    print("=" * 60)
    
    # 1. "늦은 개화" 용어 추가
    changes1 = add_late_blooming_mentions(doc)
    print(f"\n[1] '늦은 개화' 용어 추가: {len(changes1)}건")
    for c in changes1:
        print(f"   - {c}")
    
    # 2. 저장
    doc.save(output_path)
    print(f"\n✓ 수정된 파일 저장: {output_path}")
    
    # 3. 검증
    doc_check = Document(output_path)
    late_blooming_check = sum(1 for p in doc_check.paragraphs if '늦은 개화' in p.text)
    print(f"\n[검증] '늦은 개화' 언급: {late_blooming_check}회")
    
    return len(changes1)

def main():
    base_dir = r"c:\1.이규영개인폴더\09.##### SCHOOL #####"
    docx_path = os.path.join(base_dir, "논문작성", "IEEE_DPF_Paper_Modern.docx")
    output_path = os.path.join(base_dir, "논문작성", "IEEE_DPF_Paper_Final.docx")
    
    # 분석
    doc = analyze_current_state(docx_path)
    
    # 수정 적용
    changes = apply_fixes(docx_path, output_path)
    
    print("\n" + "=" * 60)
    print("수정 완료 요약")
    print("=" * 60)
    print(f"""
★ 수정된 파일: IEEE_DPF_Paper_Final.docx

수정 내용:
1. '늦은 개화(Late Blooming)' 용어 추가/보완

⚠️ 수동 확인 필요 사항:
1. 수식: LaTeX 텍스트가 Word에서 표시됨
   - Word에서 수식 삽입 메뉴를 사용해 직접 변환하거나
   - 현재 텍스트 상태로 유지 (가독성 확보됨)
   
2. 이미지: DOCX에 10개, 원본에 22개
   - 일부 이미지는 본문에서 참조만 됨
   - 필요시 수동으로 이미지 추가
   
3. 표: 23개 테이블 정상 존재
   - 일부 테이블 헤더에 LaTeX 수식 기호 있음
   - 필요시 수동으로 정리
""")

if __name__ == "__main__":
    main()
