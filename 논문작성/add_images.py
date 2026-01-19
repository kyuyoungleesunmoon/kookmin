from docx import Document
from docx.shared import Inches, Pt
import os
import shutil

print("="*80)
print("이미지 확인 및 추가")
print("="*80)

# 원본 파일 로드
original_doc = Document(r'C:\1.이규영개인폴더\09.##### SCHOOL #####\논문작성\IEEE_DPF_Paper_Draft_v2.docx')

# 이미지 추출
print("\n1. 원본 문서의 이미지 확인...")

image_count = 0
image_rels = []

# 문서의 이미지 관계 확인
for rel_id, rel in original_doc.part.rels.items():
    if "image" in rel.target_ref:
        image_count += 1
        image_rels.append({
            'rel_id': rel_id,
            'target': rel.target_ref,
            'type': rel.reltype
        })
        print(f"  이미지 {image_count}: {rel.target_ref}")

print(f"\n총 {image_count}개 이미지 발견")

# 이미지 파일 추출
print("\n2. 이미지 추출 중...")

temp_img_dir = r'C:\1.이규영개인폴더\09.##### SCHOOL #####\논문작성\temp_images'
os.makedirs(temp_img_dir, exist_ok=True)

extracted_images = []

try:
    # DOCX는 ZIP 파일이므로 압축 해제
    import zipfile
    
    docx_path = r'C:\1.이규영개인폴더\09.##### SCHOOL #####\논문작성\IEEE_DPF_Paper_Draft_v2.docx'
    
    with zipfile.ZipFile(docx_path, 'r') as zip_ref:
        # word/media/ 폴더의 이미지 추출
        for file_info in zip_ref.filelist:
            if file_info.filename.startswith('word/media/'):
                # 이미지 파일 추출
                extracted_path = zip_ref.extract(file_info, temp_img_dir)
                image_filename = os.path.basename(file_info.filename)
                
                # temp_images/word/media/image1.png -> temp_images/image1.png
                final_path = os.path.join(temp_img_dir, image_filename)
                shutil.move(extracted_path, final_path)
                
                extracted_images.append(final_path)
                print(f"  추출: {image_filename}")
        
        # word/media 폴더 삭제
        try:
            shutil.rmtree(os.path.join(temp_img_dir, 'word'))
        except:
            pass

except Exception as e:
    print(f"  추출 오류: {e}")

print(f"\n총 {len(extracted_images)}개 이미지 추출 완료")

# 수정된 문서에 이미지 추가
print("\n3. 수정된 문서에 이미지 추가...")

revised_doc = Document(r'C:\1.이규영개인폴더\09.##### SCHOOL #####\논문작성\IEEE_DPF_Paper_Revised_v3.docx')

# 이미지를 삽입할 위치 찾기 (그림 번호가 있는 문단 뒤)
added_count = 0

for i, para in enumerate(revised_doc.paragraphs):
    text = para.text.strip()
    
    # "그림 X:" 또는 "Fig X:" 패턴 찾기
    import re
    if re.search(r'그림\s*\d+[:：]', text) or re.search(r'Fig\s*\d+[:：]', text):
        print(f"\n  발견: {text[:50]}...")
        
        # 이미지 삽입 (다음 문단에)
        if added_count < len(extracted_images):
            try:
                # 새 문단 추가
                new_para = revised_doc.paragraphs[i].insert_paragraph_before()
                
                # 이미지 추가
                run = new_para.add_run()
                run.add_picture(extracted_images[added_count], width=Inches(5.5))
                
                # 중앙 정렬
                new_para.alignment = 1  # CENTER
                
                print(f"    → 이미지 {added_count+1} 삽입 완료")
                added_count += 1
                
            except Exception as e:
                print(f"    → 삽입 실패: {e}")

# 남은 이미지들은 문서 끝에 추가
if added_count < len(extracted_images):
    print(f"\n남은 {len(extracted_images) - added_count}개 이미지를 문서 끝에 추가...")
    
    revised_doc.add_page_break()
    revised_doc.add_heading('추가 이미지', level=1)
    
    for idx in range(added_count, len(extracted_images)):
        try:
            para = revised_doc.add_paragraph()
            para.add_run(f'\n그림 {idx+1}\n')
            
            img_para = revised_doc.add_paragraph()
            run = img_para.add_run()
            run.add_picture(extracted_images[idx], width=Inches(5.5))
            img_para.alignment = 1
            
            revised_doc.add_paragraph()
            added_count += 1
            print(f"  이미지 {idx+1} 추가 완료")
            
        except Exception as e:
            print(f"  이미지 {idx+1} 추가 실패: {e}")

# 저장
output_path = r'C:\1.이규영개인폴더\09.##### SCHOOL #####\논문작성\IEEE_DPF_Paper_Revised_Final.docx'
revised_doc.save(output_path)

print("\n" + "="*80)
print(f"✅ 완료!")
print(f"   저장 위치: {output_path}")
print(f"   추가된 이미지: {added_count}개")
print("="*80)

# 임시 폴더 정리 여부 확인
print(f"\n임시 이미지 폴더: {temp_img_dir}")
print("  (수동으로 삭제 가능)")
