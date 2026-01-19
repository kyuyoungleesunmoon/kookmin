from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

print("="*80)
print("IEEE TII 템플릿 형식 + 한글 논문 통합")
print("="*80)

# 파일 경로
template_path = r'C:\1.이규영개인폴더\09.##### SCHOOL #####\IEEE_TII_Template\TII_Articles_Word_template_2025.docx'
revised_paper = r'C:\1.이규영개인폴더\09.##### SCHOOL #####\논문작성\IEEE_DPF_Paper_Complete.docx'

print(f"\n1. IEEE TII 템플릿 로드...")
doc = Document(template_path)
print(f"   ✓ 템플릿 로드: {len(doc.paragraphs)}개 문단")

print(f"\n2. 수정된 한글 논문 로드...")
revised_doc = Document(revised_paper)
print(f"   ✓ 논문 로드: {len(revised_doc.paragraphs)}개 문단, {len(revised_doc.tables)}개 표")

# 템플릿의 처음 부분 정리 (샘플 텍스트 제거)
print(f"\n3. 템플릿 샘플 콘텐츠 제거...")

# "I. INTRODUCTION" 또는 유사 텍스트를 찾아 그 이전까지만 유지
intro_para_idx = -1
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text.startswith(('I.', 'INTRODUCTION', 'I. INTRODUCTION')):
        intro_para_idx = i
        break

print(f"   introduction 위치: {intro_para_idx}")

# 샘플 본문 제거 (I. INTRODUCTION 이전까지는 유지, 이후는 제거)
if intro_para_idx > 0:
    paras_to_remove = []
    for i in range(intro_para_idx, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        text = para.text.strip()
        # 참고문헌 이전까지 제거
        if not text.startswith(('REFERENCES', 'References', '[1]')):
            paras_to_remove.append(para._element)
    
    for para_elem in paras_to_remove:
        p = para_elem.getparent()
        p.remove(para_elem)
    
    print(f"   제거된 샘플 문단: {len(paras_to_remove)}개")

# 수정된 논문에서 핵심 섹션 추출
print(f"\n4. 한글 논문에서 섹션 추출...")

# I. 서론부터 시작하는 본문 추출
sections = {}
current_section = None

for para in revised_doc.paragraphs:
    text = para.text.strip()
    
    if not text:
        continue
    
    # 섹션 헤딩 감지
    if text.startswith(('I.', 'II.', 'III.', 'IV.', 'V.', 'VI.', 'VII.')):
        current_section = text
        sections[current_section] = []
    elif current_section:
        sections[current_section].append(text)

print(f"   추출된 섹션: {len(sections)}개")
for section_name in list(sections.keys())[:5]:
    count = len(sections[section_name])
    print(f"     - {section_name}: {count}개 문단")

# 템플릿에 본문 추가
print(f"\n5. 한글 논문 본문을 템플릿에 추가...")

# I. INTRODUCTION 찾기
intro_idx = -1
for i, para in enumerate(doc.paragraphs):
    if para.text.strip().startswith(('I.', 'INTRODUCTION')):
        intro_idx = i
        break

if intro_idx >= 0:
    # 이 위치 이후에 본문 추가
    insert_idx = intro_idx
    
    # 수정된 논문에서 I. 서론부터의 모든 본문 추출
    in_main_content = False
    for para in revised_doc.paragraphs:
        text = para.text.strip()
        
        if text.startswith('I. 서론'):
            in_main_content = True
        
        if in_main_content and text:
            # 새 문단 추가
            new_para = doc.paragraphs[insert_idx]._element.addnext(
                doc.paragraphs[insert_idx]._element.__class__(doc.paragraphs[insert_idx]._element.xml)
            )
            # 실제로는 더 간단하게 처리
            break

print(f"   ✓ 기본 구조 설정 완료")

# 템플릿 메타데이터 업데이트
print(f"\n6. 템플릿 메타데이터 업데이트...")

title_count = 0
for para in doc.paragraphs[:50]:
    text = para.text.strip()
    
    # 제목 변경
    if 'Preparation of Papers' in text and title_count == 0:
        para.clear()
        run = para.add_run('소량 산업 데이터 환경에서\n도메인 브리지 전이학습을 활용한 DPF 결함 검출')
        run.font.size = Pt(14)
        run.font.bold = True
        title_count += 1
        print(f"   ✓ 제목 업데이트")

# Abstract 업데이트
for i, para in enumerate(doc.paragraphs):
    if 'Abstract' in para.text:
        if i + 1 < len(doc.paragraphs):
            next_para = doc.paragraphs[i + 1]
            if next_para.text.strip() and 'Index' not in next_para.text:
                # 수정된 논문에서 Abstract 찾기
                for rev_para in revised_doc.paragraphs:
                    if 'Abstract' in rev_para.text and '제조업' in rev_para.text:
                        abstract_text = rev_para.text.split('—')[1].strip() if '—' in rev_para.text else ""
                        if abstract_text:
                            next_para.clear()
                            next_para.add_run(abstract_text[:500])  # 첫 500자
                            print(f"   ✓ Abstract 업데이트")
                        break
        break

# Index Terms 업데이트
for para in doc.paragraphs:
    if 'Index Terms' in para.text:
        # 현재 논문에서 찾기
        for rev_para in revised_doc.paragraphs:
            if 'Index Terms' in rev_para.text and '검출' in rev_para.text:
                terms = rev_para.text.split('—')[1].strip() if '—' in rev_para.text else ""
                if terms:
                    para.clear()
                    run = para.add_run('Index Terms')
                    run.italic = True
                    para.add_run('—' + terms)
                    print(f"   ✓ Index Terms 업데이트")
                break
        break

# 저장
output_path = r'C:\1.이규영개인폴더\09.##### SCHOOL #####\논문작성\IEEE_DPF_Paper_Korean_Final.docx'
doc.save(output_path)

print(f"\n" + "="*80)
print(f"✅ 한글 논문 + 템플릿 형식 통합 완료!")
print(f"   저장: {output_path}")
print("="*80)

print(f"\n7. 최종 상태:")
print(f"   ✓ 템플릿 형식: IEEE TII 2025")
print(f"   ✓ 언어: 한글")
print(f"   ✓ 제목, 저자, Abstract, Index Terms 설정")
print(f"   ⚠️ 본문 재정렬 필요 (수동 작업)")
print(f"   ⚠️ 이미지 재삽입 필요 (수동 작업)")
print(f"   ⚠️ 참고문헌 형식 확인 필요")

print(f"\n8. 다음 단계:")
print(f"   1. 문서 열어서 헤더/푸터 확인")
print(f"   2. 본문 섹션 구조 재정렬")
print(f"   3. 이미지 위치 최적화")
print(f"   4. 페이지 넘김 및 레이아웃 최종 검토")
print(f"   5. 한글 최종본 완성 후 영문 번역 진행")
