from docx import Document
from docx.shared import Pt, Inches
import re
import os
import copy

print("="*80)
print("IEEE TII 논문 양식 자동 수정")
print("="*80)

# 파일 경로
paper_path = r'C:\1.이규영개인폴더\09.##### SCHOOL #####\논문작성\IEEE_DPF_Paper_Draft_v2.docx'
output_path = r'C:\1.이규영개인폴더\09.##### SCHOOL #####\논문작성\IEEE_DPF_Paper_Corrected.docx'

# 문서 로드
print("\n1. 문서 로드...")
doc = Document(paper_path)
print(f"   ✓ {len(doc.paragraphs)}개 문단, {len(doc.tables)}개 표")

# ================== 과도한 표현 수정 ==================
print("\n2. 과도한 표현 수정...")

corrections = {
    # 과장된 표현
    '성능 폭발': '성능 가속',
    '늦은 개화': '후반부 성능 향상',
    '치명적 결과': '심각한 영향',
    '치명적 결과를 초래': '심각한 영향을 미칠',
    '시너지 발현': '상승 효과',
    '압도적': '우수한',
    '혁신적': '향상된',
    '획기적': '효과적인',
    '폭발적': '급격한',
    
    # 불필요한 기호
    '(2배 가속!)': '(약 2배)',
    '!)': ')',
    ' ⭐': '',
    ' ★': '',
    '⚠️': '',
    '✅': '',
    '🔥': '',
    
    # 화살표 정리
    ' ← Best checkpoint': ' (Best checkpoint)',
    ' ←, 이는 조기 종료 시 놓치게 되는 구간이다.': ' - 조기 종료 시 놓치게 되는 구간',
    ' ←, 이는 최종 성능 도약 구간이다.': ' - 최종 성능 도약 구간',
    '←에서     ': '에서 ',
    '← ': '',
    ' ←': '',
    
    # 코드 스타일 주석 정리
    '# 충분한 수렴 시간 확보 ★': '# 충분한 수렴 시간 확보',
    '# Stage 1보다 높은 인내 ★': '# Stage 1보다 높은 인내',
}

correction_count = 0

for para in doc.paragraphs:
    original_text = para.text
    modified_text = original_text
    
    for old_expr, new_expr in corrections.items():
        if old_expr in modified_text:
            modified_text = modified_text.replace(old_expr, new_expr)
    
    # 텍스트가 변경되었으면 적용
    if modified_text != original_text:
        # 문단의 모든 run을 합쳐서 새 텍스트로 교체
        for run in para.runs:
            run.text = ''
        if para.runs:
            para.runs[0].text = modified_text
        else:
            para.add_run(modified_text)
        correction_count += 1

print(f"   ✓ {correction_count}개 문단 수정")

# ================== 수식 앞 콜론 추가 ==================
print("\n3. 수식 배치 개선...")

equation_fixes = 0
prev_para = None

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    # 수식 시작 ($$로 시작)
    if text.startswith('$$'):
        if prev_para:
            prev_text = prev_para.text.strip()
            # 이전 문단이 콜론으로 끝나지 않고, 빈 줄이 아닐 때
            if prev_text and not prev_text.endswith(':') and not prev_text.endswith('다:') and not prev_text.endswith('다.'):
                # 콜론으로 끝나지 않는 문장은 그대로 유지 (수동 검토 필요)
                pass
    
    prev_para = para

print(f"   ✓ 수식 배치 검토 완료 (수동 확인 권장)")

# ================== 참고문헌 형식 수정 ==================
print("\n4. 참고문헌 형식 수정...")

refs_fixed = 0
for para in doc.paragraphs:
    text = para.text.strip()
    
    # Roboflow 참고문헌 수정
    if text.startswith('[39]') and 'Roboflow' in text:
        new_text = '[39] Roboflow, Inc., "X-ray Defects Dataset v5," Roboflow Universe, 2023. [Online]. Available: https://universe.roboflow.com/dataset'
        for run in para.runs:
            run.text = ''
        if para.runs:
            para.runs[0].text = new_text
        refs_fixed += 1
    
    if text.startswith('[40]') and 'Roboflow' in text:
        new_text = '[40] Roboflow, Inc., "Casting Defects Dataset v1," Roboflow Universe, 2023. [Online]. Available: https://universe.roboflow.com/dataset'
        for run in para.runs:
            run.text = ''
        if para.runs:
            para.runs[0].text = new_text
        refs_fixed += 1

print(f"   ✓ {refs_fixed}개 참고문헌 수정")

# ================== 표에서도 수정 ==================
print("\n5. 표 내용 수정...")

table_fixes = 0
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            original = cell.text
            modified = original
            
            for old_expr, new_expr in corrections.items():
                if old_expr in modified:
                    modified = modified.replace(old_expr, new_expr)
            
            if modified != original:
                # 셀 내용 교체
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = ''
                    if para.runs:
                        para.runs[0].text = modified
                    else:
                        para.add_run(modified)
                table_fixes += 1

print(f"   ✓ {table_fixes}개 표 셀 수정")

# ================== 저장 ==================
doc.save(output_path)

print("\n" + "="*80)
print("✅ 수정 완료!")
print("="*80)

print(f"\n저장 위치: {output_path}")

print(f"\n수정 요약:")
print(f"   - 과도한 표현: {correction_count}개 문단")
print(f"   - 참고문헌: {refs_fixed}개")
print(f"   - 표 내용: {table_fixes}개 셀")

print(f"\n⚠️ 수동 확인 권장:")
print(f"   1. 수식 앞 문장이 자연스러운지 확인")
print(f"   2. 화살표(←) 제거 후 문맥 확인")
print(f"   3. 전체적인 톤 일관성 검토")
print(f"   4. 이미지 위치 확인")
