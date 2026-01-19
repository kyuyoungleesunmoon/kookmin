
import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- Configuration ---
MARKDOWN_FILE = r"C:\1.이규영개인폴더\09.##### SCHOOL #####\converted_md\이규영_국민대_DPF_논문.md"
TEMPLATE_FILE = r"C:\1.이규영개인폴더\09.##### SCHOOL #####\IEEE_TII_Template\TII_Articles_Word_template_2025.docx"
OUTPUT_FILE = r"C:\1.이규영개인폴더\09.##### SCHOOL #####\논문작성\IEEE_DPF_Paper_Draft_v2.docx"
IMAGE_BASE_DIR = r"C:\1.이규영개인폴더\09.##### SCHOOL #####\converted_md"

# Sections to exclude (Source Code, Raw Data, Redundant Images)
SKIP_KEYWORDS = ["완전한 학습 코드", "데이터셋 설정", "모니터링 및 검증", "소스 코드", "패키지 버전", "requirements.txt", "추출된 이미지"]

def clean_markdown_formatting(text):
    """Remove **bold** formatting for simple text replacement."""
    return text.replace('**', '')

def parse_markdown():
    """Parses the markdown file into metadata and sections."""
    with open(MARKDOWN_FILE, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    data = {
        "title": "Domain-Bridged Transfer Learning Framework for Manufacturing Defect Detection", # Default/Fallback
        "korean_title": "",
        "author": "Kyu-Young Lee",
        "affiliation": "",
        "date": "",
        "abstract": "",
        "keywords": "Diesel Particulate Filter (DPF), Deep Learning, Transfer Learning, Domain Adaptation, X-ray Inspection, Defect Detection, Smart Factory",
        "body_blocks": [] # List of (type, content) tuples. type: 'header', 'text', 'image', 'table'
    }

    skipping = False
    
    # Simple state machine
    # 0: Metadata (Title/Author) - defined by H1/Line 1-20
    # 1: Abstract - defined by header '### 초록' or 'Abstract'
    # 2: Body - defined by headers '### I.', 'I.', etc.
    
    current_section = "meta"
    
    current_table = []
    in_table = False

    for line in lines:
        line = line.strip()
        
        # --- Metadata Parsing (Heuristic) ---
        if current_section == "meta":
            if line.startswith('# '): # Title
                data['korean_title'] = line[2:] 
                # If we want detailed title parsing, we can check line numbers.
                continue
            if '저자:' in line:
                data['author'] = line.replace('저자:', '').strip()
                continue
            if '소속:' in line:
                data['affiliation'] = line.replace('소속:', '').strip()
                continue
            if '작성일:' in line:
                data['date'] = line.replace('작성일:', '').strip()
                continue
            if '핵심 키워드:' in line:
                data['keywords'] = line.replace('핵심 키워드:', '').strip()
                continue
            if '### 초록' in line or 'Abstract' in line:
                current_section = "abstract"
                continue
            if '### I. 서론' in line or 'Introduction' in line:
                current_section = "body"
                # Fall through to process this header in body
        
        # --- Skipping Logic ---
        if line.startswith('#'):
            header_text = line.lstrip('#').strip()
            if any(k in header_text for k in SKIP_KEYWORDS):
                print(f"[Skip] Entering skipped section: {header_text}")
                skipping = True
                continue
            
            # Resume on standard sections (I, II, III, IV, A, B)
            # Or specialized sub-sections that are NOT in skip list
            if skipping:
                # Resume if it looks like a paper section "I. ", "II. ", "A. "
                if re.match(r'^(I|II|III|IV|V|VI|VII|A|B)[.]\s', header_text) or \
                   '결론' in header_text or 'Related Work' in header_text:
                    print(f"[Resume] Resuming at section: {header_text}")
                    skipping = False
        
        if skipping:
            continue

        # --- Abstract Parsing ---
        if current_section == "abstract":
            # Continue capturing abstract until we hit Body start
            # The 'skipping' logic handles excluding unwanted headers, but we want to capture text.
            if '### I. 서론' in line or 'Introduction' in line:
                current_section = "body"
                # Fall through to process header
            elif line.startswith('#'): 
                # Subheaders in abstract (e.g. #### 핵심 방법론) should be kept as text or bold text?
                # IEEE abstract is usually one paragraph. We will append subheaders as Bold text.
                header_content = line.lstrip('#').strip()
                data['abstract'] += "\n**" + header_content + "** "
            elif line:
                data['abstract'] += line + " "

        # --- Body Parsing ---
        if current_section == "body":
            if not line:
                if in_table:
                    data['body_blocks'].append(('table', current_table))
                    current_table = []
                    in_table = False
                continue

            # Table
            if line.startswith('|'):
                in_table = True
                current_table.append(line)
                continue
            
            if in_table: # Broken table
                data['body_blocks'].append(('table', current_table))
                current_table = []
                in_table = False
            
            # Header
            if line.startswith('#'):
                level = line.count('#')
                text = line.lstrip('#').strip()
                data['body_blocks'].append(('header', (level, text)))
            
            # Image
            elif line.startswith('![') and '](' in line:
                alt = line[line.find('[')+1 : line.find(']')]
                path = line[line.find('](')+2 : line.find(')')]
                data['body_blocks'].append(('image', (alt, path)))
            
            # Text
            else:
                data['body_blocks'].append(('text', line))
    
    # Process Title: Combine KR and EN if possible or use KR
    if data['korean_title']:
        # Extract English subtitle if available in raw file text (hardcoded in extraction for now to keep simple)
        pass

    return data

def replace_paragraph_text(paragraph, new_text):
    """Replaces text in a paragraph while preserving the first run's style if possible."""
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    
    # Warning: specific character formatting might be lost, but paragraph style is kept.
    paragraph.text = new_text

def process_word_conversion():
    print("Parsing Source Markdown...")
    parsed_data = parse_markdown()
    
    print(f"Loading Template: {TEMPLATE_FILE}")
    doc = Document(TEMPLATE_FILE)
    
    # 1. Replace Title
    # Heuristic: Find specific placeholder text
    print("Replacing Title & Author...")
    title_set = False
    author_set = False
    
    for p in doc.paragraphs:
        if "Preparation of Papers for" in p.text:
            # This is likely the title placeholder or pre-title
            replace_paragraph_text(p, parsed_data['korean_title'])
            title_set = True
            continue
        
        # Often title is "Preparation..." in IEEE templates. 
        # If the template has dummy title "Title of the Paper", we look for that.
        
        if "First A. Author" in p.text:
            replace_paragraph_text(p, parsed_data['author'])
            author_set = True
            continue

        if "Abstract—" in p.text:
            # IEEE Abstract format: "Abstract—Content..."
            # We want to keep "Abstract—" bold/italic but replace content.
            # Simplified: Just replace text.
            replace_paragraph_text(p, "Abstract—" + parsed_data['abstract'])
            continue

        if "This paragraph of the first footnote" in p.text:
            # Replace footnote with affiliation info
            footnote_text = f"Manuscript received {parsed_data.get('date', 'October 14, 2025')}.\n"
            footnote_text += f"{parsed_data.get('author', 'Author')} is with {parsed_data.get('affiliation', 'Kookmin University')}."
            replace_paragraph_text(p, footnote_text)
            continue

        if "Index Terms—" in p.text:
            # Replace Keywords
            replace_paragraph_text(p, "Index Terms—" + parsed_data.get('keywords', ''))
            continue
            
    # 3. Footnote Replacement (Standard Footnotes XML)
    # Replaced by post-processing step to handle stubborn XML
            
    # 2. Body Replacement
    # Strategy: Find "Introduction" header. Delete everything after. Append parsed body.
    print("Replacing Body Content...")
    start_index = -1
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Introduction" or p.text.strip() == "I. Introduction":
            start_index = i
            break
            
    if start_index != -1:
        # Delete paragraphs from end to start_index (reverse to avoid index shift issues)
        # Note: python-docx doesn't support list deletion easily on doc.paragraphs.
        # We have to access the xml element.
        
        total_p = len(doc.paragraphs)
        print(f"Found Introduction at index {start_index}. Removing {total_p - start_index} dummy paragraphs.")
        
        # Remove elements
        # Identify the parent body element
        body_ele = doc._body._body
        
        # Using lxml to remove elements
        # We iterate safely?
        
        # Easier way: Just clear the text of subsequent paragraphs? No, that leaves empty lines.
        # Direct XML removal:
        ps_to_remove = doc.paragraphs[start_index:]
        for p in ps_to_remove:
            p._element.getparent().remove(p._element)
            
    else:
        print("Warning: 'Introduction' header not found in template. Appending to end.")

    # 3. Insert New Body
    insert_body_blocks(doc, parsed_data['body_blocks'])

    print(f"Saving to {OUTPUT_FILE}...")
    doc.save(OUTPUT_FILE)
    
    # 4. Post-Process Footnotes (Direct XML Patch)
    print("Post-processing footnotes...")
    try:
        post_process_footnotes(OUTPUT_FILE, parsed_data)
    except Exception as e:
        print(f"Error in post-processing: {e}")
        
    print("Done.")

def insert_body_blocks(doc, blocks):
    for kind, content in blocks:
        if kind == 'header':
            level, text = content
            # Map Markdown level to Word Style
            # H1 -> Heading 1
            # H2 -> Heading 2
            # H3 -> Heading 3
            # IEEE Template might have specific names like 'Heading 1', 'Heading 2'
            
            if level == 1: style = 'Heading 1' # Often Section Header (I, II...)
            elif level == 2: style = 'Heading 1' # Map ## to Heading 1 to start sections
            elif level == 3: style = 'Heading 2' # Map ### to Heading 2
            else: style = 'Heading 3'
            
            p = doc.add_paragraph(text)
            try: p.style = style
            except: pass
            
        elif kind == 'text':
            # Handle Bold **text**
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', content)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

        elif kind == 'image':
            alt, rel_path = content
            abs_path = os.path.join(IMAGE_BASE_DIR, rel_path)
            if os.path.exists(abs_path):
                print(f"Inserting Image: {rel_path}")
                try:
                    # Insert Image
                    doc.add_picture(abs_path, width=Inches(3.4)) # Column width ~3.5 inches
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Insert Caption
                    caption = doc.add_paragraph(f"{alt}")
                    caption.style = 'Caption'
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    doc.add_paragraph(f"[Image Error: {alt}]")
            else:
                doc.add_paragraph(f"[Missing Image: {alt}]")

        elif kind == 'table':
            create_word_table(doc, content)



def create_word_table(doc, table_lines):
    rows = []
    for line in table_lines:
        if '---' in line: continue
        cells = [c.strip() for c in line.split('|')]
        if len(cells) > 2 and cells[0] == '' and cells[-1] == '':
             cells = cells[1:-1]
        rows.append(cells)
    
    if not rows: return

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    
    # Try to set style, but force borders anyway
    # Try to set style, but force borders anyway
    try: 
        table.style = 'Table Grid'
    except:
        try: table.style = 'Normal Table'
        except: pass
    
    # Force borders for IEEE style (Top/Bottom specific)
    # But user asked for "All borders gone" -> So let's ensure at least basic grid or IEEE style is visible.
    # IEEE usually has Top/Bottom of table and Bottom of Header.
    # Let's apply a simple grid first as requested to "fix missing borders".
    set_table_borders(table)
    
    for r, row_data in enumerate(rows):
        row = table.rows[r]
        for c, text in enumerate(row_data):
            if c < len(row.cells):
                row.cells[c].text = text


def post_process_footnotes(docx_path, data):
    """
    Directly modifies the word/footnotes.xml in the docx zip file
    to handle cases where python-docx fails to find text due to XML structure.
    """
    import zipfile
    import tempfile
    import shutil

    # Prepare new text
    date_str = data.get('date', '2025년 10월 14일')
    author_str = data.get('author', 'Author')
    affil_str = data.get('affiliation', 'Kookmin University')
    new_text = f"Manuscript received {date_str}. {author_str} is with {affil_str}."

    # Phrases to remove or replace
    # We replace the first one with our content, and others with empty string
    target_start = "This paragraph of the first footnote"
    targets_to_remove = [
        "The next few paragraphs should contain",
        "S. B. Author, Jr., was with",
        "T. C. Author is with"
    ]
    
    temp_dir = tempfile.mkdtemp()
    temp_zip = os.path.join(temp_dir, "temp.docx")
    
    try:
        # Create a new zip file by copying content
        with zipfile.ZipFile(docx_path, 'r') as zin:
            with zipfile.ZipFile(temp_zip, 'w') as zout:
                for item in zin.infolist():
                    content = zin.read(item.filename)
                    
                    if item.filename == 'word/footnotes.xml':
                        # Decode
                        xml_str = content.decode('utf-8')
                        
                        # Replace logic match
                        if target_start in xml_str:
                            print(f"[Post-Process] Patching footnotes.xml...")
                            
                            # Replace the main sentence
                            # Regex or simple string replace? 
                            # Simple replace works if the text is contiguous (which we confirmed via zip check previously)
                            # To be safe, we split by the target string.
                            # But we want to remove the WHOLE dummy paragraph.
                            # "This paragraph ... under Grant BS123456.”"
                            
                            # Let's just replace the unique starting phrase with our text
                            # and hope the rest of the sentence doesn't look too weird if not removed?
                            # No, we should try to remove the rest.
                            # But exact matching long strings is risky.
                            # Let's replace "This paragraph of the first footnote" with our New Text.
                            # And replace "will contain the date on which you submitted your paper for review." with ""
                            
                            xml_str = xml_str.replace(target_start, new_text)
                            xml_str = xml_str.replace("will contain the date on which you submitted your paper for review.", "")
                            xml_str = xml_str.replace("It will also contain support information, including sponsor and financial support acknowledgment.", "")
                            
                            for t in targets_to_remove:
                                # We can't replace logical text easily if we don't know exact chunks.
                                # But let's try to find them.
                                if t in xml_str:
                                    # Replace the sentence part we found
                                    xml_str = xml_str.replace(t, "")
                                    
                        content = xml_str.encode('utf-8')
                        
                    zout.writestr(item, content)
        
        # Replace original file
        shutil.move(temp_zip, docx_path)
        print("[Post-Process] Footnotes patched successfully.")
        
    finally:
        shutil.rmtree(temp_dir)

def set_table_borders(table):
    """
    Sets the table borders using OXML to allow specific formatting 
    when standard styles are missing or broken.
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    
    # Check if tblBorders exists, if not create it
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
    
    # Define border types
    # Top: Single, thick
    # Bottom: Single, thick
    # InsideH: Single (for rows)
    # InsideV: None (IEEE style usually has no vertical borders)
    # But user said "All borders gone", so they might prefer standard grid.
    # Let's stick to IEEE standard: Top/Bottom strong, Header weak.
    
    # Helper to create border element
    def create_border(name, val="single", sz="4", space="0", color="auto"):
        el = OxmlElement(f'w:{name}')
        el.set(qn('w:val'), val)
        el.set(qn('w:sz'), sz) # 4 = 1/2 pt
        el.set(qn('w:space'), space)
        el.set(qn('w:color'), color)
        return el

    # Clear existing
    for child in list(tblBorders):
        tblBorders.remove(child)
        
    # Add IEEE-like borders
    # top
    tblBorders.append(create_border('top', val="single", sz="12")) # 1.5pt
    # bottom
    tblBorders.append(create_border('bottom', val="single", sz="12")) # 1.5pt
    # insideH
    tblBorders.append(create_border('insideH', val="single", sz="4")) # 0.5pt
    # left/right/insideV -> None (omitted means none or inherited)
    
    # Also need to handle Header Row specially? 
    # Hard to do quickly with OXML on table-level. 
    # This sets strict Top/Bottom/InsideH which is clean.

if __name__ == "__main__":
    process_word_conversion()
