from docx import Document
from docx.shared import Pt
import re
import os

# Paths
md_path = r"C:\국민대프로젝트\converted_md\이규영_국민대_DPF_논문.md"
template_path = r"C:\국민대프로젝트\IEEE_TII_Template\TII_Articles_Word_template_2025.docx"
output_path = r"C:\국민대프로젝트\논문작성\IEEE_DPF_Paper_Submission_Ready.docx"

print(f"Reading Markdown: {md_path}")
with open(md_path, 'r', encoding='utf-8') as f:
    lines = f.readlines()

print(f"Opening Template: {template_path}")
doc = Document(template_path)
print(f"Initial Template Paragraphs: {len(doc.paragraphs)}")

# Clear body content safely
section = doc.sections[0]
header = section.header
footer = section.footer
# We only want to clear the Main contents
if len(doc.paragraphs) > 0:
    print("Clearing template body content...")
    doc._body.clear_content()

# Helper to add styled paragraph safely (No duplication)
def add_para(text, style_name=None):
    # Always add the paragraph first with default style (usually Normal)
    p = doc.add_paragraph(text)
    
    if style_name:
        try:
            p.style = style_name
        except KeyError:
            # Style not found, try to map or ignore
            # print(f"Warning: Style '{style_name}' not found. Keeping Normal.")
            pass
        except Exception as e:
            print(f"Warning: Could not apply style '{style_name}': {e}")
            
    return p

print("Processing content...")
for line in lines:
    line = line.strip()
    if not line:
        continue
        
    # Headers
    if line.startswith('# '):
        text = line.replace('# ', '').replace('*', '')
        add_para(text, style_name='Title')
        
    elif line.startswith('## '):
        text = line.replace('## ', '').replace('*', '')
        # Special case for Abstract/Index Terms if we want specific mapping
        add_para(text, style_name='Heading 1') # Or Subtitle if appropriate
            
    elif line.startswith('### '): # e.g. "I. 서론", "초록"
        text = line.replace('### ', '').replace('*', '')
        if "초록" in text or "Abstract" in text:
             add_para(text, style_name='Abstract') 
        elif "I." in text or "II." in text or "III." in text: # Roman numerals likely sections
             add_para(text, style_name='Heading 1')
        else:
             add_para(text, style_name='Heading 1')
             
    elif line.startswith('#### '): # e.g. "A. 연구 배경"
        text = line.replace('#### ', '').replace('*', '')
        add_para(text, style_name='Heading 2')
        
    elif line.startswith('##### '):
        text = line.replace('##### ', '').replace('*', '')
        add_para(text, style_name='Heading 3')
        
    # Images
    elif line.startswith('!['):
        match = re.search(r'!\[(.*?)\]\((.*?)\)', line)
        if match:
            caption = match.group(1)
            img_path = match.group(2)
            if not os.path.isabs(img_path):
                 base_dir = os.path.dirname(md_path)
                 img_full_path = os.path.join(base_dir, img_path)
            else:
                img_full_path = img_path
                
            if os.path.exists(img_full_path):
                try:
                    # Add picture in its own run/paragraph
                    p = doc.add_paragraph()
                    run = p.add_run()
                    run.add_picture(img_full_path, width=Pt(240)) # 3.3 inches roughly
                    p.alignment = 1 # Center
                    
                    # Add Caption
                    add_para(f"Fig. {caption}", style_name='Figure Caption')
                except Exception as e:
                    add_para(f"[Image Error: {caption}]")
            else:
                 add_para(f"[Image Missing: {caption}]")
                 
    # Tables
    elif line.startswith('|'):
        # Simple monospaced text for tables for now, or just Normal
        add_para(line, style_name='Normal')
        
    # Lists
    elif line.startswith('- ') or line.startswith('* '):
         add_para(line[2:], style_name='List Bullet') # Might default to Normal if missing
    elif re.match(r'\d+\.', line):
         add_para(line, style_name='List Number')
         
    # Normal Text and Metadata
    else:
        if line.startswith('저자:') or line.startswith('소속:'):
            add_para(line, style_name='Authors') # Corrected Style Name
        elif line.startswith('제출 목표:') or line.startswith('작성일:') or line.startswith('부제:'):
             add_para(line, style_name='Normal') # Keep metadata visible but normal
        else:
            add_para(line, style_name='Normal')

print(f"Saving to: {output_path}")
doc.save(output_path)
print("Done.")
