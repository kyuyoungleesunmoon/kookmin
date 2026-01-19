# -*- coding: utf-8 -*-
"""
원본 파일에서 소스코드 블록 내용 추출
"""

from docx import Document
import os
import re

def extract_code_blocks(docx_path):
    """소스코드 블록 추출"""
    doc = Document(docx_path)
    
    code_blocks = []
    current_block = []
    in_code = False
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # 코드 블록 시작 패턴
        if re.search(r'^\w+\s*=\s*[\{\[]', text) or text.startswith('stage') or text.startswith('augmentation'):
            in_code = True
            current_block = [text]
        elif in_code:
            # 코드 계속
            if re.search(r"'[a-z_]+'\s*:", text) or text in ['}', ']', '},', '],']:
                current_block.append(text)
            elif text.startswith('#') or text == '':
                current_block.append(text)
            else:
                # 코드 블록 종료
                if len(current_block) > 2:
                    code_blocks.append({
                        'start_idx': i - len(current_block),
                        'lines': current_block
                    })
                in_code = False
                current_block = []
    
    return code_blocks

def main():
    base_dir = r"c:\1.이규영개인폴더\09.##### SCHOOL #####"
    input_path = os.path.join(base_dir, "논문작성", "IEEE_DPF_Paper_Final_v2.docx")
    
    doc = Document(input_path)
    
    print("=" * 70)
    print("원본 파일 소스코드 블록 분석")
    print("=" * 70)
    
    # 모든 코드 관련 단락 추출
    code_paragraphs = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # 코드 패턴 감지
        patterns = [
            r"^\w+_config\s*=",  # config 변수
            r"^augmentation\s*=",  # augmentation
            r"'[a-z_]+'\s*:\s*",  # dict key
            r"^\s*\d+\s*$",  # 숫자만
            r"^[\[\{]$",  # 괄호만
            r"^[\]\}],?$",  # 닫는 괄호
        ]
        
        for pattern in patterns:
            if re.search(pattern, text):
                code_paragraphs.append((i, text))
                break
    
    print(f"\n코드 관련 단락 수: {len(code_paragraphs)}\n")
    
    # 그룹핑
    groups = []
    current_group = []
    prev_idx = -10
    
    for idx, text in code_paragraphs:
        if idx - prev_idx <= 2:
            current_group.append((idx, text))
        else:
            if current_group:
                groups.append(current_group)
            current_group = [(idx, text)]
        prev_idx = idx
    
    if current_group:
        groups.append(current_group)
    
    print(f"코드 블록 그룹 수: {len(groups)}\n")
    
    for gi, group in enumerate(groups):
        if len(group) >= 3:
            print(f"\n--- 그룹 {gi+1} (단락 {group[0][0]}-{group[-1][0]}) ---")
            for idx, text in group[:15]:  # 최대 15줄만
                print(f"  [{idx}] {text[:80]}")
            if len(group) > 15:
                print(f"  ... ({len(group)-15}줄 더)")

if __name__ == "__main__":
    main()
