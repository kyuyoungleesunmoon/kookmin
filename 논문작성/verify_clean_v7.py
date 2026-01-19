from docx import Document
import re

file_path = r"C:\국민대프로젝트\논문작성\IEEE_DPF_Paper_Clean_v7.docx"
doc = Document(file_path)

code_patterns = [
    r'^\s*(def |class |import |from |if |for |while |try:|except:|return |print\()',
    r'^\s*\w+\s*=\s*\{',
    r"^\s*'[\w_]+'\s*:",
    r'^\s*plt\.',
    r'^\s*├─|└─',
]

equation_patterns = [
    r'\$\$',
    r'\\tag\{',
    r'\\begin\{',
    r'\\frac\{',
]

code_count = 0
equation_count = 0
code_samples = []
eq_samples = []

for para in doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue
    
    for pattern in code_patterns:
        if re.search(pattern, text):
            code_count += 1
            if len(code_samples) < 5:
                code_samples.append(text[:60])
            break
            
    for pattern in equation_patterns:
        if re.search(pattern, text):
            equation_count += 1
            if len(eq_samples) < 10:
                eq_samples.append(text[:60])
            break

print(f"Remaining code-like blocks: {code_count}")
print(f"Remaining equation blocks: {equation_count}")
print(f"Total paragraphs: {len(doc.paragraphs)}")

if code_samples:
    print("\nCode samples:")
    for s in code_samples:
        print(f"  - {s}...")
        
if eq_samples:
    print("\nEquation samples:")
    for s in eq_samples:
        print(f"  - {s}...")
