from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# 원본 파일 로드
doc = Document(r'C:\1.이규영개인폴더\09.##### SCHOOL #####\논문작성\IEEE_DPF_Paper_Draft_v2.docx')

# 새 문서 생성
new_doc = Document()

# 섹션 설정
section = new_doc.sections[0]
section.page_height = Inches(11)
section.page_width = Inches(8.5)
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(1)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)

print("문서 변환 시작...")

# AI 톤 제거 함수
def remove_ai_tone(text):
    """AI 스타일 표현 제거 및 학술적 표현으로 변환"""
    replacements = {
        '성능 폭발': '성능 가속',
        '늦은 개화': '후반부 성능 향상',
        '시너지 발현': '상승 효과',
        '치명적 결과': '심각한 영향',
        '(2배 가속!)': '(약 2배)',
        '⚠️': '',
        '★': '',
        '🔥': '',
        '→': '→',  # 유지
        '⭐': '',
        '✅': '',
        '←': '',
        ' ⭐': '',
    }
    
    result = text
    for old, new in replacements.items():
        result = result.replace(old, new)
    
    return result

# 제목 추가
title = new_doc.add_paragraph()
title_run = title.add_run('소량 산업 데이터 환경에서의 도메인 브리지 전이학습을 활용한\nDPF 결함 검출 시스템')
title_run.font.size = Pt(14)
title_run.font.bold = True
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 저자 정보
author = new_doc.add_paragraph()
author_run = author.add_run('\n이규영\n국민대학교')
author_run.font.size = Pt(11)
author.alignment = WD_ALIGN_PARAGRAPH.CENTER

new_doc.add_paragraph()

# Abstract 작성
abstract_heading = new_doc.add_paragraph()
abstract_heading.add_run('Abstract').bold = True
abstract_heading.add_run('—제조업에서 딥러닝 기반 결함 검출은 데이터 부족과 도메인 특화성으로 인해 실용화에 어려움을 겪어왔다. 본 연구는 디젤 미립자 필터(DPF) 결함 검출을 위한 도메인 브리지 전이학습 프레임워크를 제안한다. 제안 방법은 ImageNet → X-ray 결함 데이터 → DPF 데이터의 3단계 계층적 전이학습을 통해, 단 339장의 제한된 데이터로 91.7% mAP50 정확도를 달성했다. 이는 직접 학습(56.9%) 대비 34.8%p, ImageNet 직접 전이(72.3%) 대비 19.4%p의 성능 향상을 보여준다. 특히 100 에포크의 충분한 학습을 통해 후반부(51-100 에포크)에서 19.8%p의 추가 개선을 확인했으며, 이는 현대 어텐션 기반 모델의 점진적 수렴 특성을 반영한다. 본 프레임워크는 Intel i5 CPU 환경에서 학습 가능하여 접근성이 높고, 완전한 학습 프로토콜 공개로 재현 가능하며, 다양한 제조 부품으로 확장 가능하다.')

# Index Terms
index_terms = new_doc.add_paragraph()
index_terms.add_run('Index Terms').italic = True
index_terms.add_run('—DPF 결함 검출, 전이학습, 도메인 적응, 제조업 AI, 소량 데이터 학습, YOLO, 품질 관리')

new_doc.add_paragraph()

# 본문 처리
skip_until = None
current_section = None
skip_sections = ['핵심 방법론', '주요 발견', '산업적 기여', '핵심 키워드']

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    # 빈 문단 건너뛰기
    if not text:
        continue
    
    # 제목, 저자, 초기 Abstract 건너뛰기
    if i < 10:
        continue
    
    # 건너뛸 섹션 체크
    if any(skip_sect in text for skip_sect in skip_sections):
        if text in skip_sections:
            skip_until = 'I. 서론'
            continue
    
    if skip_until:
        if text == skip_until:
            skip_until = None
        else:
            continue
    
    # AI 톤 제거
    cleaned_text = remove_ai_tone(text)
    
    # 섹션 번호 감지 (I., II., A., 1. 등)
    is_heading = False
    if re.match(r'^[IVX]+\.\s+', text) or re.match(r'^[A-Z]\.\s+', text):
        is_heading = True
    
    # 문단 추가
    new_para = new_doc.add_paragraph()
    
    if is_heading:
        run = new_para.add_run(cleaned_text)
        run.bold = True
        run.font.size = Pt(11)
    else:
        # 들여쓰기나 특수 형식 유지
        if text.startswith('├─') or text.startswith('└─') or text.startswith('│'):
            # ASCII 트리를 간단한 불릿으로 변환
            cleaned_text = re.sub(r'^[├└│─\s]+', '  • ', cleaned_text)
        
        new_para.add_run(cleaned_text)
        new_para.paragraph_format.line_spacing = 1.15
    
    # 진행 상황 출력
    if i % 200 == 0:
        print(f"처리 중... {i}/{len(doc.paragraphs)}")

print("\n표 복사 중...")

# 표 복사 (간단한 버전)
table_count = 0
for table in doc.tables[:10]:  # 처음 10개 표만
    try:
        new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        new_table.style = 'Table Grid'
        
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_table.rows[i].cells[j].text = cell.text
        
        new_doc.add_paragraph()
        table_count += 1
    except Exception as e:
        print(f"표 {table_count + 1} 처리 중 오류: {e}")

print(f"\n총 {table_count}개 표 복사 완료")

# 저장
output_path = r'C:\1.이규영개인폴더\09.##### SCHOOL #####\논문작성\IEEE_DPF_Paper_Revised_Korean.docx'
new_doc.save(output_path)

print(f"\n✅ 수정된 논문이 저장되었습니다:")
print(f"   {output_path}")
print(f"\n주요 수정 사항:")
print(f"  - 제목 및 저자 정보 정리")
print(f"  - Abstract 재작성 (간결하게)")
print(f"  - 초반 요약 섹션 제거 (핵심 방법론, 주요 발견, 산업적 기여)")
print(f"  - AI 톤 제거 (성능 폭발 → 성능 가속, 이모지 제거)")
print(f"  - ASCII 트리를 불릿 포인트로 변환")
print(f"  - 총 {len(doc.paragraphs)} 문단 → 정리된 버전")
