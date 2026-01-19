# -*- coding: utf-8 -*-
"""
남은 코드 패턴을 모두 서술형으로 변환
"""

from docx import Document
import os
import re

def convert_remaining_code(input_path, output_path):
    """남은 코드 패턴을 서술형으로 변환"""
    doc = Document(input_path)
    
    changes = 0
    
    # 특정 코드 블록들을 서술형으로 변환
    conversions = {
        # 모델 로딩 코드
        "model = YOLO('yolo11s.pt')": "모델은 YOLO11s 사전 학습 가중치를 사용하여 초기화하였다.",
        
        # close_mosaic 설정
        "close_mosaic = 10": "마지막 10 에포크(Epoch 91-100)에서는 Mosaic 증강을 해제하여 실제 이미지에 대한 미세 조정을 수행하였다.",
    }
    
    # 하이퍼파라미터 테이블 영역 서술형 변환
    hyperparams_prose_stage1 = """Stage 1 하이퍼파라미터 설정: 초기 학습률(lr0)은 0.01, 최종 학습률(lrf)은 0.01로 설정하여 cosine decay를 적용하였다. SGD 모멘텀은 0.937(AdamW에서는 beta1), weight decay는 0.0005로 설정하였다. 학습률 워밍업은 3 에포크 동안 적용하였다. 손실 함수 가중치는 Box loss 7.5, Class loss 0.5, DFL loss 1.5로 설정하였다. 조기 종료 patience는 10 에포크로 설정하였으며, Mosaic 증강은 전 학습 과정에서 유지하였다."""
    
    hyperparams_prose_stage2 = """Stage 2 하이퍼파라미터 설정: 학습 에포크는 100회로 Stage 1의 2배로 확장하였다. 대부분의 하이퍼파라미터는 Stage 1과 동일하게 유지하여 학습 안정성을 확보하였다. 다만 조기 종료 patience는 15 에포크로 증가시켜 후반부 성능 개선(늦은 개화 현상)의 기회를 보장하였다. 또한 마지막 10 에포크에서는 Mosaic 증강을 해제하여 실제 이미지 분포에 대한 미세 조정이 이루어지도록 하였다."""
    
    # 단락별 처리
    i = 0
    while i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        text = para.text.strip()
        
        # 직접 변환
        for old_text, new_text in conversions.items():
            if old_text in text:
                para.text = new_text
                changes += 1
                break
        
        # 하이퍼파라미터 테이블 영역 (lr0: 0.01 등)
        if re.match(r'^lr0:\s*\d', text):
            # 이 영역 전체를 서술형으로 변환
            para.text = hyperparams_prose_stage1
            changes += 1
            
            # 후속 코드 라인들 제거
            for j in range(i+1, min(i+15, len(doc.paragraphs))):
                next_text = doc.paragraphs[j].text.strip()
                if re.match(r'^(lrf|momentum|weight_decay|warmup_epochs|box|cls|dfl|patience|close_mosaic):', next_text):
                    doc.paragraphs[j].text = ""
                    changes += 1
                elif next_text.startswith('epochs:'):
                    # Stage 2 시작
                    doc.paragraphs[j].text = hyperparams_prose_stage2
                    changes += 1
        
        # epochs: 100 등 Stage 2 영역
        elif re.match(r'^epochs:\s*100', text):
            para.text = hyperparams_prose_stage2
            changes += 1
            
            # 후속 코드 라인들 제거
            for j in range(i+1, min(i+20, len(doc.paragraphs))):
                next_text = doc.paragraphs[j].text.strip()
                if re.match(r'^(patience|close_mosaic):\s*\d', next_text):
                    doc.paragraphs[j].text = ""
                    changes += 1
        
        # 개별 파라미터 라인 (위에서 처리 안 된 것)
        elif re.match(r'^(lrf|momentum|weight_decay|warmup_epochs|box|cls|dfl|patience|close_mosaic):\s*[\d\.]+', text):
            para.text = ""
            changes += 1
        
        i += 1
    
    # 최종 정리: 남은 코드 패턴
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # logs = { 패턴
        if text == 'logs = {' or text == 'logs = {':
            para.text = "학습 과정 모니터링을 위해 각 에포크별로 손실, 정밀도, 재현율, mAP, 학습률 등의 지표를 기록하였다."
            changes += 1
        
        # 단독 괄호
        if text in ['{', '}', '[', ']', '},', '],']:
            para.text = ""
            changes += 1
    
    print(f"변환 완료: {changes}건")
    
    doc.save(output_path)
    return doc

def final_verification(docx_path):
    """최종 검증"""
    doc = Document(docx_path)
    
    print("\n" + "=" * 60)
    print("최종 검증")
    print("=" * 60)
    
    # 코드 패턴 검색
    code_issues = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text or len(text) > 100:
            continue
        
        # 코드 패턴
        if re.search(r'^[a-z_]+:\s*[\d\.]+\s*#', text):
            code_issues.append((i, text))
        elif re.search(r"^model\s*=\s*YOLO", text):
            code_issues.append((i, text))
        elif re.match(r'^[a-z_]+\s*=\s*\d+\s*#', text):
            code_issues.append((i, text))
    
    if code_issues:
        print(f"\n⚠️ 잔존 코드 패턴: {len(code_issues)}건")
        for idx, text in code_issues[:10]:
            print(f"  [{idx}] {text[:60]}")
    else:
        print("\n✓ 코드 패턴 없음")
    
    # 핵심 내용 확인
    full_text = ' '.join([p.text for p in doc.paragraphs])
    
    checks = [
        ('91.7%', 'mAP50'),
        ('늦은 개화', 'Late Blooming'),
        ('도메인 브리지', '핵심 방법론'),
        ('34.8%p', '총 성능 향상'),
        ('Stage 1 하이퍼파라미터', '서술형 변환'),
        ('Stage 2 하이퍼파라미터', '서술형 변환'),
    ]
    
    print("\n핵심 내용 확인:")
    for term, desc in checks:
        found = term in full_text
        status = "✓" if found else "✗"
        print(f"  {status} {term} ({desc})")
    
    return len(code_issues)

def main():
    base_dir = r"c:\1.이규영개인폴더\09.##### SCHOOL #####"
    input_path = os.path.join(base_dir, "논문작성", "IEEE_DPF_Paper_Final_v4_clean.docx")
    output_path = os.path.join(base_dir, "논문작성", "IEEE_DPF_Paper_Final_v5.docx")
    
    print("=" * 60)
    print("남은 코드 패턴 서술형 변환")
    print("=" * 60)
    
    convert_remaining_code(input_path, output_path)
    
    print(f"\n✓ 저장: {output_path}")
    
    issues = final_verification(output_path)
    
    if issues > 0:
        print("\n추가 정리 수행...")
        # 한번 더 정리
        doc = Document(output_path)
        for para in doc.paragraphs:
            text = para.text.strip()
            if re.search(r'^[a-z_]+:\s*[\d\.]+\s*#', text):
                para.text = ""
            elif re.search(r'^[a-z_]+\s*=\s*\d+\s*#', text):
                para.text = ""
        doc.save(output_path)
        final_verification(output_path)

if __name__ == "__main__":
    main()
