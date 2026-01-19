# -*- coding: utf-8 -*-
"""
IEEE_DPF_Paper_Modern.docx와 원본 마크다운 파일 심층 비교
수식, 표, 이미지, 섹션 구조 상세 비교
"""

import zipfile
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Inches
import os
import re
from collections import defaultdict

def extract_all_text_from_docx(docx_path):
    """DOCX에서 모든 텍스트 추출"""
    doc = Document(docx_path)
    all_text = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            all_text.append(text)
    return all_text

def extract_equations_from_docx(docx_path):
    """DOCX에서 수식 관련 텍스트 추출"""
    doc = Document(docx_path)
    equations = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        # 수식 번호 패턴 (1), (2), ... (20)
        if re.search(r'\(\d{1,2}\)', text):
            equations.append({
                'para_index': i,
                'text': text[:200]
            })
    
    # OOXML에서 수식 요소 찾기
    with zipfile.ZipFile(docx_path, 'r') as z:
        doc_xml = z.read('word/document.xml')
        # oMath 요소 수 계산
        omath_count = doc_xml.count(b'<m:oMath')
        omath_para_count = doc_xml.count(b'<m:oMathPara')
    
    return equations, omath_count, omath_para_count

def extract_tables_from_docx(docx_path):
    """DOCX에서 표 상세 추출"""
    doc = Document(docx_path)
    tables = []
    
    for i, table in enumerate(doc.tables):
        table_data = {
            'index': i,
            'rows': len(table.rows),
            'cols': len(table.columns),
            'header': [],
            'content_preview': []
        }
        
        # 첫 번째 행 (헤더)
        if len(table.rows) > 0:
            table_data['header'] = [cell.text.strip()[:30] for cell in table.rows[0].cells]
        
        # 내용 미리보기
        for row_idx, row in enumerate(table.rows[:5]):
            table_data['content_preview'].append([cell.text.strip()[:50] for cell in row.cells])
        
        tables.append(table_data)
    
    return tables

def check_equation_formatting(docx_path):
    """수식 서식 확인"""
    issues = []
    
    with zipfile.ZipFile(docx_path, 'r') as z:
        doc_xml = z.read('word/document.xml').decode('utf-8')
        
        # 수식 요소 분석
        # 1. oMath 블록 수식
        omath_blocks = len(re.findall(r'<m:oMath>', doc_xml))
        
        # 2. oMathPara (단락 내 수식)
        omath_para = len(re.findall(r'<m:oMathPara>', doc_xml))
        
        # 3. 텍스트 기반 수식 패턴 확인 (LaTeX 스타일 잔존)
        latex_patterns = [
            (r'\$\$', 'LaTeX 블록 수식 ($$ ... $$)'),
            (r'\$[^$]+\$', 'LaTeX 인라인 수식 ($ ... $)'),
            (r'\\frac\{', 'LaTeX frac 명령'),
            (r'\\sum', 'LaTeX sum 명령'),
            (r'\\int', 'LaTeX int 명령'),
            (r'\\mathcal\{', 'LaTeX mathcal 명령'),
            (r'\\begin\{equation\}', 'LaTeX equation 환경'),
            (r'\\tag\{', 'LaTeX tag 명령'),
        ]
        
        for pattern, desc in latex_patterns:
            matches = re.findall(pattern, doc_xml)
            if matches:
                issues.append(f"  ⚠️ {desc} 발견: {len(matches)}개")
    
    return omath_blocks, omath_para, issues

def extract_md_equations(md_path):
    """마크다운에서 수식 추출"""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    equations = {
        'block': [],  # $$ ... $$
        'inline': [],  # $ ... $
        'numbered': [],  # \tag{...}
    }
    
    # 블록 수식
    block_eq = re.findall(r'\$\$([\s\S]*?)\$\$', content)
    equations['block'] = block_eq
    
    # 번호가 붙은 수식 (tag)
    numbered = re.findall(r'\\tag\{(\d+)\}', content)
    equations['numbered'] = sorted(set(numbered), key=int)
    
    return equations

def compare_sections(docx_path, md_path):
    """섹션 구조 비교"""
    doc = Document(docx_path)
    
    # DOCX 섹션 추출
    docx_sections = []
    for para in doc.paragraphs:
        if para.style and 'Heading' in para.style.name:
            docx_sections.append(para.text.strip())
    
    # 마크다운 섹션 추출
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    md_sections = re.findall(r'^#{1,4}\s+(.+)$', content, re.MULTILINE)
    
    return docx_sections, md_sections

def check_key_content_presence(docx_path, md_path):
    """핵심 내용 존재 여부 확인"""
    doc = Document(docx_path)
    docx_text = ' '.join([p.text for p in doc.paragraphs])
    
    # 핵심 키워드/수치 목록
    key_items = [
        ('91.7%', 'mAP50 최종 성능'),
        ('339장', 'DPF 데이터 수'),
        ('310장', 'X-ray 데이터 수'),
        ('34.8%p', '전체 성능 향상'),
        ('19.4%p', '도메인 브리지 효과'),
        ('14.8%p', '후반부 학습 효과'),
        ('72.6%', 'mAP50-95'),
        ('92.8%', 'Precision'),
        ('82.2%', 'Recall'),
        ('YOLO11', '모델명'),
        ('YOLOv8', '비교 모델'),
        ('C2PSA', '어텐션 모듈'),
        ('도메인 브리지', '핵심 방법론'),
        ('늦은 개화', 'Late Blooming'),
        ('ImageNet', '사전학습'),
        ('X-ray', '중간 도메인'),
    ]
    
    results = []
    for keyword, desc in key_items:
        found = keyword in docx_text
        results.append({
            'keyword': keyword,
            'description': desc,
            'found': found
        })
    
    return results

def check_images_in_docx(docx_path):
    """DOCX 내 이미지 확인"""
    with zipfile.ZipFile(docx_path, 'r') as z:
        media_files = [f for f in z.namelist() if f.startswith('word/media/')]
        
        image_info = []
        for f in media_files:
            info = z.getinfo(f)
            image_info.append({
                'name': f,
                'size': info.file_size
            })
    
    return image_info

def validate_equation_numbers(docx_path):
    """수식 번호 순서 검증"""
    doc = Document(docx_path)
    
    equation_numbers = []
    for para in doc.paragraphs:
        text = para.text
        # (1), (2), ... 패턴 찾기
        matches = re.findall(r'\((\d+)\)', text)
        for m in matches:
            if m.isdigit() and int(m) <= 25:  # 수식 번호로 보이는 것만
                equation_numbers.append(int(m))
    
    # 중복 제거 및 정렬
    unique_numbers = sorted(set(equation_numbers))
    
    # 연속성 확인
    expected = list(range(1, max(unique_numbers) + 1)) if unique_numbers else []
    missing = set(expected) - set(unique_numbers)
    
    return unique_numbers, missing

def main():
    base_dir = r"c:\1.이규영개인폴더\09.##### SCHOOL #####"
    docx_path = os.path.join(base_dir, "논문작성", "IEEE_DPF_Paper_Modern.docx")
    md_path = os.path.join(base_dir, "converted_md", "이규영_국민대_DPF_논문.md")
    
    print("=" * 80)
    print("IEEE_DPF_Paper_Modern.docx 심층 분석 및 원본 비교")
    print("=" * 80)
    
    # 1. 기본 정보
    doc = Document(docx_path)
    print(f"\n[1] 기본 정보")
    print(f"   - 총 단락 수: {len(doc.paragraphs)}")
    print(f"   - 총 테이블 수: {len(doc.tables)}")
    
    # 2. 이미지 분석
    print(f"\n[2] 이미지 분석")
    images = check_images_in_docx(docx_path)
    print(f"   - 총 이미지 수: {len(images)}")
    for img in images:
        print(f"      {img['name']}: {img['size']:,} bytes")
    
    # 마크다운 이미지 참조 수 확인
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    md_images = re.findall(r'!\[.*?\]\(images/.*?\)', md_content)
    print(f"   - 마크다운 이미지 참조: {len(md_images)}개")
    
    if len(images) < len(md_images):
        print(f"   ⚠️ 이미지 누락 가능성: DOCX({len(images)}) < MD({len(md_images)})")
    
    # 3. 수식 분석
    print(f"\n[3] 수식 분석")
    omath_blocks, omath_para, issues = check_equation_formatting(docx_path)
    print(f"   - OOXML oMath 블록: {omath_blocks}개")
    print(f"   - OOXML oMathPara: {omath_para}개")
    
    if issues:
        print("   수식 서식 문제:")
        for issue in issues:
            print(issue)
    else:
        print("   ✓ LaTeX 스타일 수식 잔존 없음")
    
    # 마크다운 수식 분석
    md_equations = extract_md_equations(md_path)
    print(f"   - 마크다운 블록 수식: {len(md_equations['block'])}개")
    print(f"   - 마크다운 번호 수식: {md_equations['numbered']}")
    
    # 수식 번호 검증
    eq_numbers, missing = validate_equation_numbers(docx_path)
    print(f"   - DOCX 수식 번호: {eq_numbers[:20]}...")
    if missing:
        print(f"   ⚠️ 누락된 수식 번호: {missing}")
    
    # 4. 테이블 분석
    print(f"\n[4] 테이블 분석")
    tables = extract_tables_from_docx(docx_path)
    for t in tables:
        print(f"   Table {t['index']+1}: {t['rows']} rows x {t['cols']} cols")
        if t['header']:
            print(f"      헤더: {t['header'][:4]}...")
    
    # 5. 섹션 구조 비교
    print(f"\n[5] 섹션 구조 비교")
    docx_sections, md_sections = compare_sections(docx_path, md_path)
    print(f"   - DOCX 헤딩 수: {len(docx_sections)}")
    print(f"   - 마크다운 헤딩 수: {len(md_sections)}")
    
    # 주요 섹션 확인
    expected_sections = [
        'I. 서론', 
        'II. 관련 연구',
        'III. 제안 방법론',
        'IV. 실험 설정',
        'V. 실험 결과',
        'VI. 토론',
        'VII. 결론'
    ]
    
    print("   주요 섹션 존재 여부:")
    for section in expected_sections:
        found = any(section in s for s in docx_sections)
        status = "✓" if found else "✗"
        print(f"      {status} {section}")
    
    # 6. 핵심 내용 확인
    print(f"\n[6] 핵심 내용 존재 확인")
    key_content = check_key_content_presence(docx_path, md_path)
    missing_content = [k for k in key_content if not k['found']]
    present_content = [k for k in key_content if k['found']]
    
    print(f"   ✓ 존재하는 핵심 내용: {len(present_content)}/{len(key_content)}")
    for item in present_content:
        print(f"      ✓ {item['keyword']} ({item['description']})")
    
    if missing_content:
        print(f"\n   ⚠️ 누락된 핵심 내용: {len(missing_content)}/{len(key_content)}")
        for item in missing_content:
            print(f"      ✗ {item['keyword']} ({item['description']})")
    
    # 7. 참고문헌 확인
    print(f"\n[7] 참고문헌 분석")
    docx_text = ' '.join([p.text for p in doc.paragraphs])
    refs = re.findall(r'\[(\d+)\]', docx_text)
    unique_refs = sorted(set([int(r) for r in refs if r.isdigit()]))
    print(f"   - 인용된 참고문헌 수: {len(unique_refs)}")
    if unique_refs:
        print(f"   - 참고문헌 범위: [{min(unique_refs)}] ~ [{max(unique_refs)}]")
    
    # 마크다운 참고문헌 확인
    md_refs = re.findall(r'\[(\d+)\]', md_content)
    unique_md_refs = sorted(set([int(r) for r in md_refs if r.isdigit()]))
    print(f"   - 마크다운 참고문헌: [{min(unique_md_refs)}] ~ [{max(unique_md_refs)}]")
    
    # 8. 종합 평가
    print(f"\n" + "=" * 80)
    print("종합 평가")
    print("=" * 80)
    
    issues_found = []
    
    # 이미지 검사
    if len(images) < 10:
        issues_found.append(f"이미지 부족: {len(images)}개 (마크다운: {len(md_images)}개)")
    
    # 수식 검사
    if omath_blocks + omath_para < 50:
        issues_found.append(f"Word 수식 블록 적음: {omath_blocks + omath_para}개")
    
    # 섹션 검사
    missing_sections = [s for s in expected_sections if not any(s in d for d in docx_sections)]
    if missing_sections:
        issues_found.append(f"누락된 섹션: {missing_sections}")
    
    # 핵심 내용 검사
    if missing_content:
        issues_found.append(f"누락된 핵심 내용: {[m['keyword'] for m in missing_content]}")
    
    if issues_found:
        print("\n⚠️ 발견된 문제:")
        for i, issue in enumerate(issues_found, 1):
            print(f"   {i}. {issue}")
    else:
        print("\n✓ 모든 검사 통과!")
    
    return issues_found

if __name__ == "__main__":
    issues = main()
