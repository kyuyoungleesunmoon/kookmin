from docx import Document
import os

file_path = r"C:\국민대프로젝트\논문작성\IEEE_DPF_Paper_Final_v5.docx"
print(f"Inspecting: {file_path}")

try:
    doc = Document(file_path)
    print(f"Paragraphs: {len(doc.paragraphs)}")
    
    print("\n--- First 20 Paragraphs ---")
    for i, p in enumerate(doc.paragraphs[:20]):
        print(f"{i}: {p.text}")
        
    print("\n--- Text Content Check ---")
    full_text = "\n".join([p.text for p in doc.paragraphs])
    if len(full_text.strip()) < 100:
        print("WARNING: File appears mostly empty.")
    else:
        print(f"Total characters: {len(full_text)}")
        
except Exception as e:
    print(f"Error reading file: {e}")
