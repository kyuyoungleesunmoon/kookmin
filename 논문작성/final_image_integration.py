from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

print("="*80)
print("이미지 최종 통합")
print("="*80)

# 파일 경로
original_docx = r'C:\1.이규영개인폴더\09.##### SCHOOL #####\논문작성\IEEE_DPF_Paper_Draft_v2.docx'
revised_docx = r'C:\1.이규영개인폴더\09.##### SCHOOL #####\논문작성\IEEE_DPF_Paper_Revised_v3.docx'
temp_img_dir = r'C:\1.이규영개인폴더\09.##### SCHOOL #####\논문작성\temp_images'
output_path = r'C:\1.이규영개인폴더\09.##### SCHOOL #####\논문작성\IEEE_DPF_Paper_Complete.docx'

# 추출된 이미지 목록
image_files = [f for f in os.listdir(temp_img_dir) if f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif'))]
image_files.sort()

print(f"\n사용 가능한 이미지: {len(image_files)}개")
for idx, img in enumerate(image_files, 1):
    print(f"  {idx}. {img}")

# 수정된 문서 로드
doc = Document(revised_docx)

# 원본 문서에서 이미지-문단 매핑 정보 수집
print("\n원본 문서의 이미지 위치 분석...")
original_doc = Document(original_docx)

image_positions = []
current_img_idx = 0

for i, para in enumerate(original_doc.paragraphs):
    # 이미지가 포함된 문단인지 확인
    if para._element.xpath('.//pic:pic'):
        # 앞뒤 텍스트 확인
        prev_text = original_doc.paragraphs[i-1].text if i > 0 else ""
        next_text = original_doc.paragraphs[i+1].text if i < len(original_doc.paragraphs)-1 else ""
        
        image_positions.append({
            'img_idx': current_img_idx,
            'prev_text': prev_text[:100],
            'next_text': next_text[:100],
            'para_idx': i
        })
        current_img_idx += 1

print(f"원본에서 {len(image_positions)}개 이미지 위치 발견")

# 새 문서에 이미지 삽입
print("\n이미지 삽입 중...")

inserted_count = 0
img_idx = 0

# 그림 캡션 패턴
fig_patterns = [
    r'그림\s*\d+[:：]',
    r'Fig\s*\d+[:：]',
    r'\[Image Error:',
]

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    # 그림 캡션 찾기
    is_fig_caption = any(re.search(pattern, text) for pattern in fig_patterns)
    
    if is_fig_caption and img_idx < len(image_files):
        print(f"\n  위치 {i}: {text[:60]}...")
        
        try:
            # 캡션 다음에 이미지 삽입
            # 새 문단 생성
            if i + 1 < len(doc.paragraphs):
                next_para = doc.paragraphs[i + 1]
                # 다음 문단이 이미 이미지를 포함하는지 확인
                if not next_para._element.xpath('.//pic:pic'):
                    # 이미지 문단 삽입
                    img_para = para.insert_paragraph_before()
                    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    img_path = os.path.join(temp_img_dir, image_files[img_idx])
                    
                    # WMF 파일 건너뛰기
                    if img_path.lower().endswith('.wmf'):
                        print(f"    → WMF 파일 건너뜀: {image_files[img_idx]}")
                        img_idx += 1
                        if img_idx < len(image_files):
                            img_path = os.path.join(temp_img_dir, image_files[img_idx])
                    
                    # 이미지 추가
                    run = img_para.add_run()
                    run.add_picture(img_path, width=Inches(5.0))
                    
                    print(f"    ✓ 이미지 삽입: {image_files[img_idx]}")
                    inserted_count += 1
                    img_idx += 1
                    
        except Exception as e:
            print(f"    ✗ 삽입 실패: {e}")
            img_idx += 1

# 남은 이미지가 있으면 부록에 추가
if img_idx < len(image_files):
    print(f"\n남은 {len(image_files) - img_idx}개 이미지를 부록에 추가...")
    
    # 부록 섹션 추가
    doc.add_page_break()
    appendix = doc.add_paragraph()
    appendix.add_run('부록: 추가 그림').bold = True
    appendix.paragraph_format.space_after = Pt(12)
    
    while img_idx < len(image_files):
        try:
            img_path = os.path.join(temp_img_dir, image_files[img_idx])
            
            # WMF 건너뛰기
            if img_path.lower().endswith('.wmf'):
                img_idx += 1
                continue
            
            # 캡션
            caption = doc.add_paragraph()
            caption.add_run(f'그림 {img_idx + 1}').bold = True
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 이미지
            img_para = doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = img_para.add_run()
            run.add_picture(img_path, width=Inches(5.0))
            
            doc.add_paragraph()  # 간격
            
            print(f"  부록 이미지 {img_idx + 1} 추가: {image_files[img_idx]}")
            inserted_count += 1
            img_idx += 1
            
        except Exception as e:
            print(f"  부록 이미지 {img_idx + 1} 실패: {e}")
            img_idx += 1

# 저장
doc.save(output_path)

print("\n" + "="*80)
print(f"✅ 완료!")
print(f"   저장: {output_path}")
print(f"   삽입된 이미지: {inserted_count}개 / {len(image_files)}개")
print("="*80)

print("\n다음 단계:")
print("  1. 문서를 열어 이미지 위치 확인")
print("  2. 필요시 이미지 위치 수동 조정")
print("  3. 그림 번호와 캡션 일치 여부 확인")
