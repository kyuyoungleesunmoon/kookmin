# -*- coding: utf-8 -*-
"""
모든 코드 패턴을 완전히 서술형으로 변환
"""

from docx import Document
import re

def complete_conversion():
    doc = Document(r'c:\1.이규영개인폴더\09.##### SCHOOL #####\논문작성\IEEE_DPF_Paper_Final_v5.docx')
    
    changes = 0
    
    # 모든 코드 패턴 변환 매핑
    conversions = {
        # 모델 관련
        "model = model_class(pretrained='imagenet')": "모델은 ImageNet에서 사전 학습된 가중치를 사용하여 초기화하였다.",
        
        # Augmentation 강도
        "augmentation_strength = 'high'": "데이터 증강 강도를 'high'로 설정하여 강한 증강을 적용하였다.",
        "augmentation_strength = 'medium'": "데이터 증강 강도를 'medium'으로 설정하여 중간 수준의 증강을 적용하였다.",
        "augmentation_strength = 'low'": "데이터 증강 강도를 'low'로 설정하여 약한 증강을 적용하였다.",
        
        # 학습 파라미터
        "epochs: 50": "",  # 이미 서술형으로 설명됨
        "batch: 8": "",
        "imgsz: 640": "",
        "warmup_momentum: 0.8": "",
        "warmup_bias_lr: 0.1": "",
    }
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # 직접 매칭
        for old, new in conversions.items():
            if text == old or old in text:
                para.text = new
                changes += 1
                if new:
                    print(f'변환: {old[:40]}... -> {new[:30]}...')
                else:
                    print(f'제거: {old[:40]}...')
                break
        
        # 패턴 매칭으로 나머지 처리
        if re.match(r'^(epochs|batch|imgsz|warmup_momentum|warmup_bias_lr):\s*[\d\.]+', text):
            para.text = ""
            changes += 1
            print(f'제거(패턴): {text[:40]}')
    
    # 저장
    output_path = r'c:\1.이규영개인폴더\09.##### SCHOOL #####\논문작성\IEEE_DPF_Paper_Final_v5.docx'
    doc.save(output_path)
    
    print(f'\n총 {changes}건 처리')
    
    # 재검증
    print('\n재검증...')
    doc2 = Document(output_path)
    
    remaining = []
    for i, para in enumerate(doc2.paragraphs):
        text = para.text.strip()
        if not text or len(text) > 150:
            continue
        
        patterns = [
            r'^[a-z_]+\s*=\s*[\d\.\{\[\'\"]',
            r'^[a-z_]+:\s*[\d\.]+',
            r'^model\s*=',
        ]
        
        for p in patterns:
            if re.match(p, text):
                remaining.append((i, text[:60]))
                break
    
    if remaining:
        print(f'잔존: {len(remaining)}건')
        for idx, text in remaining:
            print(f'  [{idx}] {text}')
    else:
        print('✓ 모든 코드 패턴 처리 완료!')

if __name__ == "__main__":
    complete_conversion()
