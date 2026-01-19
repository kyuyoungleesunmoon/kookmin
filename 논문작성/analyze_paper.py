from docx import Document
from docx.oxml.ns import qn
import re

# DOCX 파일 로드
doc = Document(r'C:\1.이규영개인폴더\09.##### SCHOOL #####\논문작성\IEEE_DPF_Paper_Draft_v2.docx')

print("="*80)
print("논문 구조 분석")
print("="*80)

# 기본 통계
print(f"\n📄 전체 문단 수: {len(doc.paragraphs)}")
print(f"📊 전체 표 수: {len(doc.tables)}")

# 이미지 카운트
image_count = 0
for rel in doc.part.rels.values():
    if "image" in rel.target_ref:
        image_count += 1
print(f"🖼️  전체 이미지 수: {image_count}")

# 섹션 구조 파악
print("\n" + "="*80)
print("섹션 구조")
print("="*80)

sections = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text and (para.style.name.startswith('Heading') or 
                 re.match(r'^[IVX]+\.|^[A-Z]\.|^\d+\.', text)):
        sections.append((i, text[:80]))

print(f"\n발견된 주요 섹션: {len(sections)}개")
for idx, (line_num, section_title) in enumerate(sections[:30]):
    print(f"{idx+1:2d}. Line {line_num:4d}: {section_title}")

# 수식 체크 (LaTeX 표기 확인)
print("\n" + "="*80)
print("수식 분석")
print("="*80)

latex_equations = []
for i, para in enumerate(doc.paragraphs):
    text = para.text
    # LaTeX 수식 패턴 찾기
    if '$$' in text or '$' in text or '\\tag{' in text or '\\begin{' in text:
        latex_equations.append((i, text[:100]))

print(f"\n수식 포함 문단: {len(latex_equations)}개")
if latex_equations:
    print("\n첫 10개 수식 예시:")
    for idx, (line_num, eq_text) in enumerate(latex_equations[:10]):
        print(f"{idx+1}. Line {line_num}: {eq_text}")
else:
    print("⚠️  LaTeX 수식이 발견되지 않았습니다!")

# 표 분석
print("\n" + "="*80)
print("표 분석")
print("="*80)

for idx, table in enumerate(doc.tables[:5]):
    print(f"\n표 {idx+1}:")
    print(f"  - 행 수: {len(table.rows)}")
    print(f"  - 열 수: {len(table.columns)}")
    if table.rows:
        first_row = [cell.text.strip() for cell in table.rows[0].cells]
        print(f"  - 첫 행: {', '.join(first_row[:5])}")

# 전체 텍스트 추출하여 파일 저장
print("\n" + "="*80)
print("전체 내용 추출 중...")
print("="*80)

with open(r'C:\1.이규영개인폴더\09.##### SCHOOL #####\논문작성\paper_content.txt', 'w', encoding='utf-8') as f:
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            f.write(f"[Para {i:4d}] {para.text}\n")

print("\n✅ 전체 내용이 'paper_content.txt'에 저장되었습니다.")
print(f"📝 총 {len([p for p in doc.paragraphs if p.text.strip()])}개 문단 추출 완료")
