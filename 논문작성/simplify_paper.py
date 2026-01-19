# -*- coding: utf-8 -*-
"""
IEEE_DPF_Paper_Clean.docx 추가 정리
1. 남은 소스코드 형태 서술형 변환
2. 불필요한 수식 제거
3. 빈 단락 정리
"""

from docx import Document
import os
import re

def comprehensive_cleanup(docx_path, output_path):
    """종합 정리"""
    doc = Document(docx_path)
    
    changes = {
        'code_converted': 0,
        'equations_simplified': 0,
        'empty_removed': 0,
    }
    
    # 핵심 수식만 유지 (나머지는 본문에 간략히 설명)
    essential_equation_numbers = [3, 7, 8, 9, 10, 11, 12]  # 손실함수, 평가지표, 학습률
    
    # 소스코드 스타일 변환 매핑
    code_to_prose = {
        # 완전한 코드 블록을 서술형으로
        "# YOLOv8: 지수 포화 모델": "",
        "# YOLO11: 시그모이드 + 후반부 가속": "",
        "P_yolov8(t) = 62.3 * (1 - e^(-t/15))": "YOLOv8의 성능 곡선은 지수 포화 모델을 따르며, 약 45 에포크에서 95% 수렴에 도달한다.",
        "P_yolo11(t) = 91.7 / (1 + e^(-(t-60)/12))": "YOLO11의 성능 곡선은 시그모이드 형태로, 약 95 에포크에서 95% 수렴에 도달한다.",
        
        # 데이터셋 YAML 형식
        "path: C:/project/data": "",
        "train: train/images": "",
        "val: valid/images": "",
        "names:": "",
        "nc: 6  # number of classes": "",
        "nc: 2  # number of classes": "",
        
        # 변수 할당 형태
        "SEED = 42": "무작위 시드는 42로 고정하였다.",
        "conf_threshold = 0.001": "",
        "iou_threshold = 0.6": "",
    }
    
    # 제거할 패턴
    patterns_to_remove = [
        # 코드 관련
        r"^\s*0:\s*\w+$",  # 0: crack
        r"^\s*1:\s*\w+$",  # 1: melting
        r"^\s*2:\s*\w+$",  # 등
        r"^\s*#\s*\d+\.\s",  # # 1. 
        r"^\s*FPS\s*=\s*\d+",  # FPS = 1000
        r"^\s*\w+_config\s*=",  # xxx_config =
        r"^\s*├─|└─|│",  # 트리 구조
        r"^\s*\d+:\s*\w+",  # 숫자: 텍스트
        r"^\s*-\s*\w+:\s",  # - key: value
        r"^\s*결과적으로\s",  # 화살표 대체 텍스트
        r"^\s*에서\s*$",  # 화살표 대체
        r"^\s*[A-Z][a-z]+\s*에서\s*[A-Z]",  # X 에서 Y 패턴
        
        # 수식 관련 (불필요한 것들)
        r"\\tag\{1\}",  # 도메인 거리 - 복잡
        r"\\tag\{2\}",  # 삼각 부등식 - 불필요
        r"\\tag\{4\}",  # Box Loss 상세 - 불필요
        r"\\tag\{5\}",  # Class Loss 상세 - 불필요  
        r"\\tag\{6\}",  # DFL Loss 상세 - 불필요
        r"\\tag\{13\}",  # 늦은 개화 모델 - 복잡
        r"\\tag\{14\}",  # C2PSA 손실 분해 - 불필요
        r"\\tag\{15\}",  # 융합 가중치 - 불필요
        r"\\tag\{16\}",  # C2PSA 출력 - 불필요
        r"\\tag\{17\}",  # 가중치 엔트로피 - 불필요
        r"\\tag\{18\}",  # 전이학습 이득 - 불필요
        r"\\tag\{19\}",  # 유사도-성능 - 불필요
        r"\\tag\{20\}",  # ROI - 본문에서 설명
    ]
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        original_text = text
        
        if not text:
            continue
        
        # 1. 직접 매핑 변환
        for code, prose in code_to_prose.items():
            if code in text:
                if prose:
                    para.text = prose
                    changes['code_converted'] += 1
                else:
                    para.text = ""
                    changes['code_converted'] += 1
                text = para.text
                break
        
        # 2. 패턴 기반 제거
        for pattern in patterns_to_remove:
            if re.search(pattern, text):
                # 수식 태그 포함 단락은 내용 단순화
                if '\\tag' in text:
                    # 수식 번호 추출
                    tag_match = re.search(r'\\tag\{(\d+)\}', text)
                    if tag_match:
                        eq_num = int(tag_match.group(1))
                        if eq_num not in essential_equation_numbers:
                            para.text = ""
                            changes['equations_simplified'] += 1
                else:
                    para.text = ""
                break
        
        # 3. LaTeX 수식 단순화 (복잡한 수식 제거)
        if '\\mathcal' in text or '\\left' in text or '\\right' in text:
            if '\\tag' not in text:  # 번호 없는 복잡 수식
                # 간단한 설명으로 대체하거나 제거
                if len(text) > 200:
                    para.text = ""
                    changes['equations_simplified'] += 1
    
    # 연속 빈 단락 정리는 Word에서 수동으로
    
    doc.save(output_path)
    return changes

def further_simplify(docx_path, output_path):
    """추가 단순화 - 코드 스타일 완전 제거"""
    doc = Document(docx_path)
    
    changes = {'simplified': 0}
    
    # 서술형 대체 텍스트
    replacements = [
        # 학습 설정 요약
        (r".*Stage 1 학습 설정.*", ""),
        (r".*Stage 2 학습 설정.*", ""),
        
        # Epoch 테이블 형식
        (r"^Epoch\s+\d+:", ""),
        (r"^\s*\d+\s+\d+\.\d+%", ""),
        
        # 코드 주석 스타일
        (r"^#\s+.+$", ""),
        
        # 변수 할당
        (r"^\w+\s*=\s*['\"]", ""),
        (r"^\w+\s*=\s*\d+", ""),
        (r"^\w+\s*=\s*\[", ""),
        (r"^\w+\s*=\s*\{", ""),
        
        # 기술적 출력 형식
        (r"^├─|^└─|^│", ""),
        (r"^-\s+[A-Za-z]+:", ""),
    ]
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        for pattern, replacement in replacements:
            if re.match(pattern, text):
                para.text = replacement
                changes['simplified'] += 1
                break
    
    doc.save(output_path)
    return changes

def final_cleanup(docx_path, output_path):
    """최종 정리 - 코드 형태 완전 제거"""
    doc = Document(docx_path)
    
    # 코드 스타일로 보이는 모든 패턴
    code_indicators = [
        r"^\s*#",  # 주석
        r"^\s*\w+\s*=\s*",  # 변수 할당
        r"^\s*\d+:\s*[a-z]",  # 인덱스: 값
        r"^\s*[a-z_]+:\s*",  # yaml 스타일
        r"^\s*├|└|│",  # 트리
        r"^\s*\$\w+",  # 변수
        r"^\s*\[\d+\]$",  # 배열 인덱스만
        r"^\s*\.\w+\(",  # 메서드 호출
        r"^\s*if\s+\w+",  # if 문
        r"^\s*for\s+\w+",  # for 문
        r"^\s*def\s+\w+",  # 함수 정의
        r"^\s*class\s+\w+",  # 클래스 정의
        r"^\s*return\s+",  # return 문
        r"^\s*import\s+",  # import 문
        r"^\s*from\s+\w+\s+import",  # from import
        r"^\s*torch\.",  # torch
        r"^\s*np\.",  # numpy
        r"^\s*plt\.",  # matplotlib
        r"^\s*model\.",  # model 메서드
        r"^\s*results\.",  # results
        r"^\s*'[a-z_]+'\s*:",  # 딕셔너리 키
        r'^\s*"[a-z_]+"\s*:',  # 딕셔너리 키
    ]
    
    changes = {'removed': 0}
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        for pattern in code_indicators:
            if re.match(pattern, text, re.IGNORECASE):
                para.text = ""
                changes['removed'] += 1
                break
    
    doc.save(output_path)
    return changes

def main():
    base_dir = r"c:\1.이규영개인폴더\09.##### SCHOOL #####"
    input_path = os.path.join(base_dir, "논문작성", "IEEE_DPF_Paper_Clean.docx")
    temp_path = os.path.join(base_dir, "논문작성", "IEEE_DPF_Paper_Temp.docx")
    output_path = os.path.join(base_dir, "논문작성", "IEEE_DPF_Paper_Simplified.docx")
    
    print("=" * 70)
    print("추가 정리: 수식 단순화 및 코드 완전 제거")
    print("=" * 70)
    
    # 1차 정리
    changes1 = comprehensive_cleanup(input_path, temp_path)
    print(f"\n1차 정리:")
    print(f"  - 코드 변환: {changes1['code_converted']}건")
    print(f"  - 수식 단순화: {changes1['equations_simplified']}건")
    
    # 2차 정리
    changes2 = further_simplify(temp_path, temp_path)
    print(f"\n2차 정리:")
    print(f"  - 추가 단순화: {changes2['simplified']}건")
    
    # 최종 정리
    changes3 = final_cleanup(temp_path, output_path)
    print(f"\n최종 정리:")
    print(f"  - 코드 패턴 제거: {changes3['removed']}건")
    
    print(f"\n✓ 저장 완료: {output_path}")
    
    # 검증
    doc = Document(output_path)
    code_count = 0
    eq_count = 0
    for para in doc.paragraphs:
        text = para.text
        if re.search(r"'[a-z_]+'\s*:|^\s*\w+\s*=\s*\{", text):
            code_count += 1
        if '\\tag' in text:
            eq_count += 1
    
    print(f"\n[검증]")
    print(f"  - 잔존 코드 패턴: {code_count}개")
    print(f"  - 번호 수식: {eq_count}개")
    
    # 임시 파일 삭제
    if os.path.exists(temp_path):
        os.remove(temp_path)

if __name__ == "__main__":
    main()
