from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

print("="*80)
print("한글 논문 상세 수정 시작")
print("="*80)

# 원본 로드
doc = Document(r'C:\1.이규영개인폴더\09.##### SCHOOL #####\논문작성\IEEE_DPF_Paper_Draft_v2.docx')

# 새 문서
new_doc = Document()

# 페이지 설정
section = new_doc.sections[0]
section.page_height = Inches(11)
section.page_width = Inches(8.5)
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(1)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)

# ==================== 제목 ====================
title = new_doc.add_paragraph()
title_run = title.add_run('소량 산업 데이터 환경에서\n도메인 브리지 전이학습을 활용한 DPF 결함 검출')
title_run.font.size = Pt(16)
title_run.font.bold = True
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 저자
new_doc.add_paragraph()
author = new_doc.add_paragraph()
author_run = author.add_run('이규영\n')
author_run.font.size = Pt(12)
affil_run = author.add_run('국민대학교')
affil_run.font.size = Pt(11)
affil_run.font.italic = True
author.alignment = WD_ALIGN_PARAGRAPH.CENTER

new_doc.add_paragraph()
new_doc.add_paragraph()

# ==================== ABSTRACT ====================
abstract = new_doc.add_paragraph()
abstract.add_run('Abstract').bold = True
abstract.add_run('—')
abstract_text = """제조업에서 딥러닝 기반 결함 검출은 데이터 부족과 도메인 특화성으로 인해 실용화에 어려움을 겪어왔다. 본 연구는 디젤 미립자 필터(DPF) 결함 검출을 위한 도메인 브리지 전이학습 프레임워크를 제안한다. 제안 방법은 ImageNet → X-ray 결함 데이터 → DPF 데이터의 3단계 계층적 전이학습을 통해, 단 339장의 제한된 데이터로 91.7% mAP50 정확도를 달성했다. 이는 직접 학습(56.9%) 대비 34.8%p, ImageNet 직접 전이(72.3%) 대비 19.4%p의 성능 향상을 보여준다. 실험 결과, 중간 도메인(X-ray)을 경유하는 것이 도메인 갭을 효과적으로 완화하며, 100 에포크의 충분한 학습이 현대 어텐션 기반 모델의 성능을 최대화함을 확인했다. 본 프레임워크는 CPU 환경에서도 학습 가능하여 접근성이 높고, 완전한 재현성을 보장하며, 다양한 제조 분야로 확장 가능하다."""
abstract.add_run(abstract_text)
abstract.paragraph_format.line_spacing = 1.15

# Index Terms
new_doc.add_paragraph()
index_terms = new_doc.add_paragraph()
index_terms.add_run('Index Terms').italic = True
index_terms.add_run('—DPF 결함 검출, 전이학습, 도메인 적응, 제조업 AI, 소량 데이터 학습, YOLOv11, 객체 탐지, 품질 관리')

new_doc.add_paragraph()
new_doc.add_paragraph()

# ==================== 본문 처리 함수 ====================
def clean_text(text):
    """AI 톤 및 과장된 표현 제거"""
    replacements = {
        # 과장된 표현
        '성능 폭발': '성능 가속',
        '늦은 개화': '후반부 성능 향상',
        '시너지 발현': '상승 효과 확인',
        '치명적 결과를 초래': '심각한 영향을 미칠',
        '치명적 결과': '심각한 결과',
        
        # 이모지 및 특수 기호
        '⚠️': '',
        '★': '',
        '🔥': '',
        '⭐': '',
        '✅': '',
        
        # 과도한 강조
        '(2배 가속!)': '(약 2배)',
        '(!': '(',
        '!)': ')',
        
        # 비공식적 표현
        ' ← ': ' ',
        '← ': '',
        
        # ASCII 아트 제거 준비
        '├─ ': '• ',
        '└─ ': '• ',
        '│  ': '  ',
    }
    
    result = text
    for old, new in replacements.items():
        result = result.replace(old, new)
    
    # 다중 공백 정리
    result = re.sub(r'\s+', ' ', result)
    result = result.strip()
    
    return result

# 건너뛸 섹션
skip_sections_content = False
skip_until_section = None

# 본문 추가
para_count = 0
skipped_count = 0

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    if not text:
        continue
    
    # 초기 제목/저자/Abstract 건너뛰기 (첫 10개 문단)
    if i < 6:
        continue
    
    # "핵심 방법론", "주요 발견", "산업적 기여" 섹션 건너뛰기
    if text in ['핵심 방법론', '주요 발견', '산업적 기여', '핵심 키워드:', '핵심 키워드']:
        skip_sections_content = True
        skip_until_section = 'I. 서론'
        skipped_count += 1
        continue
    
    if skip_sections_content:
        if text == skip_until_section or text.startswith('I. 서론'):
            skip_sections_content = False
            skip_until_section = None
        else:
            skipped_count += 1
            continue
    
    # 텍스트 정리
    cleaned = clean_text(text)
    
    if not cleaned:
        continue
    
    # 섹션 헤딩 감지
    is_major_heading = bool(re.match(r'^[IVX]+\.\s+', text))
    is_minor_heading = bool(re.match(r'^[A-Z]\.\s+', text) and not re.match(r'^[IVX]+', text))
    is_numbered = bool(re.match(r'^\d+\.\s+', text) and len(text) < 200)
    
    # 문단 생성
    new_para = new_doc.add_paragraph()
    
    if is_major_heading:
        # I., II., III. 등 주요 섹션
        run = new_para.add_run(cleaned)
        run.font.size = Pt(12)
        run.font.bold = True
        new_para.paragraph_format.space_before = Pt(12)
        new_para.paragraph_format.space_after = Pt(6)
    elif is_minor_heading:
        # A., B., C. 등 부섹션
        run = new_para.add_run(cleaned)
        run.font.size = Pt(11)
        run.font.bold = True
        new_para.paragraph_format.space_before = Pt(6)
        new_para.paragraph_format.space_after = Pt(3)
    elif is_numbered:
        # 1., 2., 3. 등
        run = new_para.add_run(cleaned)
        run.font.bold = True
        new_para.paragraph_format.space_before = Pt(3)
    else:
        # 일반 본문
        new_para.add_run(cleaned)
        new_para.paragraph_format.line_spacing = 1.15
    
    para_count += 1
    
    if para_count % 200 == 0:
        print(f"처리: {para_count} 문단, 건너뜀: {skipped_count}")

print(f"\n총 {para_count} 문단 처리 완료 (건너뜀: {skipped_count})")

# ==================== 표 복사 ====================
print("\n표 복사 중...")
table_count = 0

for idx, table in enumerate(doc.tables):
    if idx >= 15:  # 처음 15개 표만
        break
    
    try:
        # 표 크기 확인
        if len(table.rows) == 0 or len(table.columns) == 0:
            continue
        
        # 새 표 생성
        new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        new_table.style = 'Light Grid Accent 1'
        
        # 셀 내용 복사
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                cell_text = clean_text(cell.text)
                new_table.rows[i].cells[j].text = cell_text
                
                # 첫 행은 굵게
                if i == 0:
                    for paragraph in new_table.rows[i].cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
        
        new_doc.add_paragraph()  # 표 뒤 공백
        table_count += 1
        print(f"  표 {table_count} 복사 완료 ({len(table.rows)}x{len(table.columns)})")
        
    except Exception as e:
        print(f"  표 {idx+1} 처리 중 오류: {e}")

print(f"\n총 {table_count}개 표 복사 완료")

# ==================== 이미지 처리 안내 ====================
new_doc.add_page_break()
image_note = new_doc.add_paragraph()
image_note.add_run('\n[이미지 삽입 안내]\n').bold = True
image_note.add_run("""
원본 문서의 이미지는 수동으로 삽입해주세요:
- 그림 1: Stage 2 전체 학습 곡선
- 그림 2-15: 기타 실험 결과 그림들

이미지 파일 위치: converted_md/images/ 폴더
""")

# ==================== 저장 ====================
output_path = r'C:\1.이규영개인폴더\09.##### SCHOOL #####\논문작성\IEEE_DPF_Paper_Revised_v3.docx'
new_doc.save(output_path)

print("="*80)
print(f"✅ 수정 완료!")
print(f"   파일: {output_path}")
print("="*80)
print(f"\n주요 수정 사항:")
print(f"  1. 제목 개선: '소량 산업 데이터 환경에서 도메인 브리지 전이학습을 활용한 DPF 결함 검출'")
print(f"  2. Abstract 간결화 및 재작성")
print(f"  3. 초반 요약 섹션 제거 (I. 서론 전까지)")
print(f"  4. AI 톤 제거:")
print(f"     - '성능 폭발' → '성능 가속'")
print(f"     - '늦은 개화' → '후반부 성능 향상'")
print(f"     - 이모지 전체 제거")
print(f"  5. ASCII 트리를 불릿 포인트로 변환")
print(f"  6. 과장된 표현 수정")
print(f"  7. 표 {table_count}개 복사 및 스타일 적용")
print(f"  8. 총 {para_count}개 문단 정리")
print(f"\n다음 단계:")
print(f"  - 이미지 수동 삽입")
print(f"  - 참고문헌 확인")
print(f"  - 최종 검토")
