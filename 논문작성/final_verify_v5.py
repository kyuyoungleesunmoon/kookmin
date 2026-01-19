# -*- coding: utf-8 -*-
"""최종 문서 검증"""

from docx import Document
import re

doc = Document(r'c:\1.이규영개인폴더\09.##### SCHOOL #####\논문작성\IEEE_DPF_Paper_Final_v5.docx')

print('=' * 65)
print('IEEE_DPF_Paper_Final_v5.docx 최종 검증')
print('=' * 65)

full_text = ' '.join([p.text for p in doc.paragraphs])

# 1. 코드 패턴 확인
code_count = 0
for para in doc.paragraphs:
    text = para.text.strip()
    if len(text) < 100 and re.match(r'^[a-z_]+\s*[=:]\s*[\d\.\{\[\'\"]', text):
        code_count += 1

print(f'\n1. 코드 패턴: {code_count}개')

# 2. 서술형 변환 확인
prose_checks = [
    ('Stage 1 학습에서는', 'Stage 1 설정'),
    ('Stage 2 학습에서는', 'Stage 2 설정'),
    ('Stage 1 하이퍼파라미터 설정', 'Stage 1 하이퍼파라미터'),
    ('Stage 2 하이퍼파라미터 설정', 'Stage 2 하이퍼파라미터'),
    ('데이터 증강 전략', 'Augmentation 설명'),
    ('Mosaic 증강', 'Mosaic 설명'),
    ('cosine decay', '학습률 스케줄'),
    ('ImageNet에서 사전 학습', '사전학습 설명'),
]

print('\n2. 서술형 변환 확인:')
for term, desc in prose_checks:
    found = term in full_text
    status = '✓' if found else '✗'
    print(f'   {status} {desc}')

# 3. 핵심 내용 확인
metrics = [
    ('91.7%', 'mAP50 최종 성능'),
    ('34.8%p', '총 성능 향상'),
    ('19.4%p', '도메인 브리지 효과'),
    ('14.8%p', '후반부 학습'),
    ('늦은 개화', 'Late Blooming'),
    ('도메인 브리지', '핵심 방법론'),
    ('YOLO11', '모델'),
    ('339장', 'DPF 데이터'),
]

print('\n3. 핵심 내용 확인:')
all_ok = True
for term, desc in metrics:
    found = term in full_text
    status = '✓' if found else '✗'
    if not found:
        all_ok = False
    print(f'   {status} {term} ({desc})')

# 4. 통계
print('\n4. 문서 통계:')
print(f'   - 총 단락: {len(doc.paragraphs)}')
print(f'   - 내용 있는 단락: {sum(1 for p in doc.paragraphs if p.text.strip())}')
print(f'   - 테이블: {len(doc.tables)}개')

# 5. 수식 확인
eq_count = sum(1 for p in doc.paragraphs if '\\tag' in p.text)
print(f'   - 번호 수식: {eq_count}개')

print('\n' + '=' * 65)
if code_count == 0 and all_ok:
    print('✅ 검증 완료! 최종 파일: IEEE_DPF_Paper_Final_v5.docx')
else:
    print('⚠️ 일부 항목 확인 필요')
print('=' * 65)
