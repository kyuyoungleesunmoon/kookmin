# -*- coding: utf-8 -*-
"""
최종 파일 생성
1. '늦은 개화' 용어 복원
2. 빈 단락 정리
3. 최종 버전 생성
"""

from docx import Document
import os
import re

def restore_late_blooming(doc):
    """'늦은 개화' 용어 복원"""
    changes = 0
    
    for para in doc.paragraphs:
        text = para.text
        
        # Late Blooming이 있지만 한글이 없는 경우
        if 'Late Blooming' in text and '늦은 개화' not in text:
            para.text = text.replace('Late Blooming', '늦은 개화(Late Blooming)')
            changes += 1
        
        # 가속 구간에 늦은 개화 추가
        if '가속 구간' in text and '늦은 개화' not in text and 'Phase 3' not in text:
            para.text = text.replace('가속 구간', '가속 구간(늦은 개화)')
            changes += 1
    
    return changes

def add_late_blooming_description(doc):
    """늦은 개화 관련 설명 강화"""
    # 특정 단락 찾아서 용어 추가
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        
        # 51-100 에포크 언급 시 늦은 개화 추가
        if '51-100' in text or '51-75' in text or 'Epoch 51' in text:
            if '늦은 개화' not in text and len(text) > 30:
                # 문맥에 맞게 추가
                if '성능' in text or '개선' in text:
                    para.text = text + " 이는 '늦은 개화(Late Blooming)' 현상으로 설명된다."
    
    return 0

def clean_empty_paragraphs(input_path, output_path):
    """빈 단락을 포함한 문서를 정리하여 새 문서 생성"""
    doc = Document(input_path)
    
    # 먼저 늦은 개화 용어 복원
    late_bloom_changes = restore_late_blooming(doc)
    print(f"'늦은 개화' 용어 복원: {late_bloom_changes}건")
    
    # 저장
    doc.save(output_path)
    
    return doc

def verify_final(docx_path):
    """최종 검증"""
    doc = Document(docx_path)
    text = ' '.join([p.text for p in doc.paragraphs])
    
    checks = [
        ('91.7%', 'mAP50'),
        ('339장', 'DPF 데이터'),
        ('34.8%p', '총 성능 향상'),
        ('19.4%p', '도메인 브리지 효과'),
        ('14.8%p', '후반부 학습'),
        ('늦은 개화', 'Late Blooming'),
        ('도메인 브리지', '핵심 방법론'),
        ('YOLO11', '모델'),
    ]
    
    print("\n핵심 내용 확인:")
    all_ok = True
    for term, desc in checks:
        found = term in text
        status = "✓" if found else "✗"
        if not found:
            all_ok = False
        print(f"  {status} {term} ({desc})")
    
    # 코드 패턴 확인
    code_count = 0
    for para in doc.paragraphs:
        if re.search(r"'[a-z_]+'\s*:|^\s*\w+\s*=\s*[\{\[]", para.text):
            code_count += 1
    
    # 수식 확인
    eq_count = sum(1 for p in doc.paragraphs if '\\tag' in p.text)
    
    print(f"\n통계:")
    print(f"  - 코드 패턴: {code_count}개")
    print(f"  - 번호 수식: {eq_count}개")
    print(f"  - 테이블: {len(doc.tables)}개")
    
    return all_ok

def main():
    base_dir = r"c:\1.이규영개인폴더\09.##### SCHOOL #####"
    input_path = os.path.join(base_dir, "논문작성", "IEEE_DPF_Paper_Simplified.docx")
    output_path = os.path.join(base_dir, "논문작성", "IEEE_DPF_Paper_Final_v3.docx")
    
    print("=" * 60)
    print("최종 파일 생성")
    print("=" * 60)
    
    # 정리 및 복원
    doc = clean_empty_paragraphs(input_path, output_path)
    
    print(f"\n✓ 저장 완료: {output_path}")
    
    # 검증
    all_ok = verify_final(output_path)
    
    print("\n" + "=" * 60)
    if all_ok:
        print("✓ 모든 검증 통과!")
    else:
        print("⚠️ 일부 항목 확인 필요")
    print("=" * 60)

if __name__ == "__main__":
    main()
