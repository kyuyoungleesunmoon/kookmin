# -*- coding: utf-8 -*-
"""
최종 검증 및 추가 정리
"""

from docx import Document
import os
import re

def detailed_check(docx_path):
    """상세 검증"""
    doc = Document(docx_path)
    
    print("=" * 70)
    print("상세 검증")
    print("=" * 70)
    
    # 코드 패턴 상세 검색
    code_patterns = [
        (r"'[a-z_]+'\s*:", "dict key 패턴"),
        (r"^\s*\w+\s*=\s*[\{\[]", "변수 할당 패턴"),
        (r"^\s*[\[\{}\]],?\s*$", "괄호만 있는 줄"),
        (r"#\s*[가-힣a-zA-Z].*\d", "주석 패턴"),
        (r":\s*\d+\.?\d*,?\s*#", "값:숫자,# 패턴"),
    ]
    
    issues = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        for pattern, desc in code_patterns:
            if re.search(pattern, text):
                # 서술형 변환된 것은 제외
                if len(text) > 100:  # 긴 문장은 서술형
                    continue
                if text.startswith('Stage') or text.startswith('학습'):
                    continue
                    
                issues.append((i, desc, text[:60]))
    
    if issues:
        print(f"\n⚠️ 잠재적 코드 패턴 발견: {len(issues)}건")
        for idx, desc, text in issues[:20]:
            print(f"  [{idx}] {desc}: {text}")
    else:
        print("\n✓ 코드 패턴 없음")
    
    # 서술형 변환 샘플 출력
    print("\n" + "-" * 70)
    print("서술형 변환 샘플:")
    print("-" * 70)
    
    keywords = ['학습 에포크는', 'Stage 1 학습에서', 'Stage 2 학습에서', '데이터 증강']
    shown = 0
    
    for para in doc.paragraphs:
        text = para.text.strip()
        for kw in keywords:
            if kw in text and len(text) > 50:
                print(f"\n{text[:200]}...")
                shown += 1
                break
        if shown >= 4:
            break
    
    return len(issues)

def final_cleanup(input_path, output_path):
    """최종 정리"""
    doc = Document(input_path)
    
    cleanup_count = 0
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # 짧은 코드 스타일 라인 제거
        if len(text) < 50:
            # dict key 패턴
            if re.search(r"'[a-z_]+'\s*:\s*[\d\.\[\]'\"]+,?$", text):
                para.text = ""
                cleanup_count += 1
            # 괄호만
            elif re.match(r"^[\[\{}\]],?\s*$", text):
                para.text = ""
                cleanup_count += 1
            # 변수 = { 또는 }
            elif re.match(r"^\w+\s*=\s*[\{\[]$", text):
                para.text = ""
                cleanup_count += 1
            elif text == '}' or text == '},':
                para.text = ""
                cleanup_count += 1
    
    if cleanup_count > 0:
        print(f"추가 정리: {cleanup_count}건")
        doc.save(output_path)
    else:
        print("추가 정리 불필요")
        # 그대로 복사
        doc.save(output_path)
    
    return cleanup_count

def main():
    base_dir = r"c:\1.이규영개인폴더\09.##### SCHOOL #####"
    v4_path = os.path.join(base_dir, "논문작성", "IEEE_DPF_Paper_Final_v4.docx")
    final_path = os.path.join(base_dir, "논문작성", "IEEE_DPF_Paper_Final_v4_clean.docx")
    
    # 상세 검증
    issues = detailed_check(v4_path)
    
    if issues > 0:
        # 추가 정리
        final_cleanup(v4_path, final_path)
        detailed_check(final_path)
    else:
        print("\n✅ 최종 파일: IEEE_DPF_Paper_Final_v4.docx")

if __name__ == "__main__":
    main()
