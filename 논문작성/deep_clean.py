from docx import Document
import re

input_path = r"C:\국민대프로젝트\논문작성\IEEE_DPF_Paper_Clean_v6.docx"
output_path = r"C:\국민대프로젝트\논문작성\IEEE_DPF_Paper_Clean_v7.docx"

doc = Document(input_path)

# More aggressive code patterns
code_patterns = [
    r'^\s*(def |class |import |from )',
    r'^\s*#.*코드|# \d+\.',  # Korean code comments
    r'^\s*\w+\s*=\s*[\{\[]',  # Any assignment to dict/list
    r"^\s*'[\w_]+'\s*:",  # Dict keys
    r'^\s*(plt|np|torch|model|df|cv2)\.',  # Library calls
    r'^\s*├─|└─|│',  # Tree structure
    r'^\s*(return|print|if|for|while|try|except|with)\s',  # Control flow
    r'\.train\(|\.fit\(|\.predict\(',  # ML method calls
    r'^\s*\w+_config\s*=',  # Config variables
    r'^\s*(augmentation|logs|dataloader)\s*=',  # ML configs
    r'^\s*SEED\s*=',  # Seed setting
    r'^\s*os\.environ',  # Environment vars
    r'^\s*random\.|np\.random',  # Random calls
    r'^\s*\d+\s*$',  # Standalone numbers (likely code indices)
    r'epochs:\s*\d+',  # Hyperparameter specs
    r'batch:\s*\d+',
    r'imgsz:\s*\d+',
    r'lr0:\s*[\d\.]+',
    r'momentum:\s*[\d\.]+',
    r'weight_decay:\s*[\d\.]+',
    r'patience:\s*\d+',
    r'close_mosaic:\s*\d+',
]

# Keep only these essential equation types
# Everything else gets removed
essential_keywords = [
    'mAP@',  # mAP definition
    'd_{\\mathcal{A}}',  # Proxy A-distance
    '\\mathcal{L}_{YOLO}',  # YOLO loss (just the main one)
    'Precision =',  # Precision formula
    'Recall =',  # Recall formula  
    'F1 =',  # F1 formula
]

def is_code(text):
    for pat in code_patterns:
        if re.search(pat, text, re.IGNORECASE):
            return True
    return False

def is_essential_eq(text):
    for kw in essential_keywords:
        if kw in text:
            return True
    return False

def is_equation(text):
    eq_markers = ['$$', '\\tag{', '\\begin{', '\\frac{', '\\sum', '\\mathcal{', '\\left', '\\eta', '\\lambda', '\\epsilon']
    for m in eq_markers:
        if m in text:
            return True
    return False

# Process
to_remove = []
in_eq_block = False
eq_block_indices = []
eq_block_essential = False

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue
    
    # Handle $$ blocks
    if text == '$$' or text.startswith('$$'):
        if not in_eq_block:
            in_eq_block = True
            eq_block_indices = [i]
            eq_block_essential = is_essential_eq(text)
        else:
            eq_block_indices.append(i)
            if not eq_block_essential:
                to_remove.extend(eq_block_indices)
            in_eq_block = False
            eq_block_indices = []
            eq_block_essential = False
        continue
    
    if in_eq_block:
        eq_block_indices.append(i)
        if is_essential_eq(text):
            eq_block_essential = True
        continue
    
    # Remove inline equations that aren't essential
    if is_equation(text) and not is_essential_eq(text):
        to_remove.append(i)
        continue
    
    # Remove code
    if is_code(text):
        to_remove.append(i)
        continue

print(f"Removing {len(to_remove)} more paragraphs...")

# Remove in reverse
for idx in sorted(set(to_remove), reverse=True):
    try:
        p = doc.paragraphs[idx]._element
        p.getparent().remove(p)
    except:
        pass

doc.save(output_path)

# Verify
doc2 = Document(output_path)
print(f"Final paragraph count: {len(doc2.paragraphs)}")
print(f"Saved to: {output_path}")
