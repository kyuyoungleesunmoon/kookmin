# -*- coding: utf-8 -*-
"""
IEEE_DPF_Paper_Final_v2.docx 수식 및 소스코드 분석
과도한 수식 제거 및 소스코드 서술형 변환
"""

from docx import Document
import os
import re

def analyze_equations_and_code(docx_path):
    """수식과 소스코드 분석"""
    doc = Document(docx_path)
    
    equations = []
    code_blocks = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        if not text:
            continue
        
        # 수식 패턴 찾기
        if re.search(r'\\tag\{(\d+)\}', text):
            equations.append({
                'index': i,
                'number': re.findall(r'\\tag\{(\d+)\}', text),
                'text': text[:200]
            })
        
        # LaTeX 수식 패턴
        if re.search(r'\$\$|\$[^$]+\$|\\frac|\\sum|\\mathcal', text):
            if not any(e['index'] == i for e in equations):
                equations.append({
                    'index': i,
                    'number': [],
                    'text': text[:200]
                })
        
        # 소스코드 패턴 찾기
        code_patterns = [
            r'^\s*(def |class |import |from |if |for |while |return |print\()',
            r'^\s*(model\s*=|results\s*=|config\s*=)',
            r'^\s*#\s*[A-Za-z]',  # 주석
            r'^\s*\w+\s*=\s*\{',  # 딕셔너리 정의
            r'^\s*\w+\s*=\s*\[',  # 리스트 정의
            r"'[a-z_]+'\s*:",  # 딕셔너리 키
            r'^\s*[a-z_]+\s*:\s*[0-9]+',  # YAML 스타일
            r'torch\.|numpy\.|np\.|pd\.',  # 라이브러리 호출
            r'\.train\(|\.eval\(|\.fit\(',  # 메서드 호출
        ]
        
        for pattern in code_patterns:
            if re.search(pattern, text):
                code_blocks.append({
                    'index': i,
                    'text': text[:300],
                    'pattern': pattern
                })
                break
    
    return equations, code_blocks

def main():
    base_dir = r"c:\1.이규영개인폴더\09.##### SCHOOL #####"
    docx_path = os.path.join(base_dir, "논문작성", "IEEE_DPF_Paper_Final_v2.docx")
    
    doc = Document(docx_path)
    
    print("=" * 70)
    print("수식 및 소스코드 분석")
    print("=" * 70)
    
    equations, code_blocks = analyze_equations_and_code(docx_path)
    
    # 수식 분석
    print(f"\n[1] 번호 있는 수식: {len([e for e in equations if e['number']])}개")
    numbered_eqs = [e for e in equations if e['number']]
    for eq in numbered_eqs:
        print(f"   ({eq['number'][0]}) 단락 {eq['index']}: {eq['text'][:80]}...")
    
    print(f"\n[2] 기타 수식 표현: {len([e for e in equations if not e['number']])}개")
    
    # 소스코드 분석
    print(f"\n[3] 소스코드 블록: {len(code_blocks)}개")
    for i, code in enumerate(code_blocks[:30]):
        print(f"   {i+1}. 단락 {code['index']}: {code['text'][:100]}...")
    
    if len(code_blocks) > 30:
        print(f"   ... 외 {len(code_blocks) - 30}개 더")
    
    # 전체 통계
    print(f"\n" + "=" * 70)
    print("분석 요약")
    print("=" * 70)
    print(f"총 단락 수: {len(doc.paragraphs)}")
    print(f"번호 수식: {len(numbered_eqs)}개")
    print(f"소스코드 블록: {len(code_blocks)}개")

if __name__ == "__main__":
    main()
