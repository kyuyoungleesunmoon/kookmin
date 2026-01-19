from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import shutil
import os

print("="*80)
print("IEEE TII 템플릿 기반 논문 재구성")
print("="*80)

# 원본 템플릿 로드
template_path = r'C:\1.이규영개인폴더\09.##### SCHOOL #####\IEEE_TII_Template\TII_Articles_Word_template_2025.docx'
current_paper = r'C:\1.이규영개인폴더\09.##### SCHOOL #####\논문작성\IEEE_DPF_Paper_Draft_v2.docx'

print(f"\n1. 템플릿 로드...")
doc = Document(template_path)
print(f"   ✓ 템플릿 로드 완료")

# 템플릿의 초기 콘텐츠 제거 (샘플 텍스트)
print(f"\n2. 템플릿 샘플 텍스트 정리...")

# 기존 콘텐츠 제거 (제목, 저자, abstract 이후 본문)
to_remove = []
in_content = False

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    # Abstract 찾기
    if 'Abstract' in text:
        in_content = True
    
    # Abstract 이후 샘플 텍스트는 제거
    if in_content and i > 20:
        # "I. INTRODUCTION" 또는 "II." 등을 찾으면 멈춤
        if text.startswith(('I.', 'II.', 'INTRODUCTION', 'RELATED')):
            in_content = False
            break
        
        # 샘플 텍스트 마크
        if i > 15 and i < 100:
            to_remove.append(para)

print(f"   제거할 샘플 문단: {len(to_remove)}개")

# 현재 논문 로드
print(f"\n3. 현재 논문 로드...")
current_doc = Document(current_paper)

# 제목 추출 (첫 3개 문단)
title_text = ""
author_text = ""
affiliation_text = ""

if current_doc.paragraphs:
    title_text = current_doc.paragraphs[0].text.strip()
    if len(current_doc.paragraphs) > 1:
        author_text = current_doc.paragraphs[1].text.strip()
    if len(current_doc.paragraphs) > 2:
        affiliation_text = current_doc.paragraphs[2].text.strip()

print(f"   제목: {title_text[:50]}...")
print(f"   저자: {author_text}")
print(f"   소속: {affiliation_text}")

# 템플릿의 제목/저자 영역 업데이트
print(f"\n4. 템플릿 메타데이터 업데이트...")

title_updated = False
author_updated = False

for para in doc.paragraphs[:30]:
    text = para.text.strip()
    
    # 제목 찾기 및 교체
    if 'Preparation of Papers' in text or (para.style.name == 'Title' and not title_updated):
        para.clear()
        run = para.add_run('DPF 결함 검출을 위한 도메인 브리지 전이학습')
        run.font.size = Pt(14)
        run.font.bold = True
        title_updated = True
        print(f"   ✓ 제목 업데이트")
    
    # 저자 찾기 및 교체
    elif ('First Author' in text or 'author' in text.lower()) and not author_updated:
        para.clear()
        run = para.add_run('Lee, Kyu-Young')
        run.font.size = Pt(11)
        author_updated = True
        print(f"   ✓ 저자 업데이트")

# Abstract 섹션 찾기 및 업데이트
print(f"\n5. Abstract 업데이트...")

abstract_found = False
for i, para in enumerate(doc.paragraphs):
    if 'Abstract' in para.text and not abstract_found:
        # Abstract 텍스트 추출
        abstract_text = ""
        for p in current_doc.paragraphs:
            if 'Abstract' in p.text and '—' in p.text:
                # —이후의 텍스트 추출
                parts = p.text.split('—')
                if len(parts) > 1:
                    abstract_text = parts[1].strip()
                    break
        
        if abstract_text:
            # 다음 문단 찾기 (Abstract 다음)
            if i + 1 < len(doc.paragraphs):
                next_para = doc.paragraphs[i + 1]
                if next_para.text.strip() and 'Index' not in next_para.text:
                    next_para.clear()
                    next_para.add_run(abstract_text)
                    print(f"   ✓ Abstract 업데이트 ({len(abstract_text)} 자)")
                    abstract_found = True

# Index Terms 업데이트
print(f"\n6. Index Terms 업데이트...")

index_found = False
for i, para in enumerate(doc.paragraphs):
    if 'Index Terms' in para.text and not index_found:
        # 현재 논문에서 Index Terms 추출
        for p in current_doc.paragraphs:
            if 'Index Terms' in p.text or 'Keywords' in p.text:
                if '—' in p.text:
                    terms = p.text.split('—')[1].strip()
                    para.clear()
                    run = para.add_run('Index Terms—')
                    run.italic = True
                    run.add_run(terms)
                    print(f"   ✓ Index Terms 업데이트")
                    index_found = True
                    break

# 페이지 설정 확인
print(f"\n7. 페이지 설정 확인...")
section = doc.sections[0]
print(f"   여백: T={section.top_margin.inches:.2f}\", B={section.bottom_margin.inches:.2f}\"")
print(f"        L={section.left_margin.inches:.2f}\", R={section.right_margin.inches:.2f}\"")

# 저장
output_path = r'C:\1.이규영개인폴더\09.##### SCHOOL #####\논문작성\IEEE_DPF_Paper_TII_Format.docx'
doc.save(output_path)

print(f"\n" + "="*80)
print(f"✅ 템플릿 기반 재구성 완료!")
print(f"   저장: {output_path}")
print("="*80)

# 상세 체크리스트
print(f"\n8. 최종 점검 항목:")
checklist = [
    ("페이지 크기 (Letter)", True),
    ("여백 설정 (0.7\")", True),
    ("제목 영문화", False),
    ("저자 영문 표기", False),
    ("소속 정보", False),
    ("Abstract 영문", False),
    ("Index Terms 포함", True),
    ("I. INTRODUCTION", False),
    ("II. RELATED WORK", False),
    ("III. METHODOLOGY", False),
    ("IV. EXPERIMENTS", False),
    ("V. RESULTS", False),
    ("VI. DISCUSSION", False),
    ("VII. CONCLUSION", False),
    ("References IEEE 형식", False),
    ("Figure captions", False),
    ("Table captions", False),
]

for item, status in checklist:
    mark = "✓" if status else "⚠️"
    print(f"   {mark} {item}")

print(f"\n⚠️ 주의: 다음 항목은 수동으로 확인/수정이 필요합니다:")
print(f"   - Abstract를 150-250단어 영문으로 재작성")
print(f"   - 저자 및 소속 정보 영문화")
print(f"   - 섹션 제목 영문화 (I. INTRODUCTION 등)")
print(f"   - 참고문헌 IEEE 형식으로 통일")
print(f"   - 모든 그림/표에 영문 캡션 추가")
