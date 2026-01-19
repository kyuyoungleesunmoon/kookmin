from docx import Document
from docx.shared import Pt
import re
import os

# Paths
md_path = r"C:\국민대프로젝트\converted_md\이규영_국민대_DPF_논문.md"
template_path = r"C:\국민대프로젝트\IEEE_TII_Template\TII_Articles_Word_template_2025.docx"
output_path = r"C:\국민대프로젝트\논문작성\IEEE_DPF_Paper_Submission_Ready.docx"

print(f"Reading Markdown: {md_path}")
with open(md_path, 'r', encoding='utf-8') as f:
    lines = f.readlines()

print(f"Opening Template: {template_path}")
try:
    doc = Document(template_path)
    print(f"Initial Template Paragraphs: {len(doc.paragraphs)}")
except Exception as e:
    print(f"Template load failed: {e}. creating new.")
    doc = Document()

# Robust Clearing of Body Content
print("Clearing template body content (keeping styles)...")
# We delete in reverse order or use a while loop to avoid index shifting issues
if len(doc.paragraphs) > 0:
    # Method: Remove all paragraphs from the document body
    # Accessing the body element directly is safer
    body = doc._body
    body.clear_content() 
    # Note: clear_content() removes all child elements (paragraphs, tables, etc.)
    # If python-docx version doesn't support clear_content on body, we fall back
    if len(doc.paragraphs) > 0:
        print("Fallback clearing...")
        for p in doc.paragraphs[:]: # Slice copy to iterate safely?
             p._element.getparent().remove(p._element)
             
    # Double check
    print(f"Post-Clearing Paragraphs: {len(doc.paragraphs)}")

# Clean logic to avoid duplicates in MD reading (just in case)
# ... lines are lines.

# Helper to add styled paragraph
def add_para(text, style=None):
    try:
        if style:
            p = doc.add_paragraph(text, style=style)
        else:
            p = doc.add_paragraph(text)
    except:
        p = doc.add_paragraph(text) # Fallback if style missing
    return p

print("Processing content...")
for line in lines:
    line = line.strip()
    if not line:
        continue
        
    # Headers
    if line.startswith('# '):
        text = line.replace('# ', '').replace('*', '')
        add_para(text, style='Title')
    elif line.startswith('## '):
        text = line.replace('## ', '').replace('*', '')
        if "초록" in text or "Abstract" in text:
            add_para("Abstract", style='Heading 1')
            add_para(text, style='Body Text')
        else:
            add_para(text, style='Heading 1')     
    elif line.startswith('### '):
        text = line.replace('### ', '').replace('*', '')
        add_para(text, style='Heading 1')
    elif line.startswith('#### '):
        text = line.replace('#### ', '').replace('*', '')
        add_para(text, style='Heading 2')
    elif line.startswith('##### '):
        text = line.replace('##### ', '').replace('*', '')
        add_para(text, style='Heading 3')
        
    # Images
    elif line.startswith('!['):
        match = re.search(r'!\[(.*?)\]\((.*?)\)', line)
        if match:
            caption = match.group(1)
            img_path = match.group(2)
            if not os.path.isabs(img_path):
                 base_dir = os.path.dirname(md_path)
                 img_full_path = os.path.join(base_dir, img_path)
            else:
                img_full_path = img_path
                
            if os.path.exists(img_full_path):
                try:
                    doc.add_picture(img_full_path, width=Pt(400))
                    # Check if last paragraph exists to prevent error
                    if doc.paragraphs:
                        doc.paragraphs[-1].alignment = 1 # Center
                    add_para(f"Fig. {caption}", style='Caption')
                except Exception as e:
                    add_para(f"[Image: {caption} - {img_path}]")
            else:
                 add_para(f"[Image Missing: {caption} - {img_path}]")
                 
    # Tables
    elif line.startswith('|'):
        add_para(line, style='Normal')
        
    # Normal Text or Lists
    else:
        # Check if Metadata
        if line.startswith('저자:') or line.startswith('소속:'):
            add_para(line, style='Author')
        elif line.startswith('- ') or line.startswith('* '):
             add_para(line[2:], style='List Bullet')
        elif re.match(r'\d+\.', line):
             add_para(line, style='List Number')
        else:
            add_para(line, style='Body Text')

print(f"Saving to: {output_path}")
doc.save(output_path)
print("Done.")
