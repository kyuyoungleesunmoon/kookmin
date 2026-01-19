# -*- coding: utf-8 -*-
"""
DOC/DOCX/PDF 파일을 Markdown으로 변환하고 이미지/표를 추출하는 스크립트
"""

import os
import fitz  # PyMuPDF
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from PIL import Image
import io
import re
from pathlib import Path

# 작업 디렉토리 설정
BASE_DIR = r"c:\1.이규영개인폴더\09.##### SCHOOL #####"
OUTPUT_DIR = os.path.join(BASE_DIR, "converted_md")
IMAGES_DIR = os.path.join(OUTPUT_DIR, "images")

# 디렉토리 생성
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(IMAGES_DIR, exist_ok=True)

def sanitize_filename(filename):
    """파일명에서 특수문자 제거"""
    return re.sub(r'[<>:"/\\|?*]', '_', filename)

def extract_table_as_markdown(table):
    """docx 테이블을 마크다운 테이블로 변환"""
    md_table = []
    
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        md_table.append("| " + " | ".join(cells) + " |")
        
        # 헤더 구분선 추가
        if i == 0:
            md_table.append("|" + "|".join(["---" for _ in cells]) + "|")
    
    return "\n".join(md_table)

def convert_docx_to_md(docx_path, output_name):
    """DOCX 파일을 Markdown으로 변환"""
    print(f"\n📄 Converting DOCX: {docx_path}")
    
    doc = Document(docx_path)
    md_content = []
    image_count = 0
    
    # 문서 제목
    base_name = os.path.splitext(output_name)[0]
    md_content.append(f"# {base_name}\n")
    md_content.append(f"*원본 파일: {os.path.basename(docx_path)}*\n")
    md_content.append("---\n")
    
    # 이미지 추출 (embedded images)
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            try:
                image_data = rel.target_part.blob
                image_ext = rel.target_ref.split('.')[-1]
                image_filename = f"{sanitize_filename(base_name)}_img_{image_count}.{image_ext}"
                image_path = os.path.join(IMAGES_DIR, image_filename)
                
                with open(image_path, 'wb') as img_file:
                    img_file.write(image_data)
                
                print(f"  ✅ Extracted image: {image_filename}")
                image_count += 1
            except Exception as e:
                print(f"  ⚠️ Image extraction error: {e}")
    
    # 문서 내용 처리
    for element in doc.element.body:
        # 테이블 처리
        if element.tag.endswith('tbl'):
            for table in doc.tables:
                if table._tbl == element:
                    md_content.append("\n" + extract_table_as_markdown(table) + "\n")
                    break
        # 단락 처리
        elif element.tag.endswith('p'):
            for para in doc.paragraphs:
                if para._p == element:
                    text = para.text.strip()
                    if text:
                        # 스타일 기반 헤딩 처리
                        style = para.style.name if para.style else ""
                        if "Heading 1" in style or "heading 1" in style.lower():
                            md_content.append(f"\n## {text}\n")
                        elif "Heading 2" in style or "heading 2" in style.lower():
                            md_content.append(f"\n### {text}\n")
                        elif "Heading 3" in style or "heading 3" in style.lower():
                            md_content.append(f"\n#### {text}\n")
                        elif "Title" in style:
                            md_content.append(f"\n# {text}\n")
                        else:
                            md_content.append(f"{text}\n")
                    break
    
    # 추출된 이미지 참조 추가
    if image_count > 0:
        md_content.append("\n---\n## 추출된 이미지\n")
        for i in range(image_count):
            image_ext = "png"  # 기본값
            for ext in ['png', 'jpg', 'jpeg', 'gif', 'emf', 'wmf']:
                if os.path.exists(os.path.join(IMAGES_DIR, f"{sanitize_filename(base_name)}_img_{i}.{ext}")):
                    image_ext = ext
                    break
            md_content.append(f"\n![Image {i+1}](images/{sanitize_filename(base_name)}_img_{i}.{image_ext})\n")
    
    # MD 파일 저장
    output_path = os.path.join(OUTPUT_DIR, f"{sanitize_filename(base_name)}.md")
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("\n".join(md_content))
    
    print(f"  ✅ Saved: {output_path}")
    print(f"  📊 Images extracted: {image_count}")
    return output_path

def convert_pdf_to_md(pdf_path, output_name):
    """PDF 파일을 Markdown으로 변환하고 이미지/표를 추출"""
    print(f"\n📄 Converting PDF: {pdf_path}")
    
    doc = fitz.open(pdf_path)
    md_content = []
    image_count = 0
    
    base_name = os.path.splitext(output_name)[0]
    md_content.append(f"# {base_name}\n")
    md_content.append(f"*원본 파일: {os.path.basename(pdf_path)}*\n")
    md_content.append("---\n")
    
    for page_num, page in enumerate(doc):
        md_content.append(f"\n## 페이지 {page_num + 1}\n")
        
        # 텍스트 추출
        text = page.get_text("text")
        if text.strip():
            # 기본 텍스트 정리
            lines = text.split('\n')
            cleaned_lines = []
            for line in lines:
                line = line.strip()
                if line:
                    cleaned_lines.append(line)
            md_content.append("\n".join(cleaned_lines))
        
        # 이미지 추출
        image_list = page.get_images()
        for img_index, img in enumerate(image_list):
            try:
                xref = img[0]
                pix = fitz.Pixmap(doc, xref)
                
                if pix.n - pix.alpha > 3:  # CMYK
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                
                image_filename = f"{sanitize_filename(base_name)}_page{page_num+1}_img{img_index}.png"
                image_path = os.path.join(IMAGES_DIR, image_filename)
                pix.save(image_path)
                
                md_content.append(f"\n![Page {page_num+1} Image {img_index+1}](images/{image_filename})\n")
                print(f"  ✅ Extracted: {image_filename}")
                image_count += 1
                
                pix = None
            except Exception as e:
                print(f"  ⚠️ Image extraction error on page {page_num+1}: {e}")
        
        # 페이지를 이미지로 렌더링 (표/그래프 캡처용)
        try:
            # 고해상도로 페이지 렌더링
            mat = fitz.Matrix(2, 2)  # 2x zoom
            pix = page.get_pixmap(matrix=mat)
            page_image_filename = f"{sanitize_filename(base_name)}_page{page_num+1}_full.png"
            page_image_path = os.path.join(IMAGES_DIR, page_image_filename)
            pix.save(page_image_path)
            
            md_content.append(f"\n### 페이지 {page_num+1} 전체 이미지\n")
            md_content.append(f"![Page {page_num+1} Full](images/{page_image_filename})\n")
            print(f"  ✅ Page rendered: {page_image_filename}")
        except Exception as e:
            print(f"  ⚠️ Page render error: {e}")
    
    doc.close()
    
    # MD 파일 저장
    output_path = os.path.join(OUTPUT_DIR, f"{sanitize_filename(base_name)}.md")
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("\n".join(md_content))
    
    print(f"  ✅ Saved: {output_path}")
    print(f"  📊 Images extracted: {image_count}")
    return output_path

def main():
    """메인 실행 함수"""
    print("=" * 60)
    print("📁 DOC/PDF to Markdown Converter")
    print("=" * 60)
    
    # 변환할 파일 목록
    files_to_convert = [
        # 논문작성 폴더
        (os.path.join(BASE_DIR, "논문작성", "IEEE_DPF_Paper_Korean_Final_Rev5.docx"), "IEEE_DPF_Paper_Korean_Final_Rev5"),
        (os.path.join(BASE_DIR, "논문작성", "이규영 국민대 DPF 논문.docx"), "이규영_국민대_DPF_논문"),
        
        # IEEE 템플릿
        (os.path.join(BASE_DIR, "IEEE_TII_Template", "TII_Articles_Word_template_2025.doc"), "TII_Articles_Word_template_2025"),
        
        # Processing 폴더 PDF
        (os.path.join(BASE_DIR, "processing", "DPF_Complete_Technical_Report.pdf"), "DPF_Complete_Technical_Report"),
        (os.path.join(BASE_DIR, "processing", "DPF_Complete_Technical_Report_For_Paper_20250814.pdf"), "DPF_Complete_Technical_Report_For_Paper"),
        
        # 참고논문 폴더
        (os.path.join(BASE_DIR, "참고논문", "2015-GDXray-Paper.pdf"), "2015_GDXray_Paper"),
        (os.path.join(BASE_DIR, "참고논문", "2408.11250v2.pdf"), "2408_11250v2"),
        (os.path.join(BASE_DIR, "참고논문", "LoHi-WELD_A_Novel_Industrial_Dataset_for_Weld_Defect_Detection_and_Classification_a_Deep_Learning_Study_and_Future_Perspectives.pdf"), "LoHi_WELD_Dataset_Paper"),
        (os.path.join(BASE_DIR, "참고논문", "nihms-1520836.pdf"), "nihms_1520836"),
    ]
    
    converted_files = []
    
    for file_path, output_name in files_to_convert:
        if not os.path.exists(file_path):
            print(f"\n⚠️ File not found: {file_path}")
            continue
        
        try:
            ext = os.path.splitext(file_path)[1].lower()
            
            if ext in ['.docx', '.doc']:
                result = convert_docx_to_md(file_path, output_name)
                converted_files.append(result)
            elif ext == '.pdf':
                result = convert_pdf_to_md(file_path, output_name)
                converted_files.append(result)
            else:
                print(f"\n⚠️ Unsupported format: {ext}")
        except Exception as e:
            print(f"\n❌ Error converting {file_path}: {e}")
            import traceback
            traceback.print_exc()
    
    print("\n" + "=" * 60)
    print("✅ Conversion Complete!")
    print(f"📁 Output directory: {OUTPUT_DIR}")
    print(f"📁 Images directory: {IMAGES_DIR}")
    print(f"📊 Total files converted: {len(converted_files)}")
    print("=" * 60)

if __name__ == "__main__":
    main()
