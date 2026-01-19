# -*- coding: utf-8 -*-
"""
IEEE 템플릿 DOCX 파일을 Markdown으로 변환
"""

import os
import re
from docx import Document

BASE_DIR = r"c:\1.이규영개인폴더\09.##### SCHOOL #####"
OUTPUT_DIR = os.path.join(BASE_DIR, "converted_md")
IMAGES_DIR = os.path.join(OUTPUT_DIR, "images")

def sanitize_filename(filename):
    """파일명에서 특수문자 제거"""
    return re.sub(r'[<>:"/\\|?*]', '_', filename)

def extract_table_as_markdown(table):
    """docx 테이블을 마크다운 테이블로 변환"""
    md_table = []
    
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        md_table.append("| " + " | ".join(cells) + " |")
        
        # 헤더 구분선 추가
        if i == 0:
            md_table.append("|" + "|".join(["---" for _ in cells]) + "|")
    
    return "\n".join(md_table)

def convert_docx_to_md(docx_path, output_name):
    """DOCX 파일을 Markdown으로 변환"""
    print(f"\n📄 Converting DOCX: {docx_path}")
    
    doc = Document(docx_path)
    md_content = []
    image_count = 0
    
    # 문서 제목
    base_name = os.path.splitext(output_name)[0]
    md_content.append(f"# {base_name}\n")
    md_content.append(f"*원본 파일: {os.path.basename(docx_path)}*\n")
    md_content.append("---\n")
    
    # 이미지 추출 (embedded images)
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            try:
                image_data = rel.target_part.blob
                image_ext = rel.target_ref.split('.')[-1]
                image_filename = f"{sanitize_filename(base_name)}_img_{image_count}.{image_ext}"
                image_path = os.path.join(IMAGES_DIR, image_filename)
                
                with open(image_path, 'wb') as img_file:
                    img_file.write(image_data)
                
                print(f"  ✅ Extracted image: {image_filename}")
                image_count += 1
            except Exception as e:
                print(f"  ⚠️ Image extraction error: {e}")
    
    # 문서 내용 처리
    for element in doc.element.body:
        # 테이블 처리
        if element.tag.endswith('tbl'):
            for table in doc.tables:
                if table._tbl == element:
                    md_content.append("\n" + extract_table_as_markdown(table) + "\n")
                    break
        # 단락 처리
        elif element.tag.endswith('p'):
            for para in doc.paragraphs:
                if para._p == element:
                    text = para.text.strip()
                    if text:
                        # 스타일 기반 헤딩 처리
                        style = para.style.name if para.style else ""
                        if "Heading 1" in style or "heading 1" in style.lower():
                            md_content.append(f"\n## {text}\n")
                        elif "Heading 2" in style or "heading 2" in style.lower():
                            md_content.append(f"\n### {text}\n")
                        elif "Heading 3" in style or "heading 3" in style.lower():
                            md_content.append(f"\n#### {text}\n")
                        elif "Title" in style:
                            md_content.append(f"\n# {text}\n")
                        else:
                            md_content.append(f"{text}\n")
                    break
    
    # 추출된 이미지 참조 추가
    if image_count > 0:
        md_content.append("\n---\n## 추출된 이미지\n")
        for i in range(image_count):
            image_ext = "png"  # 기본값
            for ext in ['png', 'jpg', 'jpeg', 'gif', 'emf', 'wmf']:
                if os.path.exists(os.path.join(IMAGES_DIR, f"{sanitize_filename(base_name)}_img_{i}.{ext}")):
                    image_ext = ext
                    break
            md_content.append(f"\n![Image {i+1}](images/{sanitize_filename(base_name)}_img_{i}.{image_ext})\n")
    
    # MD 파일 저장
    output_path = os.path.join(OUTPUT_DIR, f"{sanitize_filename(base_name)}.md")
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("\n".join(md_content))
    
    print(f"  ✅ Saved: {output_path}")
    print(f"  📊 Images extracted: {image_count}")
    return output_path

def main():
    """메인 실행 함수"""
    print("=" * 60)
    print("📁 IEEE Template DOCX to Markdown Converter")
    print("=" * 60)
    
    # IEEE 템플릿 파일 경로
    docx_file = os.path.join(BASE_DIR, "IEEE_TII_Template", "TII_Articles_Word_template_2025.docx")
    
    if os.path.exists(docx_file):
        convert_docx_to_md(docx_file, "TII_Articles_Word_template_2025")
        print("\n✅ Conversion complete!")
    else:
        print(f"⚠️ File not found: {docx_file}")

if __name__ == "__main__":
    main()
