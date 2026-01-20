import re
import os
import io
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.styles.style import _ParagraphStyle

# --- Configuration ---
MD_FILE = "04_IEEE_DPF_Paper_Final_Extended.md"
OUTPUT_FILE = "IEEE_DPF_Paper_Final_Submission.docx"
IMAGE_DIR = "images"

# --- OXML Helper Functions for Columns ---
def set_columns(section, cols, space_twips=425):
    """
    Sets the number of columns for a section.
    space_twips: 720 twips = 0.5 inch. 425 twips ~= 0.3 inch (IEEE standard gap)
    """
    sectPr = section._sectPr
    cols_el = sectPr.xpath('./w:cols')
    if cols_el:
        cols_el = cols_el[0]
    else:
        cols_el = OxmlElement('w:cols')
        sectPr.append(cols_el)
    
    cols_el.set(qn('w:num'), str(cols))
    cols_el.set(qn('w:space'), str(space_twips)) # 425 twips = ~0.3 inch

# --- Styles Setup ---
def setup_ieee_styles(doc):
    # Base Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    
    # Title Style
    if 'IEEE Title' not in doc.styles:
        title_style = doc.styles.add_style('IEEE Title', 1) # 1 = Paragraph Style
        title_style.font.name = 'Times New Roman'
        title_style.font.size = Pt(24)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)

    # Author Style
    if 'IEEE Author' not in doc.styles:
        author_style = doc.styles.add_style('IEEE Author', 1)
        author_style.font.name = 'Times New Roman'
        author_style.font.size = Pt(11)
        author_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_style.paragraph_format.space_after = Pt(24)

    # Heading 1 (Section)
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(12) # Slightly larger or SC
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)

    # Heading 2 (Subsection)
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(10)
    h2.font.bold = False
    h2.font.italic = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(4)

    # Caption Style
    if 'IEEE Caption' not in doc.styles:
        caption_style = doc.styles.add_style('IEEE Caption', 1)
        caption_style.font.name = 'Times New Roman'
        caption_style.font.size = Pt(8)
        caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_style.paragraph_format.space_after = Pt(12)

# --- LaTeX Rendering ---
def render_latex(formula, fontsize=12):
    """
    Renders a LaTeX formula to a PNG image in memory using matplotlib.
    """
    # Configure Matplotlib for MathText (closer to Times New Roman)
    plt.rcParams['mathtext.fontset'] = 'stix'
    plt.rcParams['font.family'] = 'STIXGeneral'
    
    fig = plt.figure(figsize=(0.1, 0.1)) # Dummy size
    # Use raw string concatenation to ensure backslashes are preserved
    text = "$" + formula + "$"
    fig.text(0, 0, text, fontsize=fontsize)
    
    buf = io.BytesIO()
    plt.axis('off')
    try:
        plt.savefig(buf, format='png', bbox_inches='tight', pad_inches=0.1, dpi=300, transparent=True)
    except Exception as e:
        print(f"Error rendering latex [{formula}]: {str(e)}")
        # Fallback: try removing potentially problematic commands if needed, or just return None
        return None
    plt.close(fig)
    buf.seek(0)
    return buf

# --- Markdown Parsing ---
def parse_markdown(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    blocks = []
    current_block = {"type": "text", "content": []}
    
    in_code_block = False
    
    for line in lines:
        line = line.strip()
        
        # Heading
        if line.startswith("#"):
            if current_block["content"]:
                blocks.append(current_block)
            level = len(line.split()[0])
            text = line.lstrip("#").strip()
            # Remove ** ** from heading
            text = text.replace("**", "")
            blocks.append({"type": "heading", "level": level, "content": text})
            current_block = {"type": "text", "content": []}
        
        # Image
        elif line.startswith("!"):
            if current_block["content"]:
                blocks.append(current_block)
            # ![Alt](path)
            match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
            if match:
                caption = match.group(1)
                path = match.group(2)
                blocks.append({"type": "image", "path": path, "caption": caption})
            current_block = {"type": "text", "content": []}
        
        # Latex Equation ($$ ... $$) - Single line assumption for now
        elif line.startswith("$$") and line.endswith("$$"):
            if current_block["content"]:
                blocks.append(current_block)
            latex = line.replace("$$", "").strip()
            blocks.append({"type": "equation", "content": latex})
            current_block = {"type": "text", "content": []}
        
        # Table (Basic check)
        elif line.startswith("|"):
            if current_block["type"] != "table":
                if current_block["content"]:
                    blocks.append(current_block)
                current_block = {"type": "table", "content": [line]}
            else:
                current_block["content"].append(line)
        
        # Horizontal Rule
        elif line.startswith("---"):
            continue # Skip separators
            
        else:
            if current_block["type"] == "table":
                 blocks.append(current_block)
                 current_block = {"type": "text", "content": []}
            
            if line:
                current_block["content"].append(line)
            # Empty lines in text generally mean paragraph breaks, 
            # but we'll handle them by appending empty strings to preserve spacing or ignore.
            # Here we just append.
    
    if current_block["content"]:
        blocks.append(current_block)
        
    return blocks

# --- Main Conversion Logic ---
def create_document():
    doc = Document()
    setup_ieee_styles(doc)
    
    blocks = parse_markdown(MD_FILE)
    
    # --- 1. Title & Abstract (One Column) ---
    title_text = ""
    author_info = []
    abstract_text = ""
    start_body_index = 0
    
    # Naive extraction of Title/Author from top blocks
    # Looking for Metadata at the top
    for i, block in enumerate(blocks):
        if block['type'] == 'heading' and block['level'] == 1:
            title_text = block['content']
        elif block['type'] == 'text':
            for line in block['content']:
                if line.startswith("**저자:**"):
                    author_info.append(line.replace("**저자:**", "").strip())
                elif line.startswith("**소속:**"):
                    author_info.append(line.replace("**소속:**", "").strip())
                elif line.startswith("**작성일:**"):
                    pass # Skip date
                elif line.startswith("**핵심어:**"):
                    pass # Will handle with abstract
        elif block['type'] == 'heading' and '요약' in block['content']: # Abstract Header
            # Next block is likely abstract
             pass
        elif block['type'] == 'text' and any('인더스트리 5.0' in l for l in block['content']):
             # This is abstract content
             abstract_text = " ".join(block['content'])
        
        if block['type'] == 'heading' and 'I. 서론' in block['content']:
            start_body_index = i
            break
            
    # Write Title
    doc.add_paragraph(title_text, style='IEEE Title')
    
    # Write Author
    if author_info:
        doc.add_paragraph("\n".join(author_info), style='IEEE Author')
        
    # Write Abstract
    p_abs = doc.add_paragraph()
    p_abs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Bold Abstract
    runner = p_abs.add_run("Abstract—")
    runner.bold = True
    runner.font.name = 'Times New Roman'
    p_abs.add_run(abstract_text.replace("**", "")) # Remove markdown bold
    
    # Add Keywords if found
    # (Simplified: hardcoded or scanned)
    
    doc.add_paragraph() # Spacer
    
    # --- 2. Body (Two Columns) ---
    section = doc.add_section(WD_SECTION.CONTINUOUS)
    set_columns(section, 2)
    
    # Process blocks from start_body_index
    for block in blocks[start_body_index:]:
        if block['type'] == 'heading':
            level = block['level']
            text = block['content']
            style = 'Heading 1' if level == 2 else 'Heading 2'
            doc.add_paragraph(text, style=style)
            
        elif block['type'] == 'text':
            full_text = " ".join(block['content'])
            if not full_text.strip(): continue
            
            # Simple bold/italic markdown parsing could be added here
            full_text = full_text.replace("**", "") # Strip bold for now
            
            p = doc.add_paragraph(full_text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Pt(10) # Paragraph indent
            
        elif block['type'] == 'equation':
            latex = block['content']
            img_buf = render_latex(latex)
            if img_buf:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_buf, height=Cm(1.5)) # Approximate height
                # Add equation number manually if needed
                
        elif block['type'] == 'image':
            path = block['path']
            caption = block['caption']
            
            # Fix path if needed (remove 'images/' prefix if it's duplicated or check existence)
            real_path = path
            if not os.path.exists(path):
                # Try finding it in current dir or subdirs
                basename = os.path.basename(path)
                real_path = os.path.join(IMAGE_DIR, basename)
            
            if os.path.exists(real_path):
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    # Fit to column width (approx 3.4 inches)
                    run.add_picture(real_path, width=Inches(3.4)) 
                    
                    # Caption
                    c_p = doc.add_paragraph(f"{caption}", style='IEEE Caption')
                except Exception as e:
                    print(f"Error adding image {real_path}: {e}")
            else:
                doc.add_paragraph(f"[Image Missing: {path}]", style='IEEE Caption')

        elif block['type'] == 'table':
            # Basic table
            rows = block['content']
            # Determine cols
            header = rows[0]
            cols_count = len(header.split('|')) - 2
            if cols_count < 1: cols_count = 1
            
            table = doc.add_table(rows=0, cols=cols_count)
            table.style = 'Table Grid'
            
            for row_md in rows:
                if '---' in row_md: continue
                cells_md = row_md.strip('|').split('|')
                row_cells = table.add_row().cells
                for i, cell_md in enumerate(cells_md):
                    if i < len(row_cells):
                        row_cells[i].text = cell_md.strip().replace("**", "") # Clean bold

    # Save
    doc.save(OUTPUT_FILE)
    print(f"Successfully generated {OUTPUT_FILE}")

if __name__ == "__main__":
    create_document()
